from docx import Document
import os

def inspect_templates():
    tdir = 'data/templates'
    if not os.path.exists(tdir):
        print(f"Directory {tdir} not found.")
        return
    
    for fn in os.listdir(tdir):
        if fn.endswith('.docx'):
            print(f"\n--- {fn} ---")
            try:
                doc = Document(os.path.join(tdir, fn))
                for i, p in enumerate(doc.paragraphs):
                    if p.text.strip():
                        print(f"P{i}: {p.text}")
                for i, table in enumerate(doc.tables):
                    print(f"Table {i}:")
                    for row in table.rows:
                        print(" | ".join([cell.text.strip() for cell in row.cells]))
            except Exception as e:
                print(f"Error reading {fn}: {e}")

if __name__ == "__main__":
    inspect_templates()
