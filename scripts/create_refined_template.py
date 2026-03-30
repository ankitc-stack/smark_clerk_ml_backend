from docx import Document
from docx.shared import Pt
import os

def create_refined_template():
    doc = Document()
    
    # Header
    p = doc.add_paragraph("Gen Staff Branch")
    p.alignment = 1 # Center
    p = doc.add_paragraph("Integrated HQ of MoD (Army)")
    p.alignment = 1
    p = doc.add_paragraph("New Delhi-01")
    p.alignment = 1
    
    doc.add_paragraph()
    
    # Station and Date
    doc.add_paragraph("Station : {{STATION|default('New Delhi')}}")
    doc.add_paragraph("Dated   : {{DATE}}")
    
    doc.add_paragraph()
    
    # Subject
    p = doc.add_paragraph()
    p.add_run("Subject : ").bold = True
    p.add_run("{{SUBJECT}}")
    
    doc.add_paragraph()
    
    # Body
    doc.add_paragraph("{{BODY}}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature Zone
    p = doc.add_paragraph("({{SIGN_NAME}})")
    p.alignment = 2 # Right
    p = doc.add_paragraph("{{DESIGNATION}}")
    p.alignment = 2
    p = doc.add_paragraph("{{RANK}}")
    p.alignment = 2

    # Distribution
    doc.add_paragraph()
    p = doc.add_paragraph("Copy to:-")
    p.add_run().bold = True
    doc.add_paragraph("1.  {{DISTR_1|default('Indl Concerned')}}")
    doc.add_paragraph("2.  {{DISTR_2|default('Office Copy')}}")

    out_path = "data/templates/Refined_Leave_Template.docx"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Created refined template: {out_path}")

if __name__ == "__main__":
    create_refined_template()
