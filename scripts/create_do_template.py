from docx import Document
from docx.shared import Pt
import os

def create_do_template():
    doc = Document()
    
    # 1. Header (Sender)
    doc.add_paragraph("{{SENDER_RANK}} {{SENDER_NAME}}")
    doc.add_paragraph("{{SENDER_APPOINTMENT}}")
    doc.add_paragraph("{{SENDER_ORGANISATION}}")
    doc.add_paragraph("Tele : {{TELE}} Fax : {{FAX}}")
    doc.add_paragraph("E-mail : {{EMAIL}}")
    doc.add_paragraph("{{SENDER_ADDR_1}}")
    doc.add_paragraph("{{SENDER_ADDR_2}}")
    
    doc.add_paragraph() # Spacer
    
    # 2. DO No and Date
    doc.add_paragraph("DO No: {{DO_NO}} Date: {{DATE}}")
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer (Extra as requested <br>)
    
    # 3. Recipient
    doc.add_paragraph("{{RECIPIENT_RANK}} {{RECIPIENT_NAME}}")
    doc.add_paragraph("{{RECIPIENT_DESIGNATION}}")
    doc.add_paragraph("{{RECIPIENT_ORGANISATION}}")
    doc.add_paragraph("{{RECIPIENT_ADDRESS}}")
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer (Extra as requested <br>)
    
    # 4. Subject
    p_subj = doc.add_paragraph()
    p_subj.add_run("{{SUBJECT_UPPER}}").bold = True
    p_subj.alignment = 1 # Center
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer (Extra as requested <br>)
    
    # 5. Salutation
    doc.add_paragraph("My dear {{RECIPIENT_SURNAME}},")
    
    doc.add_paragraph() # Spacer
    
    # 6. Body
    doc.add_paragraph("{{BODY}}")
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer (Extra as requested <br>)
    
    # 7. Closing
    p_closing1 = doc.add_paragraph("Warm regards,")
    p_closing1.alignment = 0 # Left
    p_closing2 = doc.add_paragraph("Yours sincerely,")
    p_closing2.alignment = 0 # Left
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer (Extra as requested <br>)
    
    # 8. Signature
    p_sig = doc.add_paragraph("(Signature)")
    p_sig.alignment = 0 # Left

    out_path = "data/templates/Refined_DO_Template.docx"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Template recreated: {out_path}")

if __name__ == "__main__":
    create_do_template()
