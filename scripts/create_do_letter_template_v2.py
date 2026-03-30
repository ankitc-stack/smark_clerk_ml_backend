#!/usr/bin/env python3
"""
scripts/create_do_letter_template_v2.py

Creates DO_Letter_Template_docxtpl_v2.docx — a clean Demi-Official (DO) letter template
using only docxtpl Jinja2 {{ }} variables (no [[BLOCK]] markers).

Template variables used:
    {{ file_reference_number }}  — reference number (top-left)
    {{ date }}                   — date (top-right)
    {{ addressee_1 }}            — recipient name/rank (bold)
    {{ addressee_2 }}            — recipient designation/dept (bold, optional)
    {{ subject }}                — subject line (centered, bold, uppercase)
    {{ salutation }}             — e.g. "My dear Sir,"
    {% for p in body_paras %}    — numbered body paragraphs
    {{ signatory_name }}         — signer name
    {{ signatory_designation }}  — signer rank/title
    {{ signatory_dept }}         — signer department/unit (optional)
    {% for c in copy_to_list %}  — copy-to distribution (optional)

Run once:
    python scripts/create_do_letter_template_v2.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _run(para, text: str, bold=False, underline=False, italic=False, size_pt: int = 12):
    r = para.add_run(text)
    r.bold = bold
    r.underline = underline
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size_pt)
    return r


def _para(doc, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6) -> "Paragraph":
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def create_template(out_path: str):
    doc = Document()

    # Page layout
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ── Header row: Reference No (left) and Date (right) ─────────────────────
    # Use a table for side-by-side layout
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    # Remove borders via XML
    for cell in tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "nil")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(lp, "{{ file_reference_number }}", bold=False)

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rp, "{{ date }}", bold=False)

    # ── Spacer ────────────────────────────────────────────────────────────────
    _para(doc, space_after=4)

    # ── Addressee block ───────────────────────────────────────────────────────
    p = _para(doc, space_after=2)
    _run(p, "{{ addressee_1 }}", bold=True)

    p = _para(doc, space_after=12)
    _run(p, "{% if addressee_2 %}{{ addressee_2 }}{% endif %}", bold=True)

    # ── Subject ───────────────────────────────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=12)
    _run(p, "{{ subject }}", bold=True, underline=True)

    # ── Salutation ────────────────────────────────────────────────────────────
    p = _para(doc, space_after=8)
    _run(p, "{{ salutation }}", bold=False)

    # ── Body paragraphs (Jinja2 for loop) ─────────────────────────────────────
    # docxtpl evaluates Jinja2 across runs, so we put the entire loop block
    # in a single paragraph to avoid XML splitting issues.
    p = _para(doc, space_after=6)
    _run(
        p,
        "{% for p in body_paras %}{{ loop.index }}.\t{{ p }}\n{% endfor %}",
        bold=False,
    )

    # ── Closing ───────────────────────────────────────────────────────────────
    _para(doc, space_after=2)
    p = _para(doc, space_after=2)
    _run(p, "Yours sincerely,")

    # ── Spacer for signature ──────────────────────────────────────────────────
    for _ in range(3):
        _para(doc, space_after=0)

    # ── Signatory block (right-aligned) ───────────────────────────────────────
    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    _run(p, "{{ signatory_name }}", bold=True)

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    _run(p, "{{ signatory_designation }}")

    p = _para(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)
    _run(p, "{% if signatory_dept %}{{ signatory_dept }}{% endif %}")

    # ── Copy-to block (optional) ──────────────────────────────────────────────
    p = _para(doc, space_after=4)
    _run(p, "{% if copy_to_list %}Copy to:-{% endif %}")

    p = _para(doc, space_after=4)
    _run(p, "{% for c in copy_to_list %}{{ loop.index }}.\t{{ c }}\n{% endfor %}")

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    repo = Path(__file__).resolve().parents[1]
    out = repo / "data" / "templates" / "DO_Letter_Template_docxtpl_v2.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    create_template(str(out))
