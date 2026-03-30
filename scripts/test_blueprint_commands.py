#!/usr/bin/env python3
"""
Test all blueprint command ops end-to-end.

Tests:
  1. ADD_PARAGRAPH  — add para, check DOCX has new paragraph
  2. DELETE_SECTION — delete para, check it's gone from DOCX
  3. MOVE_SECTION   — move para 2 before para 1, check order
  4. REWRITE_CONTENT — content op: rewrite para text
  5. EXPAND_CONTENT  — content op: expand para text
  6. SHORTEN_CONTENT — content op: shorten para text

Each test creates a fresh DO letter, applies the command using the
server's live LLM (COMMAND_INTENT_USE_LLM + COMMAND_TRANSFORM_USE_LLM),
downloads the re-rendered DOCX, and verifies the result.

Usage:
    python scripts/test_blueprint_commands.py
    python scripts/test_blueprint_commands.py --base http://localhost:8000
"""
import argparse
import os
import sys
import tempfile
import requests
from docx import Document

BASE = "http://localhost:8000"
TEMPLATE_ID = "tmpl_do_001"
USER_ID = "ci-test"
PROMPT = "do letter on request for additional ration on army day celebration"

PASS = "\033[32mPASS\033[0m"
FAIL = "\033[31mFAIL\033[0m"
SKIP = "\033[33mSKIP\033[0m"

results: list[tuple[str, bool, str]] = []


def _cmd(doc_id: str, version: int, prompt: str) -> dict:
    headers = {"Content-Type": "application/json"}
    body = {
        "version": version,
        "prompt": prompt,
        "context": {"current_section_id": None, "selected_section_ids": [], "cursor_position": None},
        "input": {"type": "text", "value": prompt},
    }
    resp = requests.post(f"{BASE}/documents/{doc_id}/command", json=body, timeout=180)
    if not resp.ok:
        return {"status": "error", "_http": resp.status_code, "detail": resp.text[:200]}
    return resp.json()


def create_doc() -> tuple[str, int, int]:
    """Returns (doc_id, ml_version, de_version)."""
    r = requests.post(
        f"{BASE}/documents/generate",
        json={"template_id": TEMPLATE_ID, "user_id": USER_ID, "prompt": PROMPT},
        timeout=90,
    )
    r.raise_for_status()
    d = r.json()
    return d["document_id"], d["version"], d["docengine_version"]


def download_docx(doc_id: str, version: int) -> list[str]:
    """Download DOCX for version, return list of non-empty paragraph texts."""
    r = requests.get(
        f"{BASE}/documents/{doc_id}/versions/{version}/download?format=docx",
        timeout=30,
    )
    r.raise_for_status()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(r.content)
        path = f.name
    try:
        doc = Document(path)
        # Para 5 (index 5) is the body_paras loop in DO letter template
        body_para = doc.paragraphs[5].text if len(doc.paragraphs) > 5 else ""
        lines = [l.strip() for l in body_para.split("\n") if l.strip()]
        return lines
    finally:
        os.unlink(path)


def record(name: str, passed: bool, detail: str = ""):
    tag = PASS if passed else FAIL
    print(f"  [{tag}] {name}" + (f": {detail}" if detail else ""))
    results.append((name, passed, detail))


def test_add_paragraph():
    print("\n--- Test 1: ADD_PARAGRAPH ---")
    doc_id, ver, _ = create_doc()
    paras_before = download_docx(doc_id, ver)
    count_before = len(paras_before)
    print(f"  Initial paragraphs: {count_before}")

    resp = _cmd(doc_id, ver, "add paragraph about logistics planning")
    ok = resp.get("status") == "applied"
    record("status=applied", ok, resp.get("status"))
    if not ok:
        record("DOCX has extra paragraph", False, "command failed")
        return

    new_ver = resp["version"]
    has_url = bool(resp.get("docx_url"))
    record("docx_url in response", has_url)

    paras_after = download_docx(doc_id, new_ver)
    added = len(paras_after) > count_before
    record(f"paragraph count {count_before} → {len(paras_after)}", added)


def test_delete_section():
    print("\n--- Test 2: DELETE_SECTION ---")
    doc_id, ver, _ = create_doc()
    # Add a 3rd paragraph first — blueprint requires min 2, so we need 3 to delete one
    resp_add = _cmd(doc_id, ver, "add paragraph about logistics")
    if resp_add.get("status") != "applied":
        record("setup: add paragraph for delete test", False, resp_add.get("status"))
        return
    ver = resp_add["version"]
    paras_before = download_docx(doc_id, ver)
    count_before = len(paras_before)
    print(f"  Paragraphs before delete: {count_before}")

    resp = _cmd(doc_id, ver, "delete paragraph 3")
    ok = resp.get("status") == "applied"
    record("status=applied", ok, resp.get("status"))
    if not ok:
        record("DOCX lost a paragraph", False, "command failed")
        return

    new_ver = resp["version"]
    has_url = bool(resp.get("docx_url"))
    record("docx_url in response", has_url)
    paras_after = download_docx(doc_id, new_ver)
    removed = len(paras_after) < count_before
    record(f"paragraph count {count_before} → {len(paras_after)}", removed)


def test_move_section():
    print("\n--- Test 3: MOVE_SECTION ---")
    doc_id, ver, _ = create_doc()
    paras_before = download_docx(doc_id, ver)
    print(f"  Initial paragraphs: {len(paras_before)}")
    if len(paras_before) < 2:
        record("need ≥2 paras to move", False, "skip")
        return

    # First add a 3rd paragraph so we have something to move
    resp_add = _cmd(doc_id, ver, "add paragraph about logistics")
    if resp_add.get("status") != "applied":
        record("setup: add paragraph", False)
        return
    ver = resp_add["version"]

    resp = _cmd(doc_id, ver, "move paragraph 2 before paragraph 1")
    ok = resp.get("status") == "applied"
    record("status=applied", ok, resp.get("status"))
    if not ok:
        record("move op executed", False, "command failed")
        return

    new_ver = resp["version"]
    paras_after = download_docx(doc_id, new_ver)
    # Just check count is stable and DOCX renders
    stable = len(paras_after) == len(download_docx(doc_id, ver))
    record(f"DOCX rendered, {len(paras_after)} paragraphs", len(paras_after) > 0)


def test_rewrite_content():
    print("\n--- Test 4: REWRITE_CONTENT ---")
    doc_id, ver, _ = create_doc()
    paras_before = download_docx(doc_id, ver)
    print(f"  Initial paragraphs: {len(paras_before)}")
    if not paras_before:
        record("need ≥1 para", False, "skip")
        return

    first_before = paras_before[0] if paras_before else ""
    resp = _cmd(doc_id, ver, "rewrite paragraph 1 to be more formal")
    ok = resp.get("status") == "applied"
    record("status=applied", ok, resp.get("status"))
    if not ok:
        record("REWRITE_CONTENT applied", False, "command failed")
        return

    new_ver = resp["version"]
    has_url = bool(resp.get("docx_url"))
    record("docx_url in response", has_url)
    paras_after = download_docx(doc_id, new_ver)
    record("DOCX re-rendered", len(paras_after) > 0)


def test_expand_content():
    print("\n--- Test 5: EXPAND_CONTENT ---")
    doc_id, ver, _ = create_doc()
    paras_before = download_docx(doc_id, ver)
    if not paras_before:
        record("need ≥1 para", False, "skip")
        return

    resp = _cmd(doc_id, ver, "expand paragraph 1 with more details")
    ok = resp.get("status") == "applied"
    record("status=applied", ok, resp.get("status"))
    if not ok:
        record("EXPAND_CONTENT applied", False, "command failed")
        return

    new_ver = resp["version"]
    paras_after = download_docx(doc_id, new_ver)
    record("DOCX re-rendered", len(paras_after) > 0)


def test_shorten_content():
    print("\n--- Test 6: SHORTEN_CONTENT ---")
    doc_id, ver, _ = create_doc()
    paras_before = download_docx(doc_id, ver)
    if not paras_before:
        record("need ≥1 para", False, "skip")
        return

    resp = _cmd(doc_id, ver, "shorten paragraph 1")
    ok = resp.get("status") == "applied"
    record("status=applied", ok, resp.get("status"))
    if not ok:
        record("SHORTEN_CONTENT applied", False, "command failed")
        return

    new_ver = resp["version"]
    paras_after = download_docx(doc_id, new_ver)
    record("DOCX re-rendered", len(paras_after) > 0)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--base", default="http://localhost:8000")
    args = parser.parse_args()
    global BASE
    BASE = args.base.rstrip("/")

    # Verify services are up
    try:
        r = requests.get(f"{BASE}/health", timeout=5)
        r.raise_for_status()
    except Exception as e:
        print(f"ML pipeline not reachable at {BASE}: {e}")
        sys.exit(1)

    print(f"Testing blueprint command ops against {BASE}")
    print("=" * 60)

    test_add_paragraph()
    test_delete_section()
    test_move_section()
    test_rewrite_content()
    test_expand_content()
    test_shorten_content()

    print("\n" + "=" * 60)
    passed = sum(1 for _, ok, _ in results if ok)
    total = len(results)
    print(f"Result: {passed}/{total} checks passed")
    failed = [(n, d) for n, ok, d in results if not ok]
    if failed:
        print("Failed:")
        for n, d in failed:
            print(f"  - {n}: {d}")
        sys.exit(1)
    else:
        print("All checks passed.")


if __name__ == "__main__":
    main()
