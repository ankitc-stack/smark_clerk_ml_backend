#!/usr/bin/env python3
"""
scripts/create_movement_order_template_v2.py

Creates Movement_Order_Template_docxtpl_v2.docx that exactly matches the
sample Indian Army movement order (IAFT-1759 format) with docxtpl Jinja2 markers.

Run once:
    python scripts/create_movement_order_template_v2.py
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_run(para, text: str, bold=False, underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.underline = underline
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return r


def add_numbered_para(doc, number: int, runs: list[tuple]):
    """
    Adds a numbered paragraph.
    runs: list of (text, bold, underline)
    Uses hanging-indent style: number + tab then content.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"{number}.\t", False, False)
    for text, bold, underline in runs:
        add_run(p, text, bold, underline)
    return p


def create_template(out_path: str):
    doc = Document()

    # Page layout
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

    # Default paragraph spacing
    from docx.styles.style import _ParagraphStyle  # noqa
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # ── Top-right: "In lieu of IAFT-1759" ────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "In lieu of IAFT-1759", bold=True, underline=True)

    # ── Title: MOVEMENT ORDER ─────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "MOVEMENT ORDER", bold=True, underline=True)

    # ── Warning motto (bold, underlined, centered) ────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    add_run(
        p,
        "Do not indulge in any such activities during your journey which disrespect your "
        "country and dignity: Maintain the dignity of the Army and yourself and do not take "
        "food/drinks from any unknown person.",
        bold=True,
        underline=True,
    )

    # ── Para 1: Person details + destination ──────────────────────────────
    # Full Jinja2 block in single run to avoid XML splitting issues.
    # Bold/underline on army_no, rank, person_name, destination.
    add_numbered_para(doc, 1, [
        ("No ", False, False),
        ("{{ army_no }}", True, True),
        (" Rank ", False, False),
        ("{{ rank }}", True, True),
        (" Name ", False, False),
        ("{{ person_name }}", True, True),
        # att_unit conditional kept in one run
        (" of {{ unit }}{% if att_unit %} att with {{ att_unit }}{% endif %},"
         " will be proceeding on temp duty to ", False, False),
        ("{{ destination }}", True, True),
        (".", False, False),
    ])

    # ── Para 2: Departure date and time ───────────────────────────────────
    add_numbered_para(doc, 2, [
        ("Departure date and time\t: ", False, False),
        ("{{ departure_date }}{% if departure_time %} ({{ departure_time }}){% endif %}", True, True),
        (".", False, False),
    ])

    # ── Para 3: Route / Destination ───────────────────────────────────────
    add_numbered_para(doc, 3, [
        ("Route/Via\t", False, False),
        ("{{ route }}", True, True),
        ("\tDestination\t: ", False, False),
        ("{{ destination_desc }}", True, True),
        (".", False, False),
    ])

    # ── Para 4: Auth (fixed) ──────────────────────────────────────────────
    add_numbered_para(doc, 4, [
        ("Auth\t: This move order.", False, False),
    ])

    # ── Para 5: Remarks ───────────────────────────────────────────────────
    add_numbered_para(doc, 5, [
        ("Remarks\t: {{ remarks }}", False, False),
        (".", False, False),
    ])

    # ── Para 6: On arrival ────────────────────────────────────────────────
    add_numbered_para(doc, 6, [
        ("On arrival at (Station) ", False, False),
        ("{{ destination }}", True, True),
        (" they will report to concerned auth.", False, False),
    ])

    # ── Para 7: SOS / SORS ───────────────────────────────────────────────
    add_numbered_para(doc, 7, [
        ("SOS _______ SORS wef {{ departure_date }} ration up to and for {{ departure_date }}.",
         False, False),
    ])

    # ── Para 8–10: Fixed content ──────────────────────────────────────────
    add_numbered_para(doc, 8, [
        ("He is in possession of his I Card.", False, False),
    ])

    add_numbered_para(doc, 9, [
        ("Provisions of Aos 782/64 and 504/67 have been explained, "
         "Anti Malaria precaution will be strictly observed.", False, False),
    ])

    add_numbered_para(doc, 10, [
        ("He will not be discussing military matter with unauthorized persons.", False, False),
    ])

    # Spacing before distribution
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # ── Distribution ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    add_run(p, "Distr\t:", bold=True, underline=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "1.\tIndl Concerned.", False, False)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "2.\t{{ distr_unit }}", False, False)

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    repo = Path(__file__).resolve().parents[1]
    out = repo / "data" / "templates" / "Movement_Order_Template_docxtpl_v2.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    create_template(str(out))
