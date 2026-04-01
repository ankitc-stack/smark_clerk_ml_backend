from __future__ import annotations
import base64
import binascii
import json
import logging
import os
import re
import time
import uuid
from pathlib import Path
from fastapi import FastAPI, Depends, UploadFile, File, Form, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.orm import Session

from app.db import Base, engine, get_db
from app.models import Template, Document, DocumentVersion, User
from app.auth import (
    create_user as auth_create_user,
    authenticate_user as auth_authenticate_user,
    create_access_token,
    get_current_user,
)
from app.schemas import (
    TemplateRegisterResponse,
    TemplateOut,
    UploadResponse,
    VersionHistoryItem,
    DocumentSummaryOut,
    DocumentListItem,
    CommandRequest,
    CommandAppliedResponse,
    CommandNeedsClarificationResponse,
    CommandErrorResponse,
    RevertDocumentRequest,
    RevertDocumentResponse,
    BlueprintDocRequest,
    BlueprintDocResponse,
    SectionPatchRequest,
    SectionPatchResponse,
    SaveAsTemplateRequest,
    SaveAsTemplateResponse,
    TemplateStoreItem,
    FeedbackRequest,
    FeedbackResponse,
    RegisterRequest,
    LoginRequest,
    TokenResponse,
    MeResponse,
)
from app.services.docengine_client import DocEngineClient, DocEngineError
from app.services.action_bridge import ml_action_to_de, BYPASS_ACTIONS, CONTENT_OPS
from app.services.zones import suggest_zones
from app.services.rag import search_rules
from app.services.render import render_docx
from app.services.export_pdf import docx_to_pdf
from app import crud
from app.config import settings
from app.services.bootstrap import bootstrap
from app.services.intent_router import route_intent
from app.services.render_adapter import doc_state_from_filled_skeleton
from app.services.patch_ops import apply_patch_ops
from app.services.command_contract import (
    ACTION_OBJECT_JSON_SCHEMA,
    IntentParseError,
    NeedsClarificationError,
    PlannerError,
    build_clarification_token,
    plan_patch_ops_from_action,
)
from app.services.intent_extractor import extract_action_object_with_meta
from app.services.stt import (
    STTError,
    transcribe,
)

Base.metadata.create_all(bind=engine)

app = FastAPI(title="Army Smart Clerk Backend (Starter)")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/ui", include_in_schema=False)
def serve_ui():
    import pathlib
    html = pathlib.Path(__file__).parent.parent / "smart_clerk_test_ui.html"
    return FileResponse(str(html), media_type="text/html")

# Structured command logs are emitted as single JSON lines for easy parsing in ops pipelines.
COMMAND_LOGGER = logging.getLogger("app.command")
# Ensure command telemetry is emitted even when root logging is not explicitly configured (for local/dev runs).
if not COMMAND_LOGGER.handlers:
    _command_log_handler = logging.StreamHandler()
    _command_log_handler.setFormatter(logging.Formatter("%(message)s"))
    COMMAND_LOGGER.addHandler(_command_log_handler)
COMMAND_LOGGER.setLevel(logging.INFO)
COMMAND_LOGGER.propagate = False

@app.on_event("startup")
async def startup_bootstrap():
    if str(getattr(settings, "AUTO_BOOTSTRAP", "false")).lower() in ("1","true","yes"):
        from app.db import SessionLocal
        db = SessionLocal()
        try:
            bootstrap(db)
        finally:
            db.close()
    # Initialise the Doc-Engine HTTP client when integration is enabled.
    if settings.DOCENGINE_ENABLED:
        app.state.docengine = DocEngineClient(
            base_url=settings.DOCENGINE_URL,
            timeout=settings.DOCENGINE_TIMEOUT_S,
        )
    else:
        app.state.docengine = None


@app.on_event("shutdown")
async def shutdown_docengine():
    client: DocEngineClient | None = getattr(app.state, "docengine", None)
    if client is not None:
        await client.aclose()


@app.get("/health")
def health():
    return {"ok": True}


# ---------------------------------------------------------------------------
# Auth endpoints
# ---------------------------------------------------------------------------

@app.post("/auth/register", response_model=TokenResponse, status_code=201)
def register(body: RegisterRequest, db: Session = Depends(get_db)):
    user = auth_create_user(db, body.email, body.password, body.full_name)
    token = create_access_token(user.id)
    return TokenResponse(
        access_token=token,
        user_id=user.id,
        email=user.email,
        full_name=user.full_name,
    )


@app.post("/auth/login", response_model=TokenResponse)
def login(body: LoginRequest, db: Session = Depends(get_db)):
    user = auth_authenticate_user(db, body.email, body.password)
    token = create_access_token(user.id)
    return TokenResponse(
        access_token=token,
        user_id=user.id,
        email=user.email,
        full_name=user.full_name,
    )


@app.get("/auth/me", response_model=MeResponse)
def me(current_user: User = Depends(get_current_user)):
    return MeResponse(
        user_id=current_user.id,
        email=current_user.email,
        full_name=current_user.full_name,
        is_active=current_user.is_active,
    )


@app.get("/contracts/action-object-v1")
def action_object_contract_v1():
    # Expose the strict JSON Schema so frontend/tests can validate intent outputs against one source.
    return ACTION_OBJECT_JSON_SCHEMA


def _command_error_response(
    status_code: int,
    version: int,
    code: str,
    message: str,
    details: dict | None = None,
):
    payload = CommandErrorResponse(
        status="error",
        version=version,
        error={
            "code": code,
            "message": message,
            "details": details or {},
        },
    )
    return JSONResponse(status_code=status_code, content=payload.model_dump())


def _command_latest_preview(version_row: DocumentVersion | None) -> str:
    if not version_row:
        return ""
    change_log = version_row.change_log or {}
    prompt = change_log.get("prompt") or change_log.get("edit_prompt") or ""
    prompt = " ".join(str(prompt).split())
    if prompt:
        return prompt if len(prompt) <= 120 else prompt[:117] + "..."

    state = version_row.doc_state or {}
    structured = state.get("structured") if isinstance(state, dict) else None
    if isinstance(structured, dict):
        sections = structured.get("sections") or []
        for sec in sections:
            if not isinstance(sec, dict):
                continue
            items = ((sec.get("content") or {}).get("items")) or []
            for item in items:
                if isinstance(item, dict) and item.get("text"):
                    text = " ".join(str(item.get("text")).split())
                    return text if len(text) <= 120 else text[:117] + "..."
    return ""


def _coerce_non_negative_int(value: object) -> int:
    # Defensive coercion keeps telemetry shape stable even if an upstream field is null/non-numeric.
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, (int, float)):
        return max(0, int(value))
    return 0


def _command_latency_breakdown(
    *,
    stt_meta: dict,
    intent_ms: int,
    transform_ms: int,
    apply_ms: int,
    total_ms: int,
) -> dict[str, int]:
    # Keep the latency schema explicit so dashboards can aggregate every stage consistently.
    return {
        "stt_ms": _coerce_non_negative_int((stt_meta or {}).get("stt_latency_ms")),
        "intent_ms": _coerce_non_negative_int(intent_ms),
        "transform_ms": _coerce_non_negative_int(transform_ms),
        "apply_ms": _coerce_non_negative_int(apply_ms),
        "total_ms": _coerce_non_negative_int(total_ms),
    }


def _build_command_meta(
    *,
    request_id: str,
    intent_confidence: float,
    intent_source: str,
    repair_applied: bool,
    prompt_version: str,
    input_source: str,
    transcript: str | None,
    stt_meta: dict,
    auto_retried: bool,
    retry_count: int,
    base_version: int,
    transform_meta: dict,
    latency_ms: dict[str, int],
) -> dict:
    # One shared constructor avoids drift between applied/clarification/error paths.
    return {
        "intent_confidence": float(intent_confidence),
        "intent_source": intent_source,
        "repair_applied": bool(repair_applied),
        "prompt_version": prompt_version,
        "input_source": input_source,
        "transcript": transcript,
        "stt_model": (stt_meta or {}).get("stt_model"),
        "stt_device": (stt_meta or {}).get("stt_device"),
        "stt_compute_type": (stt_meta or {}).get("stt_compute_type"),
        "stt_language_detected": (stt_meta or {}).get("stt_language_detected"),
        "stt_confidence": (stt_meta or {}).get("stt_confidence"),
        "stt_latency_ms": (stt_meta or {}).get("stt_latency_ms"),
        "auto_retried": bool(auto_retried),
        "retry_count": int(retry_count),
        "base_version": int(base_version),
        "request_id": request_id,
        "trace_id": request_id,
        "transform_source": (transform_meta or {}).get("transform_source"),
        "transform_prompt_version": (transform_meta or {}).get("transform_prompt_version"),
        "transform_repair_applied": (transform_meta or {}).get("transform_repair_applied"),
        "latency_ms": latency_ms,
    }


def _emit_command_log(
    *,
    request_id: str,
    document_id: str,
    version: int,
    status: str,
    input_source: str,
    intent_source: str,
    prompt_version: str,
    repair_applied: bool,
    transform_meta: dict,
    latency_ms: dict[str, int],
    error_code: str | None = None,
    reason_code: str | None = None,
) -> None:
    # Exactly one log line per request makes incident triage and replay straightforward.
    payload = {
        "event": "command_request",
        "request_id": request_id,
        "trace_id": request_id,
        "document_id": document_id,
        "version": int(version),
        "status": status,
        "error_code": error_code,
        "reason_code": reason_code,
        "input_source": input_source,
        "intent_source": intent_source,
        "prompt_version": prompt_version,
        "repair_applied": bool(repair_applied),
        "transform_source": (transform_meta or {}).get("transform_source"),
        "transform_prompt_version": (transform_meta or {}).get("transform_prompt_version"),
        "transform_repair_applied": (transform_meta or {}).get("transform_repair_applied"),
        "latency_ms": latency_ms,
    }
    COMMAND_LOGGER.info(json.dumps(payload, ensure_ascii=False))


@app.post("/templates/register", response_model=TemplateRegisterResponse)
def register_template(name: str, doc_type: str, file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Template must be .docx")
    os.makedirs(settings.STORAGE_DIR, exist_ok=True)
    out_path = os.path.join(settings.STORAGE_DIR, "templates", f"{doc_type}_{file.filename}")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(file.file.read())

    zones = suggest_zones(out_path)
    t = Template(name=name, doc_type=doc_type, docx_path=out_path, zones_json=zones)
    db.add(t)
    db.commit()
    db.refresh(t)
    return TemplateRegisterResponse(template_id=t.id, name=t.name, doc_type=t.doc_type)

@app.get("/templates", response_model=list[TemplateOut])
def list_templates(doc_type: str | None = None, db: Session = Depends(get_db)):
    items = crud.list_templates(db, doc_type)
    return [TemplateOut(id=i.id, name=i.name, doc_type=i.doc_type, version=i.version, zones_json=i.zones_json) for i in items]


@app.get("/templates/{template_id}", response_model=TemplateOut)
def get_template(template_id: int, db: Session = Depends(get_db)):
    t = db.get(Template, template_id)
    if not t:
        raise HTTPException(404, "Template not found")
    return TemplateOut(id=t.id, name=t.name, doc_type=t.doc_type, version=t.version, zones_json=t.zones_json)


def normalize_doc_type(dt) -> str:
    """
    Ensure doc_type is a plain uppercase string like 'GOI_LETTER'.
    Handles Enum, 'DocType.X', whitespace, etc.
    """
    if dt is None:
        return "UNKNOWN"
    # Enum -> value or name
    if hasattr(dt, "value"):
        dt = dt.value
    elif hasattr(dt, "name"):
        dt = dt.name
    s = str(dt).strip()
    if "." in s:  # e.g., 'DocType.GOI_LETTER'
        s = s.split(".")[-1]
    
    s = s.upper()
    aliases = {
        "GOI": "GOI_LETTER",
        "GOVT_LETTER": "GOI_LETTER",
        "GOV_LETTER": "GOI_LETTER",
        "GOVERNMENT_OF_INDIA_LETTER": "GOI_LETTER",
        "DO": "DO_LETTER",
        "DEMI_OFFICIAL": "DO_LETTER",
        "MOVEMENT": "MOVEMENT_ORDER",
        "MOVE_ORDER": "MOVEMENT_ORDER",
        "LEAVE": "LEAVE_CERTIFICATE",
        "LEAVE_CERT": "LEAVE_CERTIFICATE",
    }
    return aliases.get(s, s)


def _build_pdf_path(docx_path: str) -> str:
    return os.path.splitext(docx_path)[0] + ".pdf"


def _rerender_version(db, doc: "Document", new_version: "DocumentVersion", new_state: dict) -> None:
    """Re-render DOCX for a new version created by a /command operation.

    For legacy docs: renders via docxtpl template.
    For blueprint docs: use _render_blueprint_docx() inline instead — this function
    handles only the legacy template path now.
    """
    import traceback as _tb
    from app.services.render_adapter import doc_state_from_filled_skeleton
    doc_type = (doc.doc_type or "").upper().replace("-", "_")
    template_obj = (
        db.query(Template)
        .filter(Template.doc_type == doc_type)
        .order_by(Template.id.desc())
        .first()
    )
    structured = (new_state.get("structured") or new_state)
    if not template_obj:
        # Blueprint doc — generate_plain_docx path
        if not doc.docengine_doc_id:
            return  # legacy doc without a template: skip
        try:
            from app.services.doc_importer import generate_plain_docx, sections_for_render
            secs = sections_for_render(structured)
            out_docx = crud.storage_path(doc.id, new_version.id, "docx")
            crud.ensure_dir(os.path.dirname(out_docx))
            generate_plain_docx(secs, doc.title or "Document", out_docx)
            new_version.docx_path = out_docx
            db.add(new_version)
            db.commit()
        except Exception as _e:
            COMMAND_LOGGER.error(
                "blueprint rerender failed doc=%s ver=%s: %s\n%s",
                doc.id, new_version.id, _e, _tb.format_exc()
            )
            try:
                db.rollback()
            except Exception:
                pass
            # Fallback: write to v0 stable slot so download doesn't 404
            try:
                from app.services.doc_importer import generate_plain_docx, sections_for_render
                secs = sections_for_render(structured)
                v0_path = crud.storage_path(doc.id, 0, "docx")
                crud.ensure_dir(os.path.dirname(v0_path))
                generate_plain_docx(secs, doc.title or "Document", v0_path)
                new_version.docx_path = v0_path
                db.add(new_version)
                db.commit()
            except Exception as _e2:
                COMMAND_LOGGER.error("blueprint rerender v0 fallback also failed: %s", _e2)
                try:
                    db.rollback()
                except Exception:
                    pass
        return
    try:
        render_state = doc_state_from_filled_skeleton(structured, doc_type)
        out_docx = crud.storage_path(doc.id, new_version.id, "docx")
        render_docx(template_obj.docx_path, render_state, out_docx)
        new_version.docx_path = out_docx
        db.add(new_version)
        db.commit()
    except Exception as _e:
        COMMAND_LOGGER.error(
            "rerender failed doc=%s ver=%s: %s\n%s",
            doc.id, new_version.id, _e, _tb.format_exc()
        )
        try:
            db.rollback()
        except Exception:
            pass


def _render_blueprint_docx_inline(
    doc_id: str,
    structured: dict,
    title: str,
    fallback_path: str,
    doc_type: str = "",
    db=None,
    template_id: int | None = None,
) -> str:
    """Render DOCX for a blueprint doc using the proper letter template when available.

    Uses `template_id` (int FK from documents.template_id) for a direct Template
    lookup when provided; falls back to doc_type heuristic otherwise.  Falls back
    to generate_plain_docx for uploaded/flexible docs or when alignment overrides
    are present (docxtpl templates have fixed alignment and cannot honour them).

    Returns the path of the written file, or fallback_path if generation fails.
    """
    import traceback as _tb
    v0_path = crud.storage_path(doc_id, 0, "docx")
    try:
        crud.ensure_dir(os.path.dirname(v0_path))
        # --- Preferred path: docxtpl template (GOI, DO, service letter, etc.) ---
        # Skip if any section has an alignment override — docxtpl templates have fixed
        # alignment and cannot reflect dynamic alignment changes.  generate_plain_docx
        # correctly applies left_indent / CENTER alignment per section.
        _has_alignment = any(
            isinstance(s, dict) and s.get("alignment")
            for s in (structured.get("sections") or [])
        )
        # If any section carries CSS styling (font/color/size set by the doc-engine or user
        # formatting commands), use generate_plain_docx so all per-run styles are preserved.
        # docxtpl templates render plain text only and cannot honour inline CSS.
        def _any_section_has_css(secs):
            for _s in (secs or []):
                if not isinstance(_s, dict):
                    continue
                _state = ((_s.get("content") or {}).get("richtext") or {}).get("state") or {}
                for _p in (_state.get("root") or {}).get("children") or []:
                    for _n in (_p.get("children") or []):
                        if _n.get("type") == "text" and (_n.get("format") or _n.get("style")):
                            return True
            return False
        _has_inline_css = _any_section_has_css(structured.get("sections") or [])
        if db and not _has_alignment and not _has_inline_css:
            doc_type_upper = doc_type.upper().replace("-", "_") if doc_type else ""
            # Prefer direct lookup by integer PK (documents.template_id FK)
            if template_id is not None:
                template_obj = db.get(Template, template_id)
            elif doc_type_upper:
                template_obj = (
                    db.query(Template)
                    .filter(Template.doc_type == doc_type_upper)
                    .order_by(Template.id.desc())
                    .first()
                )
            else:
                template_obj = None
            if template_obj:
                from app.services.render_adapter import doc_state_from_filled_skeleton
                render_state = doc_state_from_filled_skeleton(
                    structured, template_obj.doc_type
                )
                render_docx(template_obj.docx_path, render_state, v0_path)
                return v0_path
        # --- Fallback: generic plain DOCX (uploaded_doc / bp_flexible_v1, or when
        #     alignment overrides are present and docxtpl cannot honour them) ---
        from app.services.doc_importer import generate_plain_docx, sections_for_render
        secs = sections_for_render(structured)
        generate_plain_docx(secs, title or "Document", v0_path)
        return v0_path
    except Exception as _e:
        COMMAND_LOGGER.error(
            "inline blueprint DOCX render failed doc=%s: %s\n%s",
            doc_id, _e, _tb.format_exc()
        )
        return fallback_path




async def _generate_para_text(topic: str, doc_context: str = "") -> str:
    """
    Generate two formal paragraph from a topic description and/or document context.

    topic       — action_obj.content, e.g. "about ration" or "based on letter"
    doc_context — plain-text of existing document paragraphs (for context-aware generation)
    """
    import requests as _req
    t = re.sub(r"^\s*about\s+", "", topic.strip(), flags=re.IGNORECASE).strip()

    # Detect context-only requests: "based on letter", "similar to above", etc.
    if re.search(r"\b(based on|similar to|continue|following|like the)\b", t, re.IGNORECASE):
        t = ""  # LLM will use doc_context only

    if settings.LLM_PROVIDER.lower() == "ollama" and settings.COMMAND_TRANSFORM_USE_LLM:
        try:
            _req.get(
                settings.OLLAMA_BASE_URL.rstrip("/") + "/api/tags",
                timeout=max(0.1, float(settings.COMMAND_TRANSFORM_HEALTH_TIMEOUT_S)),
            )
            from app.ml.ollama_client import ollama_chat
            ctx_block = (
                f"\n\nExisting document paragraphs for context:\n{doc_context}"
                if doc_context else ""
            )
            instruction = (
                f"Write a paragraph about: {t}" if t
                else "Write the next logical paragraph that continues this formal military document."
            )
            text = await ollama_chat(
                "Write a single formal military document paragraph. No headers, no bullets. Return only the paragraph.",
                f"{instruction}{ctx_block}",
            )
            if text and text.strip():
                return text.strip()
        except Exception:
            pass

    # Stub fallback
    if t:
        return (
            f"The matter concerning {t} is noted. Necessary action will be taken "
            "as per existing orders and guidelines."
        )
    return "This paragraph will be updated with relevant details as directed."


@app.post("/documents/{document_id}/command")
async def command_document(document_id: str, req: CommandRequest, db: Session = Depends(get_db)):
    """v1 command contract endpoint for text/voice command execution.

    Why this endpoint exists:
    - Provides one stable interface for downstream modules (intent extraction/planner/frontend).
    - Keeps command outcomes deterministic with explicit statuses:
      applied | needs_clarification | error.
    """
    request_started = time.perf_counter()
    request_id = f"cmd_{uuid.uuid4().hex}"

    # Track request-stage timings in one place so every response and log line shares identical metrics.
    intent_ms = 0
    transform_ms = 0
    apply_ms = 0

    requested_version = int(req.version)
    command_version = requested_version
    auto_retried = False
    retry_count = 0

    input_source = "text"
    transcript: str | None = None
    stt_meta = {
        "stt_model": None,
        "stt_device": None,
        "stt_compute_type": None,
        "stt_language_detected": None,
        "stt_confidence": None,
        "stt_latency_ms": None,
    }

    intent_confidence = 0.0
    intent_source = "fallback_rule"
    repair_applied = False
    prompt_version = "intent_extraction_v1"
    transform_meta: dict = {
        "transform_source": None,
        "transform_prompt_version": None,
        "transform_repair_applied": None,
    }

    def _latency_payload() -> dict[str, int]:
        return _command_latency_breakdown(
            stt_meta=stt_meta,
            intent_ms=intent_ms,
            transform_ms=transform_ms,
            apply_ms=apply_ms,
            total_ms=int((time.perf_counter() - request_started) * 1000),
        )

    def _meta_payload() -> dict:
        return _build_command_meta(
            request_id=request_id,
            intent_confidence=intent_confidence,
            intent_source=intent_source,
            repair_applied=repair_applied,
            prompt_version=prompt_version,
            input_source=input_source,
            transcript=transcript,
            stt_meta=stt_meta,
            auto_retried=auto_retried,
            retry_count=retry_count,
            base_version=requested_version,
            transform_meta=transform_meta,
            latency_ms=_latency_payload(),
        )

    def _emit(
        *,
        status: str,
        version: int,
        error_code: str | None = None,
        reason_code: str | None = None,
    ) -> None:
        _emit_command_log(
            request_id=request_id,
            document_id=document_id,
            version=version,
            status=status,
            input_source=input_source,
            intent_source=intent_source,
            prompt_version=prompt_version,
            repair_applied=repair_applied,
            transform_meta=transform_meta,
            latency_ms=_latency_payload(),
            error_code=error_code,
            reason_code=reason_code,
        )

    def _error(
        *,
        status_code: int,
        version: int,
        code: str,
        message: str,
        details: dict | None = None,
    ):
        _emit(status="error", version=version, error_code=code)
        return _command_error_response(
            status_code=status_code,
            version=version,
            code=code,
            message=message,
            details=details,
        )

    def _needs_clarification(*, version: int, reason_code: str, question: str, options: list[dict]):
        # Unsupported actions are tracked as a standardized code in logs while keeping UX in clarification mode.
        log_error_code = "unsupported_action" if reason_code == "unsupported_action" else None
        _emit(
            status="needs_clarification",
            version=version,
            error_code=log_error_code,
            reason_code=reason_code,
        )
        resp = CommandNeedsClarificationResponse(
            status="needs_clarification",
            version=version,
            clarification={
                "question": question,
                "options": options,
                "clarification_token": build_clarification_token(),
                "reason_code": reason_code,
            },
            meta=_meta_payload(),
        )
        return resp.model_dump()

    if not settings.ENABLE_COMMAND_API:
        return _error(
            status_code=404,
            version=command_version,
            code="unsupported_action",
            message="Command API is disabled by ENABLE_COMMAND_API=false",
            details={},
        )

    doc = db.get(Document, document_id)
    if not doc or not doc.current_version_id:
        return _error(
            status_code=404,
            version=command_version,
            code="patch_apply_failed",
            message="Document not found",
            details={"document_id": document_id},
        )

    effective_version = int(doc.current_version_id)
    command_version = effective_version

    # Optimistic concurrency guard with optional bounded auto-retry.
    if requested_version != effective_version:
        latest_version_row = db.get(DocumentVersion, effective_version)
        latest_preview = _command_latest_preview(latest_version_row)
        auto_retry_eligible = requested_version < effective_version
        if not (req.auto_retry and auto_retry_eligible):
            return _error(
                status_code=409,
                version=requested_version,
                code="version_conflict",
                message="Client version is stale",
                details={
                    "server_version": effective_version,
                    "latest_preview": latest_preview,
                    "auto_retry_eligible": auto_retry_eligible,
                    "hint": "Document changed, please retry",
                },
            )
        auto_retried = True
        retry_count = 1

    cur_v = db.get(DocumentVersion, effective_version)
    if not cur_v:
        return _error(
            status_code=404,
            version=command_version,
            code="patch_apply_failed",
            message="Current version missing",
            details={},
        )

    combined = cur_v.doc_state or {}
    structured = combined.get("structured") if isinstance(combined, dict) else None
    if not isinstance(structured, dict):
        return _error(
            status_code=422,
            version=command_version,
            code="invalid_command_payload",
            message="Document is not in structured mode",
            details={},
        )

    prompt = req.input.value or ""
    if req.input.type.value == "voice":
        if not settings.ENABLE_VOICE_INPUT:
            return _error(
                status_code=400,
                version=command_version,
                code="could_not_understand_audio",
                message="Voice input is disabled in this backend",
                details={},
            )
        try:
            audio_bytes = base64.b64decode(req.input.audio_base64 or "", validate=True)
        except (binascii.Error, ValueError):
            return _error(
                status_code=400,
                version=command_version,
                code="could_not_understand_audio",
                message="Could not understand audio payload",
                details={},
            )

        try:
            transcript, stt_meta = transcribe(audio_bytes=audio_bytes, mime_type=req.input.mime_type or "")
        except STTError as ex:
            return _error(
                status_code=400,
                version=command_version,
                code="could_not_understand_audio",
                message=str(ex),
                details={},
            )

        input_source = "voice"
        prompt = transcript

        stt_confidence = stt_meta.get("stt_confidence")
        if isinstance(stt_confidence, (int, float)) and stt_confidence < float(settings.STT_CONFIRM_CONFIDENCE):
            return _needs_clarification(
                version=command_version,
                reason_code="stt_low_confidence_confirm",
                question=f'I heard: "{prompt}". Is that correct?',
                options=[
                    {"label": "Yes, continue", "token": "confirm_transcript_yes"},
                    {"label": "No, repeat", "token": "confirm_transcript_no"},
                ],
            )

    intent_started = time.perf_counter()
    try:
        intent_result = await extract_action_object_with_meta(prompt, req.context, structured)
        action_obj = intent_result.action_object
        intent_source = intent_result.intent_source
        repair_applied = intent_result.repair_applied
        prompt_version = intent_result.prompt_version
        intent_confidence = float(action_obj.confidence)
    except IntentParseError as ex:
        intent_ms = int((time.perf_counter() - intent_started) * 1000)
        return _error(
            status_code=500,
            version=command_version,
            code="intent_parse_error",
            message=str(ex),
            details={},
        )
    except Exception as ex:
        intent_ms = int((time.perf_counter() - intent_started) * 1000)
        return _error(
            status_code=500,
            version=command_version,
            code="intent_parse_error",
            message=f"Failed to parse command intent: {ex}",
            details={},
        )
    intent_ms = int((time.perf_counter() - intent_started) * 1000)

    if action_obj.needs_clarification:
        clarification = action_obj.clarification or {
            "question": "Which target should I update?",
            "options": [],
        }
        question = clarification.question if hasattr(clarification, "question") else clarification["question"]
        options = clarification.options if hasattr(clarification, "options") else clarification["options"]
        return _needs_clarification(
            version=command_version,
            reason_code="intent_needs_clarification",
            question=question,
            options=options,
        )

    # --- Blueprint path: forward to doc-engine microservice ---
    if doc.docengine_doc_id:
        # Fetch fresh sections from doc-engine so action_bridge has live data
        # (cached DB structured state may be stale, causing wrong insertion positions).
        _live_sections: list = (structured.get("sections") or []) if isinstance(structured, dict) else []
        _de_client_early: DocEngineClient | None = app.state.docengine
        if _de_client_early is not None:
            try:
                _fresh = await _de_client_early.get_document(doc.docengine_doc_id)
                _live_sections = (_fresh.get("data") or {}).get("sections") or _live_sections
            except Exception:
                pass  # fall back to cached sections

        de_action = ml_action_to_de(action_obj, _live_sections)
        action_val = action_obj.action.value if hasattr(action_obj.action, "value") else str(action_obj.action)
        de_version = int((cur_v.doc_state or {}).get("docengine_version", requested_version))
        de_client: DocEngineClient | None = app.state.docengine

        if de_action is None:
            # --- Content ops: ML LLM generates new text, then PATCH /sections ---
            # REPLACE_TEXT and INSERT_TEXT with pre-filled content act like REWRITE_CONTENT
            # (content is already set by the rule resolver — no LLM call needed).
            _is_text_set_with_content = (
                action_val in {"REPLACE_TEXT", "INSERT_TEXT"}
                and bool(action_obj.content)
            )
            if _is_text_set_with_content:
                from app.schemas import CommandAction as _CA
                action_obj = action_obj.model_copy(update={"action": _CA.REWRITE_CONTENT})
                action_val = "REWRITE_CONTENT"
            if action_val in CONTENT_OPS:
                from app.services.lexical_wrapper import (
                    lexical_to_plain_text, text_to_lexical_node, apply_format_to_word,
                )
                from app.services.content_transform import (
                    call_llm_transform, TransformError,
                    prepare_input_text, postprocess_text,
                    _supported_transform_prompt,
                )

                section_id = action_obj.target.section_id

                # --- "create new section + set content" signal ---
                # _try_section_content_set encodes the target type in para_id when the section
                # doesn't exist yet: para_id = "__create:{sec_type}"
                _para_id_str = str(action_obj.target.para_id or "")
                if _para_id_str.startswith("__create:") and action_obj.content:
                    # para_id format: "__create:{sec_type}" or "__create:{sec_type}:{pos_hint}"
                    # pos_hint: "after_para:N" (1-based), "after_last", or absent (default)
                    _create_parts = _para_id_str[len("__create:"):].split(":", 1)
                    _create_sec_type = _create_parts[0]
                    _pos_hint = _create_parts[1] if len(_create_parts) > 1 else ""
                    if de_client is None:
                        return _error(status_code=503, version=command_version, code="patch_apply_failed", message="Doc-Engine client unavailable", details={})
                    _new_sec_id: str | None = None
                    try:
                        # Resolve insertion position from pos_hint or default heuristic
                        _body_types = {"paragraph", "table_block", "remarks_block"}
                        _ins_position: dict
                        _all_sections = structured.get("sections") or []
                        _para_sections = [s for s in _all_sections
                                          if isinstance(s, dict) and s.get("type") == "paragraph"]
                        if _pos_hint.startswith("after_para:"):
                            # Insert after Nth paragraph (1-based)
                            _n = int(_pos_hint.split(":")[-1])
                            _anchor = _para_sections[_n - 1] if 0 < _n <= len(_para_sections) else None
                            _ins_position = (
                                {"policy": "after", "section_id": _anchor["id"]}
                                if _anchor else {"policy": "end"}
                            )
                        elif _pos_hint == "after_last":
                            # After all body content — use the last body-type section as anchor
                            _last_body = next(
                                (s for s in reversed(_all_sections)
                                 if isinstance(s, dict) and s.get("type") in _body_types),
                                None,
                            )
                            _ins_position = (
                                {"policy": "after", "section_id": _last_body["id"]}
                                if _last_body else {"policy": "end"}
                            )
                        elif _create_sec_type in _body_types:
                            # Default: after last paragraph
                            _last_para = _para_sections[-1] if _para_sections else None
                            _ins_position = (
                                {"policy": "after", "section_id": _last_para["id"]}
                                if _last_para else {"policy": "end"}
                            )
                        else:
                            _ins_position = {"policy": "end"}
                        _ins = await de_client.apply_command(
                            doc.docengine_doc_id, de_version,
                            {"action": "INSERT_SECTION", "target_type": _create_sec_type,
                             "target_ref": None, "position": _ins_position, "ai_instruction": None},
                        )
                        if _ins.get("status") != "applied":
                            return _error(status_code=422, version=command_version, code="patch_apply_failed", message=f"INSERT_SECTION for {_create_sec_type} not applied", details={})
                        de_version = _ins.get("version", de_version)
                        _new_sec_id = next(
                            (u.get("section", {}).get("id") for u in (_ins.get("updates") or [])
                             if u.get("op") == "insert"),
                            None,
                        )
                        if not _new_sec_id:
                            return _error(500, version=command_version, code="patch_apply_failed", message="INSERT_SECTION returned no section id", details={})
                    except DocEngineError as exc:
                        # If section already exists (stale local state missed it), fall back to
                        # finding and patching the existing section instead of failing.
                        if exc.status_code == 422 and "single_instance_violated" in str(exc.body):
                            try:
                                _cur_doc = await de_client.get_document(doc.docengine_doc_id)
                                de_version = _cur_doc.get("version", de_version)
                                _existing = next(
                                    (s for s in (_cur_doc.get("data", {}).get("sections") or [])
                                     if isinstance(s, dict) and s.get("type") == _create_sec_type),
                                    None,
                                )
                                if _existing:
                                    _new_sec_id = _existing["id"]
                                else:
                                    return _error(status_code=exc.status_code or 500, version=command_version, code="patch_apply_failed", message=str(exc.body), details={})
                            except DocEngineError:
                                return _error(status_code=exc.status_code or 500, version=command_version, code="patch_apply_failed", message=str(exc.body), details={})
                        else:
                            return _error(status_code=exc.status_code or 500, version=command_version, code="patch_apply_failed", message=str(exc.body), details={})
                    # Inject into structured so PATCH response below sees it.
                    # Use the full section object from the INSERT response if available.
                    _new_sec_full = next(
                        (u.get("section") for u in (_ins.get("updates") or [])
                         if u.get("op") == "insert" and isinstance(u.get("section"), dict)),
                        None,
                    )
                    _new_sec = _new_sec_full or {"id": _new_sec_id, "type": _create_sec_type, "content": {}}
                    if isinstance(structured, dict):
                        structured.setdefault("sections", [])
                        if isinstance(structured["sections"], list):
                            if not any(s.get("id") == _new_sec_id for s in structured["sections"]):
                                # Insert at the correct position to keep local state in sync
                                # with doc-engine ordering (after anchor, not always at end)
                                if _ins_position.get("policy") == "after" and _ins_position.get("section_id"):
                                    _anchor_local_idx = next(
                                        (i for i, _s in enumerate(structured["sections"])
                                         if _s.get("id") == _ins_position["section_id"]),
                                        None,
                                    )
                                    if _anchor_local_idx is not None:
                                        structured["sections"].insert(_anchor_local_idx + 1, _new_sec)
                                    else:
                                        structured["sections"].append(_new_sec)
                                else:
                                    structured["sections"].append(_new_sec)
                    section_id = _new_sec_id  # route to the single-section PATCH path

                # Build the list of sections to transform.
                # section_id=None (scope=DOCUMENT) means "apply to all paragraph sections".
                if not section_id:
                    _all_secs = [
                        s for s in (structured.get("sections") or [])
                        if isinstance(s, dict) and s.get("type") == "paragraph"
                    ]
                    if not _all_secs:
                        return _error(
                            status_code=422,
                            version=command_version,
                            code="intent_parse_error",
                            message="No paragraph sections found in this document",
                            details={},
                        )
                else:
                    _tgt = next(
                        (s for s in (structured.get("sections") or [])
                         if isinstance(s, dict) and s.get("id") == section_id),
                        None,
                    )
                    if not _tgt:
                        return _error(
                            status_code=422,
                            version=command_version,
                            code="patch_apply_failed",
                            message=f"Section {section_id} not found in document",
                            details={},
                        )
                    _all_secs = [_tgt]

                if de_client is None:
                    return _error(status_code=503, version=command_version, code="patch_apply_failed", message="Doc-Engine client unavailable", details={})

                try:
                    _prompt_ver = _supported_transform_prompt(action_obj)
                except Exception:
                    _prompt_ver = "rewrite_v1"

                updates = []
                new_de_version = de_version
                transform_started = time.perf_counter()

                for _sec in _all_secs:
                    _sec_id = _sec.get("id")
                    richtext_state = ((_sec.get("content") or {}).get("richtext") or {}).get("state")
                    source_text = lexical_to_plain_text(richtext_state) if richtext_state else ""

                    if action_obj.content:
                        # Pre-filled content (e.g. "add signee Brig RS Sharma") — skip LLM
                        new_text = action_obj.content
                        # For signee_block, auto-format "RS Bhati" → "Yours sincerely,\n(RS Bhati)"
                        if _sec.get("type") == "signee_block" and new_text:
                            from app.services.doc_importer import _format_signee_lines
                            _lines = _format_signee_lines(new_text)
                            # Prepend "Yours sincerely," if not already present
                            if _lines and not _lines[0].lower().startswith("yours"):
                                _lines = ["Yours sincerely,"] + _lines
                            new_text = "\n\n".join(_lines)
                        # For table_block, content is "NxM" dims → build markdown table
                        elif _sec.get("type") == "table_block" and re.match(r"^\d+x\d+$", new_text or ""):
                            _dims = new_text.split("x")
                            _rows, _cols = int(_dims[0]), int(_dims[1])
                            _col_headers = " | ".join(f"Col {i+1}" for i in range(_cols))
                            _separator = " | ".join("-------" for _ in range(_cols))
                            _empty_row = " | ".join("       " for _ in range(_cols))
                            _table_lines = [
                                f"| {_col_headers} |",
                                f"|{_separator.replace(' | ', '|')}|",
                            ] + [f"| {_empty_row} |"] * _rows
                            new_text = "\n".join(_table_lines)
                    else:
                        prepared_text, prep_meta = prepare_input_text(
                            source_text,
                            preserve_numbering=bool(action_obj.params.preserve_numbering),
                        )
                        try:
                            new_text, t_meta = await call_llm_transform(
                                action_object=action_obj,
                                source_text=prepared_text,
                                prompt_version=_prompt_ver,
                            )
                        except TransformError as ex:
                            transform_ms = int((time.perf_counter() - transform_started) * 1000)
                            return _error(
                                status_code=500,
                                version=command_version,
                                code="patch_apply_failed",
                                message=f"Content transform failed: {ex}",
                                details={},
                            )
                        new_text = postprocess_text(new_text, prep_meta)
                    if _sec.get("type") == "subject":
                        new_lexical = text_to_lexical_node(new_text, bold=True, underline=True, align="center")
                    else:
                        new_lexical = text_to_lexical_node(new_text)
                    # FIX_GRAMMAR: re-apply inline formatting (bold/italic/underline) from
                    # the original Lexical state so word-level formatting is preserved.
                    if action_val == "FIX_GRAMMAR" and richtext_state:
                        FLAG_BOLD, FLAG_ITALIC, FLAG_UNDERLINE = 1, 2, 8
                        for _lp in (richtext_state.get("root") or {}).get("children") or []:
                            for _node in (_lp.get("children") or []):
                                if _node.get("type") != "text":
                                    continue
                                _fmt = _node.get("format", 0)
                                if not _fmt:
                                    continue
                                _word = _node.get("text", "").strip()
                                if not _word:
                                    continue
                                _style: dict = {}
                                if _fmt & FLAG_BOLD:      _style["bold"] = True
                                if _fmt & FLAG_ITALIC:    _style["italic"] = True
                                if _fmt & FLAG_UNDERLINE: _style["underline"] = True
                                if _style and _word.lower() in lexical_to_plain_text(new_lexical).lower():
                                    new_lexical = apply_format_to_word(new_lexical, _word, _style)
                    lexical_content = {"richtext": {"format": "lexical", "state": new_lexical}}
                    try:
                        patch_result = await de_client.patch_section(
                            doc.docengine_doc_id, _sec_id, new_de_version, lexical_content
                        )
                        new_de_version = patch_result.get("version", new_de_version)
                    except DocEngineError as exc:
                        if exc.status_code == 409:
                            return _error(status_code=409, version=command_version, code="version_conflict", message="Doc-Engine version conflict", details={})
                        return _error(exc.status_code or 500, version=command_version, code="patch_apply_failed", message=str(exc.body), details={})
                    _sec.setdefault("content", {}).setdefault("richtext", {})["state"] = new_lexical
                    updates.append({"op": "update_section", "section_id": _sec_id, "new_lexical_state": new_lexical})

                transform_ms = int((time.perf_counter() - transform_started) * 1000)

                new_state = dict(cur_v.doc_state or {})
                new_state["docengine_version"] = new_de_version
                change_log = {
                    "action": "content_op_blueprint",
                    "request_id": request_id,
                    "prompt": prompt,
                    "action_object": action_obj.model_dump(),
                    "section_ids": [s.get("id") for s in _all_secs],
                    "intent_source": intent_source,
                    "docengine_doc_id": doc.docengine_doc_id,
                    "docengine_version_before": de_version,
                    "docengine_version_after": new_de_version,
                }
                _fallback_docx = cur_v.docx_path or crud.storage_path(doc.id, 0, "docx")
                new_docx_path = _render_blueprint_docx_inline(
                    doc.id, new_state.get("structured") or new_state,
                    doc.title or "Document", _fallback_docx,
                    doc_type=doc.doc_type or "", db=db,
                )
                try:
                    new_version = crud.add_version(db, doc.id, new_state, change_log, new_docx_path)
                except Exception as ex:
                    return _error(status_code=500, version=command_version, code="patch_apply_failed", message=f"Failed to persist: {ex}", details={})
                resp = CommandAppliedResponse(status="applied", version=new_version.id, updates=updates, meta=_meta_payload())
                _emit(status="applied", version=new_version.id)
                result = resp.model_dump()
                result["docengine_version"] = new_de_version
                result["docx_url"] = f"/documents/{doc.id}/versions/{new_version.id}/download?format=docx"
                _co_fresh: dict = {}
                try:
                    _co_de_doc = await de_client.get_document(doc.docengine_doc_id)
                    _co_fresh = dict(_co_de_doc.get("data") or {})
                    _co_fresh["_slots"] = (new_state.get("structured") or {}).get("_slots", {})
                except Exception:
                    _co_fresh = dict(new_state.get("structured") or {})
                result["data"] = _co_fresh
                return result

            # SET_FORMAT for blueprint docs: apply Lexical format flags + alignment + PATCH to doc-engine
            if action_val == "SET_FORMAT":
                style = dict(action_obj.params.style_params or {})
                if not style:
                    return _error(status_code=422, version=command_version, code="intent_parse_error",
                                  message="SET_FORMAT requires style params (bold, italic, align, etc.)", details={})
                target_word    = style.pop("target_word",    None)
                document_wide  = style.pop("document_wide",  False)
                # Alignment is a block-level (section-level) property — not a Lexical text-run flag.
                # Extract it separately and pass it to patch_section as a dedicated field.
                align = style.pop("align", None)
                section_id = action_obj.target.section_id
                if not section_id and not document_wide:
                    return _error(status_code=422, version=command_version, code="intent_parse_error",
                                  message="SET_FORMAT requires a target section_id", details={})
                target_sec = next(
                    (s for s in (structured.get("sections") or [])
                     if isinstance(s, dict) and s.get("id") == section_id),
                    None,
                ) if section_id else None
                if not target_sec and not document_wide:
                    return _error(status_code=422, version=command_version, code="section_not_found",
                                  message=f"Section {section_id} not found in document", details={})
                richtext_state = ((target_sec.get("content") or {}).get("richtext") or {}).get("state") \
                    if target_sec else None
                if not isinstance(richtext_state, dict) and not document_wide:
                    return _error(status_code=422, version=command_version, code="section_not_found",
                                  message="Section has no Lexical state to format", details={})

                from app.services.lexical_wrapper import (
                    apply_format_to_lexical, apply_format_to_word, lexical_to_plain_text,
                )

                if de_client is None:
                    return _error(status_code=503, version=command_version, code="patch_apply_failed",
                                  message="Doc-Engine client is not available", details={})

                new_de_version = de_version
                _updated_sections: list[dict] = []
                _live_sections_after: list | None = None   # set in document_wide branch

                if document_wide and style:
                    # --- Apply format to ALL sections (font, size, color, bold, etc.) ---
                    # Always fetch fresh doc state so we apply CSS on the actual current
                    # content, not a potentially stale cached version (stale state would
                    # silently wipe paragraph content on PATCH).
                    try:
                        _fresh = await de_client.get_document(doc.docengine_doc_id)
                        _live_sections = (_fresh.get("data") or {}).get("sections") or []
                        new_de_version = _fresh.get("version", new_de_version)
                    except Exception:
                        _live_sections = structured.get("sections") or []
                    for _ds in _live_sections:
                        if not isinstance(_ds, dict):
                            continue
                        _rs = ((_ds.get("content") or {}).get("richtext") or {}).get("state")
                        if not isinstance(_rs, dict):
                            continue
                        _new_lx = apply_format_to_lexical(_rs, dict(style))
                        _ds_id = _ds.get("id")
                        try:
                            _pr = await de_client.patch_section(
                                doc.docengine_doc_id, _ds_id, new_de_version,
                                {"richtext": {"format": "lexical", "state": _new_lx}},
                                alignment=align,
                            )
                            new_de_version = _pr.get("version", new_de_version)
                            _ds.setdefault("content", {}).setdefault("richtext", {})["state"] = _new_lx
                            if align:
                                _ds["alignment"] = align
                            _updated_sections.append({"op": "update_section", "section_id": _ds_id})
                        except DocEngineError as exc:
                            if exc.status_code == 409:
                                return _error(status_code=409, version=command_version,
                                              code="version_conflict",
                                              message="Doc-Engine version conflict", details={})
                            logging.warning("set_format doc_wide: patch sec %s err=%s", _ds_id, exc)
                    # Keep the updated sections so new_state can persist correct font.
                    _live_sections_after = _live_sections

                elif target_word and style:
                    # --- Apply word-format to ALL sections that contain the word ---
                    _all_format_secs = [
                        s for s in (structured.get("sections") or [])
                        if isinstance(s, dict)
                    ]
                    for _fs in _all_format_secs:
                        _rs = ((_fs.get("content") or {}).get("richtext") or {}).get("state")
                        if not isinstance(_rs, dict):
                            continue
                        # Quick text-presence check to avoid unnecessary patching
                        if target_word.lower() not in lexical_to_plain_text(_rs).lower():
                            continue
                        _new_lx = apply_format_to_word(_rs, target_word, dict(style))
                        _fs_id = _fs.get("id")
                        try:
                            _pr = await de_client.patch_section(
                                doc.docengine_doc_id, _fs_id, new_de_version,
                                {"richtext": {"format": "lexical", "state": _new_lx}},
                            )
                            new_de_version = _pr.get("version", new_de_version)
                            _fs.setdefault("content", {}).setdefault("richtext", {})["state"] = _new_lx
                            _updated_sections.append({"op": "update_section", "section_id": _fs_id,
                                                       "new_lexical_state": _new_lx})
                        except DocEngineError as exc:
                            if exc.status_code == 409:
                                return _error(status_code=409, version=command_version,
                                              code="version_conflict",
                                              message="Doc-Engine version conflict", details={})
                            logging.warning("set_format word: patch sec %s err=%s", _fs_id, exc)
                    # Alignment still applies to the resolved single section (if any)
                    if align and target_sec:
                        try:
                            _pr2 = await de_client.patch_section(
                                doc.docengine_doc_id, section_id, new_de_version,
                                {"richtext": {"format": "lexical", "state": richtext_state}},
                                alignment=align,
                            )
                            new_de_version = _pr2.get("version", new_de_version)
                            target_sec["alignment"] = align
                        except DocEngineError as exc:
                            logging.warning("set_format align: patch err=%s", exc)
                else:
                    # Single-section path (whole-section format or alignment only)
                    if style:
                        new_lexical_state = apply_format_to_lexical(richtext_state, style)
                    else:
                        new_lexical_state = richtext_state  # alignment-only
                    try:
                        patch_result = await de_client.patch_section(
                            doc.docengine_doc_id, section_id, new_de_version,
                            {"richtext": {"format": "lexical", "state": new_lexical_state}},
                            alignment=align,
                        )
                    except DocEngineError as exc:
                        if exc.status_code == 409:
                            return _error(status_code=409, version=command_version,
                                          code="version_conflict",
                                          message="Doc-Engine version conflict", details={})
                        return _error(status_code=exc.status_code or 500, version=command_version,
                                      code="patch_apply_failed", message=str(exc.body), details={})
                    new_de_version = patch_result.get("version", new_de_version)
                    target_sec.setdefault("content", {}).setdefault("richtext", {})["state"] = new_lexical_state
                    if align:
                        target_sec["alignment"] = align
                    _updated_sections.append({"op": "update_section", "section_id": section_id,
                                               "new_lexical_state": new_lexical_state})

                new_state = dict(cur_v.doc_state or {})
                new_state["docengine_version"] = new_de_version
                # For document-wide format (font size, color, etc.) persist the updated
                # Lexical states into new_state so re-fetches return the correct data.
                if _live_sections_after is not None:
                    _new_struct = dict(new_state.get("structured") or {})
                    _new_struct["sections"] = _live_sections_after
                    # Also update style_defaults so the frontend global renderer
                    # reflects the new values (e.g. font_size_pt, font_family).
                    _sd = dict(_new_struct.get("style_defaults") or {})
                    if style.get("size"):
                        _sd["font_size_pt"] = style["size"]
                    if style.get("font"):
                        _sd["font_family"] = style["font"]
                    if style.get("color"):
                        _sd["text_color"] = style["color"]
                    _new_struct["style_defaults"] = _sd
                    new_state["structured"] = _new_struct
                if align and isinstance(new_state.get("structured"), dict):
                    for _s in (new_state["structured"].get("sections") or []):
                        if isinstance(_s, dict) and _s.get("id") == section_id:
                            _s["alignment"] = align
                            break
                change_log = {"action": "set_format_blueprint", "request_id": request_id,
                              "prompt": prompt, "section_id": section_id,
                              "style": style, "align": align, "target_word": target_word,
                              "docengine_version_after": new_de_version}
                _fallback_docx = cur_v.docx_path or crud.storage_path(doc.id, 0, "docx")
                new_docx_path = _render_blueprint_docx_inline(
                    doc.id, new_state.get("structured") or new_state,
                    doc.title or "Document", _fallback_docx,
                    doc_type=doc.doc_type or "", db=db,
                    template_id=doc.template_id,
                )
                try:
                    new_version = crud.add_version(db, doc.id, new_state, change_log, new_docx_path)
                except Exception as ex:
                    return _error(status_code=500, version=command_version, code="patch_apply_failed",
                                  message=f"Failed to persist: {ex}", details={})
                resp = CommandAppliedResponse(status="applied", version=new_version.id,
                                              updates=_updated_sections or [{"op": "update_section",
                                                                              "section_id": section_id}],
                                              meta=_meta_payload())
                _emit(status="applied", version=new_version.id)
                result = resp.model_dump()
                result["docengine_version"] = new_de_version
                result["docx_url"] = f"/documents/{doc.id}/versions/{new_version.id}/download?format=docx"
                # Include fresh document data so the frontend can re-render with
                # the updated Lexical states (font size, color, etc.) immediately.
                _fmt_fresh: dict = {}
                try:
                    _fmt_de_doc = await de_client.get_document(doc.docengine_doc_id)
                    _fmt_fresh = dict(_fmt_de_doc.get("data") or {})
                    _new_struct_ref = new_state.get("structured") or {}
                    _fmt_fresh["_slots"] = _new_struct_ref.get("_slots", {})
                    # Merge updated style_defaults into the response data so the
                    # frontend sees the new font_size_pt / font_family immediately.
                    if _new_struct_ref.get("style_defaults"):
                        _fmt_fresh["style_defaults"] = _new_struct_ref["style_defaults"]
                except Exception:
                    _fmt_fresh = dict(new_state.get("structured") or {})
                result["data"] = _fmt_fresh
                return result

            # UNDO and other local-only actions not supported on blueprint docs
            return _error(
                status_code=422,
                version=command_version,
                code="unsupported_action",
                message=f"Action '{action_val}' is not supported via /command for blueprint documents.",
                details={"action": action_val},
            )

        if de_client is None:
            return _error(
                status_code=503,
                version=command_version,
                code="patch_apply_failed",
                message="Doc-Engine client is not available",
                details={},
            )

        apply_started = time.perf_counter()
        try:
            de_result = await de_client.apply_command(doc.docengine_doc_id, de_version, de_action)
        except DocEngineError as exc:
            apply_ms = int((time.perf_counter() - apply_started) * 1000)
            if exc.status_code == 409:
                return _error(
                    status_code=409,
                    version=command_version,
                    code="version_conflict",
                    message="Doc-Engine version conflict — client version is stale",
                    details={"docengine_doc_id": doc.docengine_doc_id},
                )
            return _error(
                status_code=exc.status_code if exc.status_code >= 400 else 500,
                version=command_version,
                code="patch_apply_failed",
                message=str(exc.body),
                details={},
            )
        apply_ms = int((time.perf_counter() - apply_started) * 1000)

        de_status = de_result.get("status")
        if de_status == "needs_clarification":
            return _needs_clarification(
                version=command_version,
                reason_code="intent_needs_clarification",
                question=de_result.get("question", "Which target should I update?"),
                options=de_result.get("options", []),
            )

        new_de_version = de_result.get("version", de_version)
        updates = de_result.get("updates") or []

        # --- Way 1: auto-fill new paragraph if content description present ---
        _DEFAULT_CONTENT = {"", "New paragraph.", "new paragraph."}
        _ao_content = (action_obj.content or "").strip()
        if (
            action_val == "ADD_PARAGRAPH"
            and _ao_content not in _DEFAULT_CONTENT
            and de_client is not None
        ):
            _new_sec_id = next(
                (u.get("section", {}).get("id") for u in updates if u.get("op") == "insert"),
                None,
            )
            if _new_sec_id:
                from app.services.lexical_wrapper import text_to_lexical_node
                from app.services.render_adapter import _section_text
                # Use live sections (fresh from DE) so paragraph count and context are accurate
                _para_sections = [
                    s for s in _live_sections
                    if isinstance(s, dict) and s.get("type") == "paragraph"
                ]
                # Calculate the correct number for the new paragraph based on the anchor.
                # action_obj.target.para_id is "pN" when inserting after paragraph N.
                _anchor_pid = (action_obj.target.para_id or "")
                _anchor_pid_m = re.match(r"^p(\d+)$", _anchor_pid)
                if _anchor_pid_m:
                    _next_para_num = int(_anchor_pid_m.group(1)) + 1
                else:
                    _next_para_num = len(_para_sections) + 1
                _doc_context = "\n\n".join(
                    _section_text(s).strip() for s in _para_sections if _section_text(s).strip()
                )
                try:
                    _gen_text = await _generate_para_text(_ao_content, _doc_context)
                    # Strip any leading number the LLM may have added
                    _gen_text = re.sub(r"^\s*\d+[\.\)]\s*", "", _gen_text).strip()
                    # Build two-node Lexical: first node is the number prefix (matches
                    # doc-engine's renumber regex "^\d+\.\s*$"), second is the content.
                    # This allows doc-engine's renumber_paragraphs to update numbers
                    # correctly when paragraphs are reordered or deleted later.
                    _inline_style = "font-family:Times New Roman;font-size:12pt;color:#000000;"
                    def _tn(t: str) -> dict:
                        return {"type": "text", "version": 1, "text": t, "format": 0,
                                "detail": 0, "mode": "normal", "style": _inline_style}
                    _new_lexical = {"root": {"type": "root", "version": 1, "children": [
                        {"type": "paragraph", "version": 1, "format": "", "indent": 0,
                         "direction": "ltr",
                         "children": [_tn(f"{_next_para_num}.\t\t\t\t"), _tn(_gen_text)]}
                    ]}}
                    _lexical_content = {"richtext": {"format": "lexical", "state": _new_lexical}}
                    _patch = await de_client.patch_section(
                        doc.docengine_doc_id, _new_sec_id, new_de_version, _lexical_content
                    )
                    new_de_version = _patch.get("version", new_de_version)
                    # Update the insert op so frontend sees the generated content in section.content,
                    # AND append a separate update_section op with new_lexical_state so the frontend
                    # can apply it the same way as a SET_FORMAT / REWRITE op.
                    for _u in updates:
                        if _u.get("op") == "insert" and _u.get("section", {}).get("id") == _new_sec_id:
                            _u.setdefault("section", {})["content"] = _lexical_content
                            break
                    updates.append({
                        "op": "update_section",
                        "section_id": _new_sec_id,
                        "new_lexical_state": _new_lexical,
                    })

                    # Renumber all paragraphs in document order so existing paragraphs that
                    # shifted position get updated numbers (e.g. old "3." becomes "4.").
                    try:
                        _fresh_renumber = await de_client.get_document(doc.docengine_doc_id)
                        _all_paras_renumber = [
                            s for s in ((_fresh_renumber.get("data") or {}).get("sections") or [])
                            if isinstance(s, dict) and s.get("type") == "paragraph"
                        ]
                        for _rpi, _rps in enumerate(_all_paras_renumber):
                            _expected = _rpi + 1
                            _rps_rt = ((_rps.get("content") or {}).get("richtext") or {})
                            _rps_lx = _rps_rt.get("state") if isinstance(_rps_rt.get("state"), dict) else None
                            if not _rps_lx:
                                continue
                            _rps_ch = ((_rps_lx.get("root") or {}).get("children") or [])
                            if not _rps_ch:
                                continue
                            _rps_para_ch = _rps_ch[0].get("children") or []
                            if not _rps_para_ch:
                                continue
                            _rps_first = _rps_para_ch[0]
                            if _rps_first.get("type") != "text":
                                continue
                            _rft = _rps_first.get("text", "")
                            _rft_m = re.match(r"^(\d+)\.", _rft)
                            if not _rft_m or int(_rft_m.group(1)) == _expected:
                                continue
                            # Number is wrong — patch it
                            _rps_first["text"] = f"{_expected}." + _rft[_rft_m.end():]
                            _rp = await de_client.patch_section(
                                doc.docengine_doc_id, _rps["id"], new_de_version,
                                {"richtext": {"format": "lexical", "state": _rps_lx}},
                            )
                            new_de_version = _rp.get("version", new_de_version)
                            updates.append({
                                "op": "update_section",
                                "section_id": _rps["id"],
                                "new_lexical_state": _rps_lx,
                            })
                    except Exception as _rne:
                        logging.warning("para renumber after add_para failed: %s", _rne)
                except DocEngineError:
                    pass  # Non-fatal: paragraph inserted but empty
        # --- End Way 1 ---

        # Update the local mirrored state with the new doc-engine version number.
        new_state = dict(cur_v.doc_state or {})
        new_state["docengine_version"] = new_de_version
        change_log = {
            "action": "command_blueprint",
            "request_id": request_id,
            "prompt": prompt,
            "action_object": action_obj.model_dump(),
            "de_action": de_action,
            "updates": updates,
            "intent_source": intent_source,
            "docengine_doc_id": doc.docengine_doc_id,
            "docengine_version_before": de_version,
            "docengine_version_after": new_de_version,
        }

        # Re-fetch current doc state from doc-engine so DOCX reflects latest sections.
        _cmd_doc_type = (doc.doc_type or "").upper().replace("-", "_")
        new_docx_path = cur_v.docx_path or crud.storage_path(doc.id, 0, "docx")
        fresh_structured: dict = {}
        try:
            de_doc = await de_client.get_document(doc.docengine_doc_id)
            fresh_data = de_doc.get("data") or new_state.get("structured") or {}
            # Merge _slots from previous state so render adapter can use them
            fresh_structured = dict(fresh_data)
            fresh_structured["_slots"] = (new_state.get("structured") or {}).get("_slots", {})
            new_state["structured"] = fresh_structured

            # Re-render DOCX with updated sections
            template_obj = (
                db.query(Template)
                .filter(Template.doc_type == _cmd_doc_type)
                .first()
            )
            if template_obj and os.path.exists(template_obj.docx_path):
                render_state = doc_state_from_filled_skeleton(fresh_structured, _cmd_doc_type)
                # Use version_id=0 as a stable "latest" slot for blueprint DOCX
                out_docx = crud.storage_path(doc.id, 0, "docx")
                crud.ensure_dir(os.path.dirname(out_docx))
                render_docx(template_obj.docx_path, render_state, out_docx)
                new_docx_path = out_docx
            else:
                # No legacy template — fall back to generate_plain_docx for blueprint docs
                from app.services.doc_importer import generate_plain_docx, sections_for_render
                _secs = sections_for_render(fresh_structured)
                out_docx = crud.storage_path(doc.id, 0, "docx")
                crud.ensure_dir(os.path.dirname(out_docx))
                generate_plain_docx(_secs, doc.title or "Document", out_docx)
                new_docx_path = out_docx
        except Exception as _re_exc:
            COMMAND_LOGGER.warning(
                "blueprint command re-render failed: %s", _re_exc, exc_info=True
            )

        try:
            new_version = crud.add_version(db, doc.id, new_state, change_log, new_docx_path)
        except Exception as ex:
            return _error(
                status_code=500,
                version=command_version,
                code="patch_apply_failed",
                message=f"Failed to persist blueprint command version: {ex}",
                details={},
            )

        resp = CommandAppliedResponse(
            status="applied",
            version=new_version.id,
            updates=updates,
            meta=_meta_payload(),
        )
        _emit(status="applied", version=new_version.id)
        result = resp.model_dump()
        result["docengine_version"] = new_de_version  # frontend needs this for PATCH /sections
        result["docx_url"] = f"/documents/{doc.id}/versions/{new_version.id}/download?format=docx"
        result["data"] = fresh_structured  # full sections so frontend can refresh Lexical editor
        return result
    # --- End blueprint path ---

    transform_started = time.perf_counter()
    try:
        updates, transform_meta = await plan_patch_ops_from_action(
            action_obj,
            structured,
            context=req.context.model_dump(),
        )
    except NeedsClarificationError as ex:
        transform_ms = int((time.perf_counter() - transform_started) * 1000)
        return _needs_clarification(
            version=command_version,
            reason_code="target_ambiguity",
            question=ex.question,
            options=ex.options,
        )
    except PlannerError as ex:
        transform_ms = int((time.perf_counter() - transform_started) * 1000)
        planner_message = str(ex)
        if planner_message.startswith("unsupported_action:"):
            return _needs_clarification(
                version=command_version,
                reason_code="unsupported_action",
                question=(
                    "That command is not supported in v1. "
                    "Try formal/concise tone, rewrite, shorten, or expand for a paragraph."
                ),
                options=[
                    {"label": "Make Para Formal", "token": "change_tone_formal"},
                    {"label": "Rewrite Paragraph", "token": "rewrite_content"},
                    {"label": "Shorten Paragraph", "token": "shorten_content"},
                    {"label": "Expand Paragraph", "token": "expand_content"},
                ],
            )
        return _error(
            status_code=500,
            version=command_version,
            code="patch_apply_failed",
            message=planner_message,
            details={},
        )
    except Exception as ex:
        transform_ms = int((time.perf_counter() - transform_started) * 1000)
        return _error(
            status_code=500,
            version=command_version,
            code="patch_apply_failed",
            message=f"Failed to build command updates: {ex}",
            details={},
        )
    transform_ms = int((time.perf_counter() - transform_started) * 1000)

    # UNDO shortcut — revert_to_previous is a sentinel that bypasses patch/transform.
    if updates and isinstance(updates[0], dict) and updates[0].get("op") == "revert_to_previous":
        prev_v = (
            db.query(DocumentVersion)
            .filter(DocumentVersion.document_id == doc.id, DocumentVersion.id < cur_v.id)
            .order_by(DocumentVersion.id.desc())
            .first()
        )
        if prev_v is None:
            return _error(
                status_code=422,
                version=command_version,
                code="patch_apply_failed",
                message="No previous version to revert to",
                details={},
            )
        undo_change_log = {
            "action": "UNDO",
            "request_id": request_id,
            "trace_id": request_id,
            "input_type": req.input.type.value,
            "input_source": input_source,
            "prompt": prompt,
            "action_object": action_obj.model_dump(),
            "reverted_from_version_id": cur_v.id,
            "reverted_to_version_id": prev_v.id,
        }
        try:
            new_version = crud.add_version(db, doc.id, prev_v.doc_state, undo_change_log, prev_v.docx_path)
        except Exception as ex:
            return _error(
                status_code=500,
                version=command_version,
                code="patch_apply_failed",
                message=f"Failed to persist undo version: {ex}",
                details={},
            )
        _undo_result = CommandAppliedResponse(
            status="applied",
            version=new_version.id,
            updates=[],
            meta=_meta_payload(),
        ).model_dump()
        _undo_structured = (prev_v.doc_state or {}).get("structured") or prev_v.doc_state or {}
        _undo_result["data"] = _undo_structured
        return _undo_result

    apply_started = time.perf_counter()
    try:
        updated_structured = apply_patch_ops(structured, updates)
    except Exception as ex:
        apply_ms = int((time.perf_counter() - apply_started) * 1000)
        return _error(
            status_code=500,
            version=command_version,
            code="patch_apply_failed",
            message=f"Failed to apply patch ops: {ex}",
            details={},
        )
    apply_ms = int((time.perf_counter() - apply_started) * 1000)

    # Persist mutation as a new version so command execution is auditable and replay-safe.
    new_state = dict(combined) if isinstance(combined, dict) else {}
    new_state["structured"] = updated_structured
    new_state["render"] = doc_state_from_filled_skeleton(updated_structured, doc.doc_type)
    latency_ms = _latency_payload()
    change_log = {
        "action": "command",
        "request_id": request_id,
        "trace_id": request_id,
        "input_type": req.input.type.value,
        "input_source": input_source,
        "transcript": transcript,
        "prompt": prompt,
        "action_object": action_obj.model_dump(),
        "updates": updates,
        "intent_source": intent_source,
        "repair_applied": repair_applied,
        "prompt_version": prompt_version,
        "stt_model": stt_meta.get("stt_model"),
        "stt_device": stt_meta.get("stt_device"),
        "stt_compute_type": stt_meta.get("stt_compute_type"),
        "stt_language_detected": stt_meta.get("stt_language_detected"),
        "stt_confidence": stt_meta.get("stt_confidence"),
        "stt_latency_ms": stt_meta.get("stt_latency_ms"),
        "auto_retried": auto_retried,
        "retry_count": retry_count,
        "base_version": requested_version,
        "transform_source": transform_meta.get("transform_source"),
        "transform_prompt_version": transform_meta.get("transform_prompt_version"),
        "transform_repair_applied": transform_meta.get("transform_repair_applied"),
        "latency_ms": latency_ms,
    }

    existing_docx = cur_v.docx_path or crud.storage_path(doc.id, 0, "docx")
    try:
        new_version = crud.add_version(db, doc.id, new_state, change_log, existing_docx)
    except Exception as ex:
        return _error(
            status_code=500,
            version=command_version,
            code="patch_apply_failed",
            message=f"Failed to persist command version: {ex}",
            details={},
        )

    resp = CommandAppliedResponse(
        status="applied",
        version=new_version.id,
        updates=updates,
        meta=_meta_payload(),
    )
    _emit(status="applied", version=new_version.id)
    result = resp.model_dump()
    # Include full updated structured data so the frontend can re-render all sections
    # immediately after any command (rewrite, expand, shorten, tone, etc.).
    result["data"] = updated_structured
    return result


@app.post("/documents/generate", response_model=BlueprintDocResponse)
async def create_blueprint_document(req: BlueprintDocRequest, db: Session = Depends(get_db)):
    """Create a new document from a template via the Doc-Engine microservice.

    The doc-engine resolves the template → blueprint → section catalog and returns
    a fully-structured document with Lexical JSON section content.

    The ML pipeline stores a lightweight Document row with a docengine_doc_id reference
    so /command can route to the correct document in subsequent requests.
    """
    client: DocEngineClient | None = app.state.docengine
    if client is None:
        raise HTTPException(503, "Doc-Engine integration is not enabled (DOCENGINE_ENABLED=false)")

    # ── Voice input: transcribe audio → text prompt ─────────────────────────────
    stt_transcript: str | None = None
    if req.input_type == "voice":
        if not settings.ENABLE_VOICE_INPUT:
            raise HTTPException(400, "Voice input is disabled in this backend")
        try:
            audio_bytes = base64.b64decode(req.audio_base64 or "", validate=True)
        except (binascii.Error, ValueError):
            raise HTTPException(400, "Could not decode audio payload")
        try:
            stt_transcript, _ = transcribe(audio_bytes=audio_bytes, mime_type=req.mime_type or "")
        except STTError as ex:
            raise HTTPException(400, f"Could not understand audio: {ex}")
        req = req.model_copy(update={"prompt": stt_transcript})

    # ── Resolve template_id to a doc-engine template ID string ─────────────────
    # Accepts three forms:
    #   "goi_letter"    — letter_type key (preferred, no DB needed)
    #   "tmpl_goi_001"  — doc-engine template ID directly (test UI)
    #   35              — integer ML pipeline templates.id PK (legacy frontend)
    ml_template_int_id: int | None = None
    de_template_id = str(req.template_id)

    # 1. letter_type key (e.g. "goi_letter") → look up in _UPLOAD_LETTER_TYPE_MAP
    if de_template_id in _UPLOAD_LETTER_TYPE_MAP:
        de_template_id, _ = _UPLOAD_LETTER_TYPE_MAP[de_template_id]

    # 2. integer PK → DB lookup → doc_type → de_template_id
    else:
        try:
            ml_template_int_id = int(req.template_id)
        except (ValueError, TypeError):
            pass  # already a doc-engine string ID like "tmpl_goi_001"

        if ml_template_int_id is not None:
            ml_tmpl_obj = db.get(Template, ml_template_int_id)
            if ml_tmpl_obj is None:
                raise HTTPException(404, f"ML template id={ml_template_int_id} not found in templates table")
            mapped = _DOC_TYPE_TO_DE_TEMPLATE.get(ml_tmpl_obj.doc_type.upper())
            if mapped is None:
                raise HTTPException(422, f"No doc-engine template mapping for doc_type={ml_tmpl_obj.doc_type!r}")
            de_template_id = mapped

    try:
        de_resp = await client.create_document(de_template_id, req.inputs)
    except DocEngineError as exc:
        raise HTTPException(exc.status_code, detail=str(exc.body))

    # Persist a Document row in the ML pipeline DB so future /command calls can look it up.
    doc_id = str(uuid.uuid4())
    doc = Document(
        id=doc_id,
        user_id=req.user_id,
        doc_type=de_resp.get("document_type", "unknown"),
        template_id=ml_template_int_id,   # int FK → templates.id (None for string IDs)
        current_version_id=None,
        docengine_doc_id=str(de_resp["document_id"]),
    )
    db.add(doc)
    db.flush()

    # Mirror the doc-engine data as a local DocumentVersion so the intent extractor
    # can read section IDs/types from doc_state["structured"] on subsequent /command calls.
    doc_data = de_resp.get("data") or {}
    de_ver = de_resp["version"]
    doc_type_upper = de_resp.get("document_type", "").upper().replace("-", "_")

    # ── Optional: LLM fill when a prompt is provided ────────────────────────────
    if req.prompt:
        from app.services.lexical_wrapper import text_to_lexical_node

        # 1. Extract structured slots via regex (fast, no extra LLM call)
        slots: dict = {}
        if doc_type_upper == "LEAVE_CERTIFICATE":
            from app.ml.slots.leave_certificate import _regex_fallback as _lc_regex
            slots = _lc_regex(req.prompt)
        elif doc_type_upper == "MOVEMENT_ORDER":
            from app.ml.slots.movement_order import _regex_fallback_mo
            slots = _regex_fallback_mo(req.prompt)
        elif doc_type_upper == "DO_LETTER":
            from app.ml.slots.do_letter import _regex_fallback_do
            slots = _regex_fallback_do(req.prompt)
        elif doc_type_upper in ("GOI_LETTER", "SERVICE_LETTER"):
            from app.ml.slots.goi_letter import _regex_fallback_goi
            slots = _regex_fallback_goi(req.prompt)
        elif doc_type_upper == "GENERAL_LETTER":
            from app.ml.slots.goi_letter import _regex_fallback_goi
            slots = _regex_fallback_goi(req.prompt)   # reuse for date + ref extraction

        # Auto-fill date with today if not extracted from prompt
        import datetime as _dt
        if not slots.get("date") and not slots.get("date_gregorian"):
            _event_date_str = slots.get("event_date", "")
            if _event_date_str and doc_type_upper == "DO_LETTER":
                # Set letter date to 7 days before the event date mentioned in prompt
                try:
                    _ev_dt = _dt.datetime.strptime(_event_date_str, "%d %b %Y").date()
                    _letter_dt = _ev_dt - _dt.timedelta(days=7)
                    slots["date"] = _letter_dt.strftime("%d %b %Y")
                except ValueError:
                    slots["date"] = _dt.date.today().strftime("%d %b %Y")
            else:
                slots["date"] = _dt.date.today().strftime("%d %b %Y")

        # Auto-generate subject if not in prompt
        if not slots.get("subject"):
            if doc_type_upper == "LEAVE_CERTIFICATE":
                slots["subject"] = "LEAVE CERTIFICATE"
            elif doc_type_upper == "MOVEMENT_ORDER":
                slots["subject"] = "MOVEMENT ORDER"
            elif doc_type_upper in ("DO_LETTER", "GOI_LETTER", "SERVICE_LETTER"):
                # Derive subject from prompt: strip leading doc-type boilerplate prefix
                _subj = re.sub(
                    r"^(?:write\s+(?:a\s+)?)?(?:service\s+letter|service|do|goi|demi[\s-]official|government\s+of\s+india)\s+letter\s*(?:for|about|to|on|regarding)?\s*",
                    "", req.prompt.strip(), flags=re.IGNORECASE,
                )
                # Strip trailing date / addressee fragments
                _subj = re.sub(r"\s+(?:dated?|on)\s+\d.*$", "", _subj, flags=re.IGNORECASE)
                # Only strip "to [the] [authority]" if it looks like an addressee (rank word or "the"),
                # NOT "to [place]" travel-route patterns like "from Pathankot to Leh"
                _subj = re.sub(
                    r"\s+to\s+(?:the\s+[A-Za-z]|(?:Lt\s+Gen|Maj\s+Gen|Brig|Col|Maj|Capt|Lt|Gen)\s+)[^\n]*$",
                    "", _subj, flags=re.IGNORECASE,
                ).strip()
                if _subj:
                    slots["subject"] = _subj.upper()[:200]
            elif doc_type_upper == "GENERAL_LETTER":
                _subj = re.sub(
                    r"^(?:write\s+(?:a\s+)?)?(?:general\s+letter|draft\s+(?:a\s+)?letter)\s*(?:for|about|on|regarding)?\s*",
                    "", req.prompt.strip(), flags=re.IGNORECASE,
                )
                _subj = re.sub(r"\s+(?:for|from|dated?|on|to)\s+.*$", "", _subj, flags=re.IGNORECASE).strip()
                if _subj:
                    slots["subject"] = _subj.upper()[:200]

        doc_data["_slots"] = slots

        # 2. Draft body paragraph text via LLM
        # Build a focused body topic: strip "do/goi letter to [rank] [name]" header so the
        # LLM receives the actual subject matter rather than echoing the instruction prefix.
        # Use the full intent phrase (e.g. "request for additional ration on army day") rather
        # than just the abbreviated subject slot so small models have enough context.
        _body_topic = req.prompt
        if doc_type_upper in ("DO_LETTER", "GOI_LETTER", "SERVICE_LETTER"):
            _bt = re.sub(
                r"^(?:write\s+(?:a\s+)?)?(?:service\s+letter|service|do|goi|demi[\s-]official|government\s+of\s+india)\s+letter\s*(?:for|about|to|on|regarding)?\s*",
                "", req.prompt.strip(), flags=re.IGNORECASE,
            ).strip()
            # Strip addressee "to [Rank] [Name]" but only when followed by a topic keyword
            _bt = re.sub(
                r"^to\s+(?:Lt\s+Col|Col|Brig|Maj\s+Gen|Lt\s+Gen|Gen|Maj|Capt|Lt)\s+[A-Za-z .]+?(?:,\s*[A-Za-z ,]+?)?\s+(?=about|on|for|regarding|dated\b)",
                "", _bt, flags=re.IGNORECASE,
            ).strip()
            if _bt:
                _body_topic = _bt
        body_texts: list = []
        try:
            if doc_type_upper in ("GOI_LETTER", "SERVICE_LETTER"):
                from app.ml.slots.goi_letter import draft_numbered_paras
                _date_for_body = slots.get("date") or slots.get("date_gregorian") or _dt.date.today().strftime("%d %b %Y")
                body_texts = await draft_numbered_paras(db, f"{_body_topic} [Letter date: {_date_for_body}]")
            elif doc_type_upper == "DO_LETTER":
                from app.ml.slots.do_letter import draft_body_paras
                body_texts = await draft_body_paras(db, _body_topic, min_paras=2, max_paras=3)
            elif doc_type_upper == "MOVEMENT_ORDER":
                # Slot-driven: build paragraph from extracted slots so Lexical shows real content
                from app.services.doc_importer import _build_movement_order_para
                body_texts = [_build_movement_order_para(slots)]
            elif doc_type_upper == "LEAVE_CERTIFICATE":
                # Slot-driven: build paragraph from extracted slots so Lexical shows real content
                from app.services.doc_importer import _build_leave_cert_para
                body_texts = [_build_leave_cert_para(slots)]
            elif doc_type_upper == "INVITATION_LETTER":
                from app.ml.slots.do_letter import draft_body_paras
                body_texts = await draft_body_paras(db, req.prompt, min_paras=2, max_paras=3)
            elif doc_type_upper == "GENERAL_LETTER":
                from app.ml.slots.general_letter import draft_body_paras_general, needs_to_whomsoever
                _is_cert = needs_to_whomsoever(req.prompt)
                body_texts = await draft_body_paras_general(db, _body_topic, is_certificate=_is_cert)
        except Exception:
            import traceback as _tb
            COMMAND_LOGGER.warning("body para draft failed doc_type=%s: %s", doc_type_upper, _tb.format_exc())
            body_texts = []

        # Build para_sections BEFORE step 3a so that the subject paragraph inserted
        # in step 3a is NOT included — body_texts[i] must map to body paragraphs only.
        para_sections = [s for s in (doc_data.get("sections") or []) if isinstance(s, dict) and s.get("type") == "paragraph"]
        _last_gen_para_id: str | None = para_sections[-1]["id"] if para_sections else None

        # 3a. DO/GOI letter: INSERT subject as a "subject"-type section after receiver_block.
        # We now refresh doc_data from the doc-engine after all operations, so the correct
        # position (order=4, after receiver_block=3) is preserved in the API response.
        # Using "subject" type means "para N" commands won't accidentally target it.
        de_doc_id = str(de_resp["document_id"])
        if doc_type_upper in ("DO_LETTER", "GOI_LETTER", "SERVICE_LETTER"):
            _subj_text = (slots.get("subject") or "").strip().upper()
            if _subj_text and client is not None:
                try:
                    _recv_id = next(
                        (s.get("id") for s in (doc_data.get("sections") or [])
                         if isinstance(s, dict) and s.get("type") == "receiver_block"),
                        None,
                    )
                    _subj_ins = await client.apply_command(de_doc_id, de_ver, {
                        "action": "INSERT_SECTION",
                        "target_type": "subject",
                        "target_ref": None,
                        "position": ({"policy": "after", "section_id": _recv_id}
                                     if _recv_id else {"policy": "start"}),
                        "ai_instruction": None,
                    })
                    if _subj_ins.get("status") == "applied":
                        de_ver = _subj_ins.get("version", de_ver)
                        _subj_sec_id = next(
                            (u.get("section", {}).get("id") for u in (_subj_ins.get("updates") or [])
                             if u.get("op") == "insert"), None,
                        )
                        if _subj_sec_id:
                            _subj_lx = text_to_lexical_node(_subj_text, bold=True, underline=True, align="center")
                            _subj_lx_c = {"richtext": {"format": "lexical", "state": _subj_lx}}
                            _sp = await client.patch_section(de_doc_id, _subj_sec_id, de_ver, _subj_lx_c)
                            de_ver = _sp.get("version", de_ver)
                            doc_data.setdefault("sections", []).append(
                                {"id": _subj_sec_id, "type": "subject", "content": _subj_lx_c}
                            )
                except Exception as _se:
                    COMMAND_LOGGER.warning("blueprint fill: subject para insert failed: %s", _se)

        # 3b. PATCH paragraph sections with drafted body text
        for i, sec in enumerate(para_sections):
            if i >= len(body_texts):
                break
            _para_text = body_texts[i]
            lx = text_to_lexical_node(_para_text)
            lx_content = {"richtext": {"format": "lexical", "state": lx}}
            try:
                patch = await client.patch_section(de_doc_id, sec["id"], de_ver, lx_content)
                de_ver = patch["version"]
                sec["content"] = lx_content
                _last_gen_para_id = sec["id"]
            except DocEngineError as _pe:
                COMMAND_LOGGER.warning("blueprint fill: patch para failed sec=%s err=%s", sec["id"], _pe)

        # INSERT extra paragraphs when LLM drafted more than the blueprint's min_instances.
        # (e.g. bp_service_letter_v1 has min_instances paragraph=1 but LLM returns 2-3)
        for extra_text in body_texts[len(para_sections):]:
            try:
                _pos = ({"policy": "after", "section_id": _last_gen_para_id}
                        if _last_gen_para_id else {"policy": "end"})
                _ins = await client.apply_command(de_doc_id, de_ver, {
                    "action": "INSERT_SECTION",
                    "target_type": "paragraph",
                    "target_ref": None,
                    "position": _pos,
                    "ai_instruction": None,
                })
                if _ins.get("status") == "applied":
                    de_ver = _ins.get("version", de_ver)
                    _new_sid = next(
                        (u.get("section", {}).get("id") for u in (_ins.get("updates") or [])
                         if u.get("op") == "insert"),
                        None,
                    )
                    if _new_sid:
                        lx = text_to_lexical_node(extra_text)
                        lx_content = {"richtext": {"format": "lexical", "state": lx}}
                        _ep = await client.patch_section(de_doc_id, _new_sid, de_ver, lx_content)
                        de_ver = _ep.get("version", de_ver)
                        _last_gen_para_id = _new_sid
                        doc_data.setdefault("sections", []).append(
                            {"id": _new_sid, "type": "paragraph", "content": lx_content}
                        )
            except DocEngineError as _pe:
                COMMAND_LOGGER.warning("blueprint fill: insert extra para failed err=%s", _pe)

        # 4. PATCH single-value sections (reference_number, date, subject, receiver_block, signee_block) from slots
        _a1 = slots.get("addressee_1", "")
        _a2 = slots.get("addressee_2", "")
        _recv = "\n".join(l for l in [_a1, _a2] if l.strip())
        _sn = slots.get("signatory_name", "")
        _sd = slots.get("signatory_designation", "")
        _sdept = slots.get("signatory_dept", "")
        if doc_type_upper == "DO_LETTER" and _sn:
            # DO letter: "Yours sincerely,\n(Name)\nRank\nUnit"
            # signatory_name may contain rank prefix: "Lt Gen Anil Kumar" → rank="Lt Gen", name="Anil Kumar"
            _rank_prefix = re.match(
                r"^(Lt\s+Gen|Maj\s+Gen|Lt\s+Col|Col|Brig|Maj|Capt|Lt|Gen)\s+",
                _sn, re.IGNORECASE,
            )
            if _rank_prefix:
                _rank_str = _rank_prefix.group(1)
                _name_only = _sn[_rank_prefix.end():].strip()
            else:
                _rank_str = ""
                _name_only = _sn
            _signee = "\n".join(l for l in [
                "Yours sincerely,",
                f"({_name_only})" if _name_only else "",
                _rank_str,
                _sd,
                _sdept,
            ] if l.strip())
        elif doc_type_upper == "GOI_LETTER":
            # GOI letter: "Yours faithfully,\nAppointment title\nfor Ministry of Defence"
            _signee = "\n".join(l for l in [
                "Yours faithfully,",
                _sd,
                f"for {_sdept}" if _sdept else "",
            ] if l.strip())
        else:
            _signee = "\n".join(l for l in [_sn, _sd, _sdept] if l.strip())
        # Slot-driven types: army_no/date are embedded in paragraph body.
        # Movement order uses the reference_number section for "In lieu of IAFT-1759"
        # (right-aligned, above title) which is a fixed form identifier.
        _slot_driven = doc_type_upper in ("LEAVE_CERTIFICATE", "MOVEMENT_ORDER")
        single_patch = {
            "reference_number": "In lieu of IAFT-1759" if doc_type_upper == "MOVEMENT_ORDER"
                                 else ("" if _slot_driven
                                       else (slots.get("file_reference_number") or slots.get("army_no") or "")),
            "date": "" if _slot_driven
                     else (slots.get("date") or slots.get("date_gregorian") or ""),
            "subject": slots.get("subject") or "",
            "receiver_block": _recv,
            "signee_block": _signee,
        }
        for sec in (doc_data.get("sections") or []):
            if not isinstance(sec, dict):
                continue
            sec_type = sec.get("type", "")
            if sec_type not in single_patch:
                continue
            val = single_patch[sec_type]
            if val is not None:  # patch even empty string to clear blueprint defaults
                _is_subj    = sec_type == "subject"
                _is_mo_ref  = (sec_type == "reference_number"
                               and doc_type_upper == "MOVEMENT_ORDER")
                _is_mo_sign = (sec_type == "signee_block"
                               and doc_type_upper == "MOVEMENT_ORDER")
                lx = text_to_lexical_node(
                    val,
                    bold=_is_subj or _is_mo_ref,
                    underline=_is_subj or _is_mo_ref,
                    align=("center" if _is_subj
                           else "right" if (_is_mo_ref or _is_mo_sign)
                           else ""),
                )
                lx_content = {"richtext": {"format": "lexical", "state": lx}}
                try:
                    patch = await client.patch_section(de_doc_id, sec["id"], de_ver, lx_content)
                    de_ver = patch["version"]
                    sec["content"] = lx_content
                except DocEngineError as _pe:
                    COMMAND_LOGGER.warning("blueprint fill: patch single sec=%s type=%s err=%s", sec["id"], sec.get("type"), _pe)

        # 4b. For leave cert remove the reference_number section (army_no is in body).
        # Movement order keeps it — used to display "In lieu of IAFT-1759" above title.
        if doc_type_upper == "LEAVE_CERTIFICATE":
            doc_data["sections"] = [
                s for s in (doc_data.get("sections") or [])
                if not (isinstance(s, dict) and s.get("type") == "reference_number")
            ]

        # 5. If signee_block not created by doc-engine (optional), insert + patch it
        _has_signee = any(s.get("type") == "signee_block" for s in (doc_data.get("sections") or []) if isinstance(s, dict))
        if not _has_signee and _signee and client is not None:
            try:
                # Always insert at end — avoids order-rule violation that occurs when
                # inserting after a paragraph that is not the last section in the
                # doc-engine document (e.g. when the subject paragraph was inserted
                # after receiver_block, making it appear earlier than body paras).
                _ins_result = await client.apply_command(
                    de_doc_id, de_ver,
                    {
                        "action": "INSERT_SECTION",
                        "target_type": "signee_block",
                        "target_ref": None,
                        "position": {"policy": "end"},
                        "ai_instruction": None,
                    },
                )
                if _ins_result.get("status") == "applied":
                    de_ver = _ins_result.get("version", de_ver)
                    _new_secs = (_ins_result.get("updates") or [])
                    _new_sec_id = next(
                        (u.get("section", {}).get("id") for u in _new_secs if u.get("op") == "insert"),
                        None,
                    )
                    if _new_sec_id:
                        _sign_align = "right" if doc_type_upper == "MOVEMENT_ORDER" else ""
                        lx = text_to_lexical_node(_signee, align=_sign_align)
                        lx_content = {"richtext": {"format": "lexical", "state": lx}}
                        patch = await client.patch_section(de_doc_id, _new_sec_id, de_ver, lx_content)
                        de_ver = patch.get("version", de_ver)
                        doc_data.setdefault("sections", []).append({
                            "id": _new_sec_id, "type": "signee_block", "content": lx_content
                        })
            except Exception as _e:
                COMMAND_LOGGER.warning("blueprint fill: signee_block insert failed: %s", _e)

        # 5b. If subject not created by doc-engine (optional in some blueprints like service_letter),
        # insert + patch it when we have a subject slot value.
        _has_subject = any(s.get("type") == "subject" for s in (doc_data.get("sections") or []) if isinstance(s, dict))
        _subj_val = (locals().get("slots") or {}).get("subject") or ""
        # DO_LETTER / GOI_LETTER / SERVICE_LETTER: subject inserted in step 3a — skip duplicate INSERT
        if not _has_subject and _subj_val and client is not None and doc_type_upper not in ("DO_LETTER", "GOI_LETTER", "SERVICE_LETTER"):
            try:
                # Find receiver_block section ID to insert subject after it
                _recv_sec_id = next(
                    (s.get("id") for s in (doc_data.get("sections") or [])
                     if isinstance(s, dict) and s.get("type") == "receiver_block"),
                    None,
                )
                _subj_pos = (
                    {"policy": "after", "section_id": _recv_sec_id}
                    if _recv_sec_id
                    else {"policy": "after", "section_type": "paragraph"}
                )
                _ins_subj = await client.apply_command(
                    de_doc_id, de_ver,
                    {
                        "action": "INSERT_SECTION",
                        "target_type": "subject",
                        "target_ref": None,
                        "position": _subj_pos,
                        "ai_instruction": None,
                    },
                )
                if _ins_subj.get("status") == "applied":
                    de_ver = _ins_subj.get("version", de_ver)
                    _new_subj_id = next(
                        (u.get("section", {}).get("id") for u in (_ins_subj.get("updates") or []) if u.get("op") == "insert"),
                        None,
                    )
                    if _new_subj_id:
                        lx = text_to_lexical_node(_subj_val, bold=True, underline=True, align="center")
                        lx_content = {"richtext": {"format": "lexical", "state": lx}}
                        patch = await client.patch_section(de_doc_id, _new_subj_id, de_ver, lx_content)
                        de_ver = patch.get("version", de_ver)
                        doc_data.setdefault("sections", []).append({
                            "id": _new_subj_id, "type": "subject", "content": lx_content
                        })
            except Exception as _e:
                COMMAND_LOGGER.warning("blueprint fill: subject insert failed: %s", _e)

        # 5c. GENERAL_LETTER certificates: INSERT "To Whomsoever It May Concern" salutation after subject
        if doc_type_upper == "GENERAL_LETTER" and client is not None:
            from app.ml.slots.general_letter import needs_to_whomsoever
            if needs_to_whomsoever(req.prompt):
                _has_sal = any(s.get("type") == "salutation"
                               for s in (doc_data.get("sections") or []) if isinstance(s, dict))
                if not _has_sal:
                    try:
                        _sal_ins = await client.apply_command(de_doc_id, de_ver, {
                            "action": "INSERT_SECTION", "target_type": "salutation",
                            "target_ref": None,
                            "position": {"policy": "after", "section_type": "subject"},
                            "ai_instruction": None,
                        })
                        if _sal_ins.get("status") == "applied":
                            de_ver = _sal_ins.get("version", de_ver)
                            _sal_id = next(
                                (u.get("section", {}).get("id") for u in (_sal_ins.get("updates") or [])
                                 if u.get("op") == "insert"), None,
                            )
                            if _sal_id:
                                _sal_lx = text_to_lexical_node(
                                    "To Whomsoever It May Concern", bold=True, underline=True, align="center"
                                )
                                _sal_c = {"richtext": {"format": "lexical", "state": _sal_lx}}
                                _sp = await client.patch_section(de_doc_id, _sal_id, de_ver, _sal_c)
                                de_ver = _sp.get("version", de_ver)
                                doc_data.setdefault("sections", []).append(
                                    {"id": _sal_id, "type": "salutation", "content": _sal_c})
                    except Exception as _e:
                        COMMAND_LOGGER.warning("blueprint fill: salutation insert failed: %s", _e)

    # ────────────────────────────────────────────────────────────────────────────

    # Refresh doc_data from the doc-engine so sections are in their actual blueprint
    # order (INSERTs like subject and signee_block were appended to the local list,
    # but the doc-engine stores them in the correct position).
    _de_doc_id_for_refresh = locals().get("de_doc_id") or str(de_resp.get("document_id", ""))
    if _de_doc_id_for_refresh and client is not None:
        try:
            _fresh = await client.get_document(_de_doc_id_for_refresh)
            _fresh_data = _fresh.get("data") or {}
            if _fresh_data.get("sections"):
                _fresh_data["_slots"] = doc_data.get("_slots", {})
                doc_data = _fresh_data
                de_ver = _fresh.get("version", de_ver)
        except Exception as _re:
            COMMAND_LOGGER.warning("blueprint fill: doc refresh failed: %s", _re)

    # Auto-generate human-readable title from slots (subject) or prompt, set on doc row.
    _slots_for_title: dict = locals().get("slots") or {}
    doc.title = _derive_title(doc_type_upper, _slots_for_title, req.prompt)

    initial_state = {
        "structured": doc_data,   # sections list (+ _slots) used by intent extractor + render_adapter
        "docengine_version": de_ver,
    }
    existing_docx = crud.storage_path(doc_id, 0, "docx")
    ml_version_row = crud.add_version(db, doc_id, initial_state, {"action": "blueprint_create", "template_id": req.template_id}, existing_docx)

    # Render DOCX immediately so download works after creation
    template_obj = (
        db.query(Template)
        .filter(Template.doc_type == doc_type_upper)
        .order_by(Template.id.desc())
        .first()
    ) if doc_type_upper else None

    # Also try exact doc_type match
    if not template_obj and doc_type_upper:
        _dt_variants = [
            doc_type_upper.lower().replace("_", ""),   # "goiletter"
            doc_type_upper.lower(),                     # "goi_letter"
            de_resp.get("document_type", ""),           # original from doc-engine
        ]
        for _dt in _dt_variants:
            template_obj = (
                db.query(Template)
                .filter(Template.doc_type == _dt)
                .order_by(Template.id.desc())
                .first()
            )
            if template_obj:
                break

    # Prefix LIKE fallback: "DO_LETTER" matches "DO_LETTER_TEMPLATE_DOCXTPL_V2"
    if not template_obj and doc_type_upper:
        template_obj = (
            db.query(Template)
            .filter(Template.doc_type.like(f"{doc_type_upper}_%"))
            .order_by(Template.id.desc())
            .first()
        )

    render_warning: str | None = None
    # Treat templates with no valid docx_path as missing (slot-driven types skip docxtpl)
    if template_obj and not (template_obj.docx_path and os.path.exists(template_obj.docx_path)):
        template_obj = None

    if template_obj:
        try:
            render_state = doc_state_from_filled_skeleton(doc_data, doc_type_upper)
            out_docx = crud.storage_path(doc_id, ml_version_row.id, "docx")
            render_docx(template_obj.docx_path, render_state, out_docx)
            ml_version_row.docx_path = out_docx

            # For slot-driven types: sync rendered DOCX text → Lexical paragraph section
            # so the editor shows exactly what the downloaded DOCX contains.
            if doc_type_upper in ("LEAVE_CERTIFICATE", "MOVEMENT_ORDER") and client:
                try:
                    from docx import Document as _DocxDoc
                    from app.services.doc_importer import extract_rich_lexical_body
                    _rdoc = _DocxDoc(out_docx)
                    _subj_val = (doc_data.get("_slots") or {}).get("subject", "").upper()
                    _bare_title = _subj_val.split(" - ")[0].strip()
                    _skip_titles = {t for t in (_subj_val, _bare_title) if t}
                    # For movement order, also skip "In lieu of IAFT-1759" from body
                    # (it is placed in the reference_number section above the title)
                    if doc_type_upper == "MOVEMENT_ORDER":
                        _skip_titles.add("IN LIEU OF IAFT-1759")
                    # Only read col-0 cells from tables (col-1 has signatory — skip it)
                    _raw_paras = list(_rdoc.paragraphs)
                    for _tbl in _rdoc.tables:
                        for _row in _tbl.rows:
                            if _row.cells:
                                _raw_paras.extend(_row.cells[0].paragraphs)
                    # Build skip set for signatory content
                    _lc_slots = doc_data.get("_slots") or {}
                    _sig_name  = (_lc_slots.get("signatory_name") or "").strip()
                    _sig_desig = (_lc_slots.get("signatory_designation") or "").strip()
                    _skip_sig = {"()"}
                    if _sig_name:
                        _skip_sig.update({f"({_sig_name})", _sig_name})
                    if _sig_desig:
                        _skip_sig.add(_sig_desig)
                    # Build rich Lexical state from DOCX runs (preserves bold/underline)
                    _rich_lx, _distr_lx = extract_rich_lexical_body(
                        _raw_paras, _skip_titles, _skip_sig
                    )
                    _para_secs = [
                        s for s in (doc_data.get("sections") or [])
                        if isinstance(s, dict) and s.get("type") == "paragraph"
                    ]
                    if _para_secs and _rich_lx:
                        _lx_c = {"richtext": {"format": "lexical", "state": _rich_lx}}
                        _pp = await client.patch_section(
                            de_doc_id, _para_secs[0]["id"], de_ver, _lx_c
                        )
                        de_ver = _pp.get("version", de_ver)
                        _para_secs[0]["content"] = _lx_c
                    # Patch distribution_list section with Distr lines from DOCX
                    if _distr_lx:
                        _distr_secs = [
                            s for s in (doc_data.get("sections") or [])
                            if isinstance(s, dict) and s.get("type") == "distribution_list"
                        ]
                        if not _distr_secs:
                            # Section doesn't exist yet — INSERT it after signee_block or at end
                            _signee_sec_id = next(
                                (s.get("id") for s in (doc_data.get("sections") or [])
                                 if isinstance(s, dict) and s.get("type") == "signee_block"),
                                None,
                            )
                            _distr_pos = (
                                {"policy": "after", "section_id": _signee_sec_id}
                                if _signee_sec_id else {"policy": "end"}
                            )
                            _ins_d = await client.apply_command(
                                de_doc_id, de_ver,
                                {"action": "INSERT_SECTION", "target_type": "distribution_list",
                                 "target_ref": None, "position": _distr_pos, "ai_instruction": None},
                            )
                            if _ins_d.get("status") == "applied":
                                de_ver = _ins_d.get("version", de_ver)
                                _new_distr_id = next(
                                    (u.get("section", {}).get("id") for u in (_ins_d.get("updates") or [])
                                     if u.get("op") == "insert"), None
                                )
                                if _new_distr_id:
                                    _distr_secs = [{"id": _new_distr_id, "type": "distribution_list"}]
                                    doc_data.setdefault("sections", []).append(_distr_secs[0])
                        if _distr_secs:
                            _dlx_c = {"richtext": {"format": "lexical", "state": _distr_lx}}
                            _dp = await client.patch_section(
                                de_doc_id, _distr_secs[0]["id"], de_ver, _dlx_c
                            )
                            de_ver = _dp.get("version", de_ver)
                            _distr_secs[0]["content"] = _dlx_c
                    ml_version_row.doc_state = {
                        **(ml_version_row.doc_state or {}),
                        "structured": doc_data,
                        "docengine_version": de_ver,
                    }
                except Exception as _sync_exc:
                    COMMAND_LOGGER.warning(
                        "post-render DOCX sync failed doc_type=%s: %s", doc_type_upper, _sync_exc
                    )

            db.add(ml_version_row)
            db.commit()
        except Exception as _render_exc:
            render_warning = f"DOCX render failed: {_render_exc}"
            logging.warning("Blueprint render failed for %s: %s", doc_id, _render_exc)
            db.rollback()
            template_obj = None  # fall through to plain render below

    if not template_obj:
        # No valid docxtpl template — generate DOCX from slots (slot-driven types) or
        # blueprint sections (GOI/DO/service letter).
        try:
            from app.services.doc_importer import (
                generate_plain_docx, sections_for_render, generate_slot_docx,
            )
            out_docx = crud.storage_path(doc_id, ml_version_row.id, "docx")
            crud.ensure_dir(os.path.dirname(out_docx))
            if not generate_slot_docx(doc_data, doc_type_upper, doc.title or "Document", out_docx):
                secs = sections_for_render(doc_data)
                generate_plain_docx(secs, doc.title or "Document", out_docx)
            ml_version_row.docx_path = out_docx
            db.add(ml_version_row)
            db.commit()
            render_warning = None
        except Exception as _plain_exc:
            render_warning = f"DOCX render failed: {_plain_exc}"
            logging.warning("Blueprint plain render failed for %s: %s", doc_id, _plain_exc)
            db.rollback()

    return BlueprintDocResponse(
        document_id=doc_id,
        docengine_doc_id=str(de_resp["document_id"]),
        version=ml_version_row.id,              # ML current_version_id — send to /command
        docengine_version=de_ver,               # doc-engine version — send to PATCH /sections
        doc_type=de_resp.get("document_type", ""),
        blueprint_id=de_resp.get("blueprint_id", ""),
        filled=doc_data.get("_slots") or {},
        data=doc_data,
        render_warning=render_warning,
    )


@app.patch("/documents/{doc_id}/sections/{section_id}", response_model=SectionPatchResponse)
async def patch_blueprint_section(
    doc_id: str,
    section_id: str,
    req: SectionPatchRequest,
    db: Session = Depends(get_db),
):
    """Replace the Lexical JSON content of a single section in a blueprint document.

    Used for SET_FORMAT (bold/italic/etc.) and direct section content edits.
    Proxies to the doc-engine PATCH /documents/{id}/sections/{section_id} endpoint
    with optimistic version locking.
    """
    client: DocEngineClient | None = app.state.docengine
    if client is None:
        raise HTTPException(503, "Doc-Engine integration is not enabled (DOCENGINE_ENABLED=false)")

    doc = db.get(Document, doc_id)
    if not doc or not doc.docengine_doc_id:
        raise HTTPException(404, f"Blueprint document not found: {doc_id}")

    try:
        result = await client.patch_section(
            doc_id=doc.docengine_doc_id,
            section_id=section_id,
            version=req.version,
            content=req.content,
        )
    except DocEngineError as exc:
        if exc.status_code == 409:
            raise HTTPException(409, "Version conflict — client version is stale")
        if exc.status_code == 404:
            raise HTTPException(404, f"Section not found: {section_id}")
        raise HTTPException(exc.status_code if exc.status_code >= 400 else 500, str(exc.body))

    new_de_version = result.get("version", req.version)

    # Update the mirrored local version to track the new doc-engine version.
    if doc.current_version_id:
        cur_v = db.get(DocumentVersion, doc.current_version_id)
        if cur_v:
            new_state = dict(cur_v.doc_state or {})
            new_state["docengine_version"] = new_de_version
            change_log = {
                "action": "patch_section",
                "section_id": section_id,
                "docengine_doc_id": doc.docengine_doc_id,
                "docengine_version_before": req.version,
                "docengine_version_after": new_de_version,
            }
            existing_docx = cur_v.docx_path or crud.storage_path(doc.id, 0, "docx")
            crud.add_version(db, doc.id, new_state, change_log, existing_docx)

    return SectionPatchResponse(ok=True, version=new_de_version)


# ---------------------------------------------------------------------------
# POST /documents/{doc_id}/checkpoint
# Called when the user clicks "Save" / "Done editing" in the frontend.
# Pulls fresh structured state from the doc-engine, re-renders the DOCX,
# and creates a new ML version so downstream commands see the latest content.
# ---------------------------------------------------------------------------

@app.post("/documents/{doc_id}/checkpoint")
async def checkpoint_document(doc_id: str, db: Session = Depends(get_db)):
    """Create a new ML version from the current doc-engine state.

    Intended to be called after a batch of manual edits so that:
    - doc_state["structured"] is refreshed with the latest section content
    - A new ML version ID is minted (version counter increments)
    - The DOCX export reflects the latest edits
    """
    client: DocEngineClient | None = app.state.docengine
    if client is None:
        raise HTTPException(503, "Doc-Engine integration is not enabled")

    doc = db.get(Document, doc_id)
    if not doc or not doc.docengine_doc_id:
        raise HTTPException(404, f"Document not found: {doc_id}")

    try:
        de_data = await client.get_document(doc.docengine_doc_id)
    except DocEngineError as exc:
        raise HTTPException(exc.status_code or 500, f"Doc-Engine error: {exc.body}")

    new_de_version = de_data.get("version", 0)
    fresh_structured = de_data.get("data") or de_data  # normalise response shape

    cur_v = db.get(DocumentVersion, doc.current_version_id) if doc.current_version_id else None
    old_state = dict(cur_v.doc_state or {}) if cur_v else {}
    new_state = {**old_state, "structured": fresh_structured, "docengine_version": new_de_version}

    fallback_docx = (cur_v.docx_path if cur_v else None) or crud.storage_path(doc.id, 0, "docx")
    new_docx_path = _render_blueprint_docx_inline(
        doc.id, fresh_structured,
        doc.title or "Document", fallback_docx,
        doc_type=doc.doc_type or "", db=db,
        template_id=doc.template_id,
    )

    change_log = {
        "action": "checkpoint",
        "docengine_doc_id": doc.docengine_doc_id,
        "docengine_version": new_de_version,
    }
    new_version = crud.add_version(db, doc.id, new_state, change_log, new_docx_path)
    db.commit()

    return {
        "ok": True,
        "version": new_version.id,
        "docengine_version": new_de_version,
        "docx_url": f"/documents/{doc.id}/versions/{new_version.id}/download?format=docx",
    }


# ---------------------------------------------------------------------------
# POST /stt/transcribe — pre-transcribe audio to text
# Used by UI panels that can't send JSON body (from-template uses query params)
# ---------------------------------------------------------------------------

from pydantic import BaseModel as _BaseModel

class _TranscribeRequest(_BaseModel):
    audio_base64: str
    mime_type: str


@app.post("/stt/transcribe")
async def stt_transcribe(req: _TranscribeRequest):
    """Convert base64-encoded audio to plain text transcript.

    Returns: { transcript: str, stt_model, stt_device, stt_confidence, stt_latency_ms, ... }
    """
    import base64 as _b64
    from app.services.stt import transcribe as _transcribe, STTError
    try:
        audio_bytes = _b64.b64decode(req.audio_base64, validate=True)
        transcript, meta = _transcribe(audio_bytes=audio_bytes, mime_type=req.mime_type)
        return {"transcript": transcript, **meta}
    except STTError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Transcription failed: {exc}")


_DOC_TYPE_DISPLAY: dict[str, str] = {
    "GOI_LETTER": "GOI Letter",
    "DO_LETTER": "DO Letter",
    "LEAVE_CERTIFICATE": "Leave Certificate",
    "MOVEMENT_ORDER": "Movement Order",
    "INVITATION_LETTER": "Invitation Letter",
    "GENERAL_LETTER": "General Letter",
}


def _derive_title(doc_type: str, slots: dict, prompt: str) -> str:
    """Derive a human-readable document title from slots/prompt, falling back to doc_type."""
    subj = (slots.get("subject") or "").strip()
    if subj:
        # Strip "SUBJECT:" prefix if present, title-case, truncate.
        subj = re.sub(r"^SUBJECT\s*:\s*", "", subj, flags=re.IGNORECASE).strip()
        return subj[:120] or subj
    if prompt:
        truncated = prompt.strip()[:80].rsplit(" ", 1)[0].strip(" ,;:")
        return truncated or prompt[:80]
    return _DOC_TYPE_DISPLAY.get(doc_type.upper(), doc_type.replace("_", " ").title())


@app.get("/documents", response_model=list[DocumentListItem])
def list_documents(
    user_id: str = "local-user",
    limit: int = 50,
    db: Session = Depends(get_db),
):
    """List all documents for a user, newest first."""
    docs = (
        db.query(Document)
        .filter(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
        .limit(max(1, min(limit, 200)))
        .all()
    )
    return [
        DocumentListItem(
            document_id=d.id,
            title=d.title or _DOC_TYPE_DISPLAY.get(d.doc_type.upper(), d.doc_type.replace("_", " ").title()),
            doc_type=d.doc_type,
            created_at=d.created_at,
            current_version_id=d.current_version_id,
            preview_url=f"/documents/{d.id}/preview",
        )
        for d in docs
    ]


@app.get("/documents/{document_id}", response_model=DocumentSummaryOut)
def get_document(document_id: str, db: Session = Depends(get_db)):
    doc = db.get(Document, document_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    de_ver = None
    _cur_data: dict | None = None
    if doc.current_version_id:
        cur_v = db.get(DocumentVersion, doc.current_version_id)
        if cur_v and isinstance(cur_v.doc_state, dict):
            de_ver = cur_v.doc_state.get("docengine_version")
            _s = cur_v.doc_state.get("structured") or cur_v.doc_state
            if isinstance(_s, dict) and _s.get("sections"):
                _cur_data = _s
    ver_count = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == doc.id)
        .count()
    )
    return DocumentSummaryOut(
        document_id=doc.id,
        doc_type=doc.doc_type,
        current_version_id=doc.current_version_id,
        template_id=doc.template_id,
        docengine_version=de_ver,
        doc_version=max(0, ver_count - 1),  # 0-based: generate=v0, first edit=v1, …
        data=_cur_data,
        preview_url=f"/documents/{doc.id}/preview",
    )


@app.get("/documents/{document_id}/versions", response_model=list[VersionHistoryItem])
def list_document_versions(document_id: str, db: Session = Depends(get_db)):
    doc = db.get(Document, document_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    versions = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == document_id)
        .order_by(DocumentVersion.id.asc())
        .all()
    )

    def _preview(change_log: dict) -> str:
        text = (change_log or {}).get("prompt") or (change_log or {}).get("edit_prompt") or ""
        text = " ".join(str(text).split())
        return text if len(text) <= 160 else text[:157] + "..."

    def _version_data(v) -> dict | None:
        """Extract structured sections from a version's doc_state for preview rendering."""
        try:
            state = v.doc_state or {}
            structured = state.get("structured") or state
            if isinstance(structured, dict) and structured.get("sections"):
                return structured
        except Exception:
            pass
        return None

    return [
        VersionHistoryItem(
            version_id=v.id,
            thread_version=idx,
            created_at=v.created_at,
            prompt_preview=_preview(v.change_log),
            docx_url=f"/documents/{document_id}/versions/{v.id}/download?format=docx",
            pdf_url=f"/documents/{document_id}/versions/{v.id}/download?format=pdf",
            preview_url=f"/documents/{document_id}/preview?version_id={v.id}",
            data=_version_data(v),
        )
        for idx, v in enumerate(versions, start=1)
    ]

# Map user-visible letter_type values → (doc-engine template_id, ML doc_type for rendering).
# letter_types NOT in this map use tmpl_flex_001 + generate_plain_docx fallback.
_UPLOAD_LETTER_TYPE_MAP: dict[str, tuple[str, str]] = {
    "goi_letter":        ("tmpl_goi_001",   "GOI_LETTER"),
    "do_letter":         ("tmpl_do_001",    "DO_LETTER"),
    "movement_order":    ("tmpl_mov_001",   "MOVEMENT_ORDER"),
    "leave_certificate": ("tmpl_leave_001", "LEAVE_CERTIFICATE"),
    "invitation_letter": ("tmpl_inv_001",   "INVITATION_LETTER"),
    "service_letter":    ("tmpl_svc_001",   "SERVICE_LETTER"),
    "general_letter":    ("tmpl_gen_001",   "GENERAL_LETTER"),
}

# Reverse map: ML doc_type (uppercase) → doc-engine template_id string.
# Used when the production frontend sends an integer ML template ID instead of
# the doc-engine string ID (e.g. template_id=35 instead of "tmpl_goi_001").
_DOC_TYPE_TO_DE_TEMPLATE: dict[str, str] = {
    ml_doc_type: de_tmpl_id
    for _, (de_tmpl_id, ml_doc_type) in _UPLOAD_LETTER_TYPE_MAP.items()
}


@app.post("/documents/upload", response_model=BlueprintDocResponse)
async def upload_document(
    file: UploadFile = File(...),
    user_id: str = Form(default="local-user"),
    letter_type: str = Form(default=""),
    db: Session = Depends(get_db),
):
    # 1. Validate file type
    ext = Path(file.filename or "").suffix.lower()
    _SUPPORTED = {".docx", ".pdf", ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"}
    if ext not in _SUPPORTED:
        raise HTTPException(400, f"Unsupported file type '{ext}'. Supported: docx, pdf, jpg, png, tiff")

    # 2. Save uploaded file to disk
    doc_id = crud.new_doc_id()
    save_ext = ext.lstrip(".")
    upload_path = crud.storage_path(doc_id, 0, save_ext)
    Path(upload_path).parent.mkdir(parents=True, exist_ok=True)
    content = await file.read()
    Path(upload_path).write_bytes(content)

    # 3. Resolve letter type → doc-engine template + ML doc_type
    lt = (letter_type or "").strip().lower()
    de_template_id, ml_doc_type = _UPLOAD_LETTER_TYPE_MAP.get(lt, ("tmpl_flex_001", "uploaded_doc"))

    # 4. Extract sections
    from app.services.doc_importer import (
        extract_sections_from_docx, extract_sections_from_pdf, extract_sections_from_image,
        generate_plain_docx,
    )
    from app.services.lexical_wrapper import text_to_lexical_node
    try:
        if ext == ".docx":
            detected_sections = extract_sections_from_docx(upload_path)
        elif ext == ".pdf":
            detected_sections = extract_sections_from_pdf(upload_path)
        else:
            detected_sections = extract_sections_from_image(upload_path)
    except Exception as exc:
        logging.warning("upload: section extraction failed: %s", exc)
        detected_sections = []

    # 5. Create doc-engine document
    client: DocEngineClient | None = app.state.docengine
    if client is None:
        raise HTTPException(503, "Doc-Engine not available")
    try:
        de_resp = await client.create_document(de_template_id, {})
    except DocEngineError as exc:
        raise HTTPException(exc.status_code, detail=str(exc.body))

    de_doc_id = str(de_resp["document_id"])
    de_ver = de_resp["version"]
    doc_data = de_resp.get("data") or {}

    # 6. PATCH each detected section into doc-engine.
    # Build a per-type pool so each available slot is consumed once (FIFO).
    # If the pool for a type is empty, INSERT_SECTION is called to add a new one.
    from collections import defaultdict as _defaultdict
    _sec_pool: dict[str, list[dict]] = _defaultdict(list)
    for _s in (doc_data.get("sections") or []):
        _sec_pool[_s.get("type", "")].append(_s)

    # Sort detected sections by blueprint order before inserting.
    # If detection returns paragraph before receiver_block (e.g. unusual layout),
    # inserting paragraph first and then receiver_block at "end" would violate
    # the blueprint order constraint → doc-engine rejects with incorrect_order.
    _FLEX_ORDER = {
        "precedence": -2, "security_classification": -1, "letterhead": 0,
        "reference_number": 1, "date": 2, "receiver_block": 3,
        "subject": 4, "salutation": 5, "paragraph": 5, "table_block": 5, "remarks_block": 7,
        "annexure_block": 8, "enclosure": 9, "endorsement": 10,
        "signee_block": 11, "distribution_list": 12, "copy_to": 13, "noo": 15,
    }
    detected_sections = sorted(
        enumerate(detected_sections),
        key=lambda t: (_FLEX_ORDER.get(t[1].get("type", ""), 99), t[0]),
    )
    detected_sections = [s for _, s in detected_sections]

    # Track the last inserted paragraph section ID so new paragraphs can be
    # inserted AFTER it (not at "end", which would place them after signee_block
    # and violate the order constraint).
    _last_para_id: str | None = None

    for sec in detected_sections:
        _sec_text = sec["text"]
        _sec_type = sec.get("type", "paragraph")
        # Apply section-type-specific Lexical formatting so the web editor
        # matches the original letter layout as closely as possible.
        # For DOCX uploads, sec["bold"]=True if all runs were bold in the source paragraph.
        # sec["align"] may be set by detection (e.g. "center" for "To whomsoever it may Concern").
        _docx_bold = sec.get("bold", False)
        _docx_underline = sec.get("underline", False)
        _detected_align = sec.get("align", "")
        if _sec_type == "subject":
            lx = text_to_lexical_node(_sec_text, bold=True, underline=True, align="center")
        elif _sec_type in ("remarks_block", "endorsement"):
            lx = text_to_lexical_node(_sec_text, bold=True, align="center")
        elif _sec_type == "date":
            lx = text_to_lexical_node(_sec_text, align="right")
        elif _sec_type == "signee_block":
            # Use _format_signee_lines() to split both multi-line AND single-line signee
            # blocks (e.g. "(RS Bhatia) Maj Addl Offr TAIC for DG TA" → 3-4 separate lines).
            # Join with \n\n so text_to_lexical_node makes one Lexical paragraph per line.
            from app.services.doc_importer import _format_signee_lines
            _signee_normalized = "\n\n".join(_format_signee_lines(_sec_text))
            lx = text_to_lexical_node(_signee_normalized, bold=_docx_bold)
        else:
            lx = text_to_lexical_node(_sec_text, bold=_docx_bold, underline=_docx_underline, align=_detected_align)
        # Sub-paragraph sections (4.1, 4.2 …) get indent=1 in the Lexical root so the
        # editor displays them indented relative to the parent paragraph.
        import re as _re
        if _sec_type == "paragraph" and _re.match(r"^\s*\d+\.\d+[\.\)]?\s+\S", _sec_text):
            for _para_node in (lx.get("root") or {}).get("children") or []:
                if isinstance(_para_node, dict):
                    _para_node["indent"] = 1
        _pool = _sec_pool.get(sec["type"], [])
        target_sec = _pool.pop(0) if _pool else None
        try:
            if target_sec:
                patch = await client.patch_section(
                    de_doc_id, target_sec["id"], de_ver,
                    {"richtext": {"format": "lexical", "state": lx}},
                )
                de_ver = patch["version"]
                target_sec["content"] = {"richtext": {"format": "lexical", "state": lx}}
                if sec["type"] in ("paragraph", "table_block"):
                    _last_para_id = target_sec["id"]
            else:
                # For paragraphs and table_blocks insert AFTER the previous body section
                # so the order constraint (paragraph/table_block < signee_block) is respected.
                if sec["type"] in ("paragraph", "table_block") and _last_para_id:
                    _position = {"policy": "after", "section_id": _last_para_id}
                else:
                    _position = {"policy": "end"}
                ins = await client.apply_command(de_doc_id, de_ver, {
                    "action": "INSERT_SECTION",
                    "target_type": sec["type"],
                    "target_ref": None,
                    "position": _position,
                    "ai_instruction": None,
                })
                if ins.get("status") == "applied":
                    de_ver = ins.get("version", de_ver)
                    new_sec_id = next(
                        (u.get("section", {}).get("id") for u in (ins.get("updates") or [])
                         if u.get("op") == "insert"),
                        None,
                    )
                    if new_sec_id:
                        patch = await client.patch_section(
                            de_doc_id, new_sec_id, de_ver,
                            {"richtext": {"format": "lexical", "state": lx}},
                        )
                        de_ver = patch["version"]
                        if sec["type"] in ("paragraph", "table_block"):
                            _last_para_id = new_sec_id
        except DocEngineError as _pe:
            logging.warning("upload: patch section type=%s err=%s", sec["type"], _pe)

    # 6b. Fetch final doc state from doc-engine after all patches
    try:
        final_de = await client.get_document(de_doc_id)
        doc_data = final_de.get("data") or doc_data
        de_ver = final_de.get("version", de_ver)
    except DocEngineError:
        pass  # use locally-tracked state if fetch fails

    # 7. Persist ML pipeline document
    doc = Document(
        id=doc_id, user_id=user_id, doc_type=ml_doc_type,
        docengine_doc_id=de_doc_id,
    )
    doc.title = Path(file.filename or "").stem.replace("_", " ").replace("-", " ").title()
    db.add(doc)
    db.flush()

    initial_state = {"structured": doc_data, "docengine_version": de_ver}
    ml_ver = crud.add_version(
        db, doc_id, initial_state,
        {"action": "upload", "filename": file.filename, "letter_type": lt or "auto"},
        upload_path,  # placeholder; overwritten below with generated DOCX
    )

    # 8. Generate DOCX for download
    render_warning: str | None = None
    out_docx = crud.storage_path(doc_id, ml_ver.id, "docx")
    try:
        template_obj = (
            db.query(Template)
            .filter(Template.doc_type == ml_doc_type.upper())
            .first()
        )
        if template_obj and template_obj.docx_path and os.path.exists(template_obj.docx_path):
            # Known letter type with a registered template — use template renderer
            render_state = doc_state_from_filled_skeleton(doc_data, ml_doc_type.upper())
            render_docx(template_obj.docx_path, render_state, out_docx)
        else:
            # No template (flexible / service letter) — build a plain DOCX from extracted sections
            from app.services.doc_importer import _SECTION_RENDER_ORDER
            sorted_sections = sorted(
                detected_sections,
                key=lambda s: _SECTION_RENDER_ORDER.get(s.get("type", ""), 7)
            )
            generate_plain_docx(sorted_sections, doc.title or "Uploaded Letter", out_docx)
        ml_ver.docx_path = out_docx
        db.add(ml_ver)
    except Exception as _render_exc:
        render_warning = f"DOCX render failed: {_render_exc}"
        logging.warning("upload: DOCX generation failed for %s: %s", doc_id, _render_exc)
        # Last resort: if the original upload was a DOCX, serve that
        if ext == ".docx":
            ml_ver.docx_path = upload_path
            db.add(ml_ver)

    db.commit()

    return BlueprintDocResponse(
        document_id=doc_id,
        docengine_doc_id=de_doc_id,
        version=ml_ver.id,
        docengine_version=de_ver,
        doc_type=ml_doc_type,
        blueprint_id=de_resp.get("blueprint_id", "bp_flexible_v1"),
        data=doc_data,
        render_warning=render_warning,
    )

@app.get("/documents/{doc_id}/export")
def export_document(doc_id: str, format: str = "docx", db: Session = Depends(get_db)):
    """Export the latest version of a document as DOCX or PDF (spec alias)."""
    doc = db.get(Document, doc_id)
    if not doc or not doc.current_version_id:
        raise HTTPException(404, "Document not found")
    return download(doc_id, doc.current_version_id, format, db)


@app.get("/documents/{doc_id}/preview")
def preview_document(doc_id: str, version_id: int | None = None, db: Session = Depends(get_db)):
    """Return the document as an inline PDF suitable for embedding in an iframe preview panel."""
    from fastapi.responses import Response
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    # Resolve which version to preview
    vid = version_id or doc.current_version_id
    if not vid:
        raise HTTPException(404, "No version available for preview")
    v = db.get(DocumentVersion, vid)
    if not v or v.document_id != doc_id:
        raise HTTPException(404, "Version not found")
    if not v.docx_path or not os.path.exists(v.docx_path):
        raise HTTPException(404, "Document file missing — please re-export first")

    # Convert to PDF if not already cached
    pdf_path = v.pdf_path or _build_pdf_path(v.docx_path)
    if not os.path.exists(pdf_path):
        docx_to_pdf(v.docx_path, pdf_path)
        v.pdf_path = pdf_path
        db.commit()

    with open(pdf_path, "rb") as fh:
        pdf_bytes = fh.read()

    # Content-Disposition: inline so the browser renders it instead of downloading
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename=\"{doc_id}_preview.pdf\""},
    )


@app.get("/documents/{doc_id}/versions/{version_id}/download")
def download(doc_id: str, version_id: int, format: str, db: Session = Depends(get_db)):
    v = db.get(DocumentVersion, version_id)
    if not v or v.document_id != doc_id:
        raise HTTPException(404, "Version not found")

    if format == "docx":
        if not v.docx_path or not os.path.exists(v.docx_path):
            raise HTTPException(404, "DOCX missing")
        dl_filename = f"{doc_id}_v{version_id}.docx"
        return FileResponse(v.docx_path, filename=dl_filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if format == "pdf":
        if not v.docx_path or not os.path.exists(v.docx_path):
            raise HTTPException(404, "DOCX missing")
        pdf_path = v.pdf_path or _build_pdf_path(v.docx_path)
        if not os.path.exists(pdf_path):
            docx_to_pdf(v.docx_path, pdf_path)
        v.pdf_path = pdf_path
        db.commit()
        dl_filename = f"{doc_id}_v{version_id}.pdf"
        return FileResponse(pdf_path, filename=dl_filename, media_type="application/pdf")

    raise HTTPException(400, "format must be docx or pdf")


@app.post("/documents/{doc_id}/revert", response_model=RevertDocumentResponse, tags=["documents"])
def revert_document(doc_id: str, req: RevertDocumentRequest, db: Session = Depends(get_db)):
    """
    Revert a document to a previous version (UI undo button).

    Creates a new DocumentVersion copying the target version's doc_state so the
    full audit trail is preserved. If target_version_id is omitted, reverts to the
    immediately previous version.
    """
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")
    if doc.current_version_id != req.version:
        raise HTTPException(409, "version_conflict")
    cur_v = db.get(DocumentVersion, doc.current_version_id)
    if cur_v is None:
        raise HTTPException(404, "Current version not found")

    if req.target_version_id is not None:
        target_v = db.get(DocumentVersion, req.target_version_id)
        if not target_v or target_v.document_id != doc_id:
            raise HTTPException(404, "Target version not found")
    else:
        target_v = (
            db.query(DocumentVersion)
            .filter(DocumentVersion.document_id == doc_id, DocumentVersion.id < cur_v.id)
            .order_by(DocumentVersion.id.desc())
            .first()
        )
        if target_v is None:
            raise HTTPException(422, "No previous version to revert to")

    new_v = crud.add_version(
        db,
        doc_id,
        target_v.doc_state,
        {
            "action": "UNDO",
            "reverted_from_version_id": cur_v.id,
            "reverted_to_version_id": target_v.id,
        },
        target_v.docx_path,
    )
    return RevertDocumentResponse(
        document_id=doc_id,
        version=new_v.id,
        reverted_from_version_id=cur_v.id,
        reverted_to_version_id=target_v.id,
    )


# ──────────────────────────────────────────────────────────────────
# Phase 7 — Letter Template Store
# ──────────────────────────────────────────────────────────────────

@app.post("/documents/{doc_id}/save-as-template", response_model=SaveAsTemplateResponse)
async def save_as_template(doc_id: str, req: SaveAsTemplateRequest, db: Session = Depends(get_db)):
    """Save the current section structure of a document as a reusable template.

    Static fields (letterhead, signee_block) are preserved; all variable fields
    (subject, paragraphs, ref_number, date, receiver) are blanked out so the user
    fills them fresh on each new document created from this template.
    """
    from app.services.template_store import save_template, build_section_schema
    from app.services.render_adapter import _section_text

    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    client: DocEngineClient | None = app.state.docengine
    if client is None:
        raise HTTPException(503, "Doc-Engine not available")

    if not doc.docengine_doc_id:
        raise HTTPException(422, "Document has no doc-engine record (cannot extract sections)")

    try:
        de_doc = await client.get_document(doc.docengine_doc_id)
    except DocEngineError as exc:
        raise HTTPException(exc.status_code, detail=str(exc.body))

    de_sections = (de_doc.get("data") or {}).get("sections") or []
    section_schema = build_section_schema(de_sections, _section_text)

    letter_type = (req.letter_type or "").strip() or (doc.doc_type or "uploaded_doc")
    display_name = (req.display_name or "").strip() or (
        doc.title or letter_type.replace("_", " ").title()
    )
    template_id = save_template(letter_type, display_name, doc_id, section_schema)

    return SaveAsTemplateResponse(
        template_id=template_id,
        letter_type=letter_type,
        display_name=display_name,
        section_count=len(section_schema),
    )


@app.get("/saved-templates", response_model=list[TemplateStoreItem])
def get_saved_templates(letter_type: str = ""):
    """List all saved letter templates, optionally filtered by letter_type."""
    from app.services.template_store import list_templates
    items = list_templates(letter_type.strip() or None)
    return [TemplateStoreItem(**item) for item in items]


@app.delete("/saved-templates/{template_id}")
def remove_saved_template(template_id: str):
    """Delete a saved letter template."""
    from app.services.template_store import delete_template
    if not delete_template(template_id):
        raise HTTPException(404, f"Template '{template_id}' not found")
    return {"ok": True}


@app.post("/documents/from-template/{template_id}", response_model=BlueprintDocResponse)
async def create_from_template(
    template_id: str,
    user_id: str = "local-user",
    prompt: str = "",
    db: Session = Depends(get_db),
):
    """Create a new document from a saved template, optionally AI-filling variable sections.

    Letterhead and signee_block always carry over from the template.
    If *prompt* is provided the subject, date and body paragraphs are filled by AI
    (same slot-extraction + LLM drafting used by /documents/generate).
    Without a prompt all variable sections start empty for manual editing.
    """
    from app.services.template_store import load_template
    from app.services.lexical_wrapper import text_to_lexical_node
    from app.services.doc_importer import generate_plain_docx

    tmpl = load_template(template_id)
    if tmpl is None:
        raise HTTPException(404, f"Template '{template_id}' not found")

    letter_type = tmpl.get("letter_type", "")
    section_schema: list[dict] = tmpl.get("section_schema") or []

    # Resolve the doc-engine template to use (same map as upload)
    de_template_id, ml_doc_type = _UPLOAD_LETTER_TYPE_MAP.get(
        letter_type, ("tmpl_flex_001", "uploaded_doc")
    )

    client: DocEngineClient | None = app.state.docengine
    if client is None:
        raise HTTPException(503, "Doc-Engine not available")

    try:
        de_resp = await client.create_document(de_template_id, {})
    except DocEngineError as exc:
        raise HTTPException(exc.status_code, detail=str(exc.body))

    de_doc_id = str(de_resp["document_id"])
    de_ver = de_resp["version"]
    doc_data = de_resp.get("data") or {}

    # PATCH / INSERT each section from the template schema.
    # Track already-patched IDs so repeatable types (e.g. multiple paragraphs)
    # INSERT a new section instead of overwriting the same one.
    _used_sec_ids: set[str] = set()
    for entry in section_schema:
        sec_type = entry.get("type", "")
        content_text = entry.get("content", "") or ""
        lx = text_to_lexical_node(content_text)

        target_sec = next(
            (s for s in (doc_data.get("sections") or [])
             if s.get("type") == sec_type and s.get("id") not in _used_sec_ids),
            None,
        )
        try:
            if target_sec:
                _used_sec_ids.add(target_sec["id"])
                patch = await client.patch_section(
                    de_doc_id, target_sec["id"], de_ver,
                    {"richtext": {"format": "lexical", "state": lx}},
                )
                de_ver = patch["version"]
                target_sec["content"] = {"richtext": {"format": "lexical", "state": lx}}
            else:
                # For repeatable sections (e.g. extra paragraphs), insert BEFORE signee_block
                # so the blueprint order rule (paragraph < signee_block) is not violated.
                _anchor_sec = next(
                    (s for s in (doc_data.get("sections") or [])
                     if s.get("type") in ("signee_block", "remarks_block", "annexure_block")
                     and s.get("type") != sec_type),
                    None,
                )
                _ins_position = (
                    {"policy": "before", "section_id": _anchor_sec["id"]}
                    if _anchor_sec else {"policy": "end"}
                )
                ins = await client.apply_command(de_doc_id, de_ver, {
                    "action": "INSERT_SECTION",
                    "target_type": sec_type,
                    "target_ref": None,
                    "position": _ins_position,
                    "ai_instruction": None,
                })
                if ins.get("status") == "applied":
                    de_ver = ins.get("version", de_ver)
                    new_sec_id = next(
                        (u.get("section", {}).get("id") for u in (ins.get("updates") or [])
                         if u.get("op") == "insert"),
                        None,
                    )
                    if new_sec_id:
                        _used_sec_ids.add(new_sec_id)
                        patch = await client.patch_section(
                            de_doc_id, new_sec_id, de_ver,
                            {"richtext": {"format": "lexical", "state": lx}},
                        )
                        de_ver = patch["version"]
        except DocEngineError as _pe:
            logging.warning("from-template: patch type=%s err=%s", sec_type, _pe)

    # Fetch final state
    try:
        final_de = await client.get_document(de_doc_id)
        doc_data = final_de.get("data") or doc_data
        de_ver = final_de.get("version", de_ver)
    except DocEngineError:
        pass

    # ── AI fill from prompt ───────────────────────────────────────────────────
    # Sticky sections (letterhead, signee_block) came from the saved template above
    # and must NOT be overwritten. Everything else (subject, date, paragraphs) is
    # filled from the prompt using the same slot-extraction + LLM drafting used by /generate.
    _STICKY = {"letterhead", "signee_block"}
    if prompt:
        import datetime as _dt
        _slots: dict = {}
        _ml_dt = ml_doc_type.upper()
        try:
            if _ml_dt == "DO_LETTER":
                from app.ml.slots.do_letter import _regex_fallback_do
                _slots = _regex_fallback_do(prompt)
            elif _ml_dt in ("GOI_LETTER", "SERVICE_LETTER"):
                from app.ml.slots.goi_letter import _regex_fallback_goi
                _slots = _regex_fallback_goi(prompt)
            elif _ml_dt == "LEAVE_CERTIFICATE":
                from app.ml.slots.leave_certificate import _regex_fallback as _lc_regex
                _slots = _lc_regex(prompt)
            elif _ml_dt == "MOVEMENT_ORDER":
                from app.ml.slots.movement_order import _regex_fallback_mo
                _slots = _regex_fallback_mo(prompt)
            elif _ml_dt in ("GENERAL_LETTER", "UPLOADED_DOC"):
                from app.ml.slots.goi_letter import _regex_fallback_goi
                _slots = _regex_fallback_goi(prompt)   # extracts date + ref
        except Exception:
            pass

        if not _slots.get("date"):
            _slots["date"] = _dt.date.today().strftime("%d %b %Y")

        if not _slots.get("subject"):
            _subj = re.sub(
                r"^(?:write\s+(?:a\s+)?)?(?:service\s+letter|service|do|goi|demi[\s-]official"
                r"|government\s+of\s+india\s+letter|invitation\s+letter|movement\s+order"
                r"|leave\s+certificate)\s*(?:for|about|to|on|regarding)?\s*",
                "", prompt.strip(), flags=re.IGNORECASE,
            ).strip()
            _subj = re.sub(r"\s+(?:dated?|on)\s+\d.*$", "", _subj, flags=re.IGNORECASE).strip()
            if _subj:
                _slots["subject"] = _subj.upper()[:80]

        # Draft body paragraphs via LLM
        _body_texts: list[str] = []
        try:
            if _ml_dt == "DO_LETTER":
                from app.ml.slots.do_letter import draft_body_paras
                _body_texts = await draft_body_paras(db, prompt, min_paras=2, max_paras=3)
            elif _ml_dt in ("GOI_LETTER", "SERVICE_LETTER"):
                from app.ml.slots.goi_letter import draft_numbered_paras
                _date_ctx = _slots.get("date") or _dt.date.today().strftime("%d %b %Y")
                _body_texts = await draft_numbered_paras(db, f"{prompt} [Letter date: {_date_ctx}]")
            elif _ml_dt == "MOVEMENT_ORDER":
                from app.ml.slots.movement_order import draft_numbered_paras as _mo_draft
                _body_texts = await _mo_draft(db, prompt, min_paras=2, max_paras=4)
            elif _ml_dt in ("LEAVE_CERTIFICATE", "INVITATION_LETTER"):
                from app.ml.slots.do_letter import draft_body_paras as _draft
                _body_texts = await _draft(db, prompt, min_paras=2, max_paras=3)
            else:
                # uploaded_doc, general_letter, inter_dep_note, or any user-saved template —
                # fall back to generic drafter so body paragraphs are filled from the prompt.
                from app.ml.slots.general_letter import draft_body_paras_general
                _body_texts = await draft_body_paras_general(db, prompt)
        except Exception:
            import traceback as _tb2
            COMMAND_LOGGER.warning("body para draft (from-template) failed doc_type=%s: %s", _ml_dt, _tb2.format_exc())
            _body_texts = []

        # PATCH single-value variable sections from slots
        _single = {
            "date":             _slots.get("date") or "",
            "subject":          _slots.get("subject") or "",
            "reference_number": _slots.get("file_reference_number") or "",
        }
        for _sec in (doc_data.get("sections") or []):
            if not isinstance(_sec, dict) or _sec.get("type") in _STICKY:
                continue
            _val = _single.get(_sec.get("type", ""))
            if _val:
                _lx = text_to_lexical_node(_val)
                _lxc = {"richtext": {"format": "lexical", "state": _lx}}
                try:
                    _p = await client.patch_section(de_doc_id, _sec["id"], de_ver, _lxc)
                    de_ver = _p["version"]
                    _sec["content"] = _lxc
                except DocEngineError as _pe:
                    logging.warning("from-template prompt: slot type=%s err=%s", _sec.get("type"), _pe)

        # PATCH paragraph sections with drafted body text
        _para_secs = [
            s for s in (doc_data.get("sections") or [])
            if isinstance(s, dict) and s.get("type") == "paragraph"
        ]
        for _i, _sec in enumerate(_para_secs):
            if _i >= len(_body_texts):
                break
            _lx = text_to_lexical_node(_body_texts[_i])
            _lxc = {"richtext": {"format": "lexical", "state": _lx}}
            try:
                _p = await client.patch_section(de_doc_id, _sec["id"], de_ver, _lxc)
                de_ver = _p["version"]
                _sec["content"] = _lxc
            except DocEngineError as _pe:
                logging.warning("from-template prompt: para err=%s", _pe)

    # Persist ML document
    doc_id = crud.new_doc_id()
    display_name = tmpl.get("display_name", "")
    doc = Document(
        id=doc_id, user_id=user_id, doc_type=ml_doc_type,
        docengine_doc_id=de_doc_id,
    )
    doc.title = display_name or letter_type.replace("_", " ").title()
    db.add(doc)
    db.flush()

    initial_state = {"structured": doc_data, "docengine_version": de_ver}
    ml_ver = crud.add_version(
        db, doc_id, initial_state,
        {"action": "from_template", "template_id": template_id, "letter_type": letter_type},
        "",
    )

    # Generate blank-structure DOCX for download
    out_docx = crud.storage_path(doc_id, ml_ver.id, "docx")
    render_warning: str | None = None
    try:
        crud.ensure_dir(os.path.dirname(out_docx))
        # Use current doc_data (with AI-filled sections if prompt was given)
        from app.services.doc_importer import sections_for_render
        plain_sections = sections_for_render(doc_data)
        generate_plain_docx(plain_sections, doc.title or "Letter", out_docx)
        ml_ver.docx_path = out_docx
        db.add(ml_ver)
    except Exception as _exc:
        render_warning = f"DOCX generation failed: {_exc}"
        logging.warning("from-template: DOCX generation failed: %s", _exc)

    db.commit()

    return BlueprintDocResponse(
        document_id=doc_id,
        docengine_doc_id=de_doc_id,
        version=ml_ver.id,
        docengine_version=de_ver,
        doc_type=ml_doc_type,
        blueprint_id=de_resp.get("blueprint_id", "bp_flexible_v1"),
        data=doc_data,
        render_warning=render_warning,
    )


@app.post("/documents/{doc_id}/feedback", response_model=FeedbackResponse)
async def submit_feedback(doc_id: str, body: FeedbackRequest, db: Session = Depends(get_db)):
    """Rate a generated document version and optionally provide a correction.

    Feedback is appended to data/ft_collected/feedback.jsonl for future fine-tuning.
    """
    import datetime as _dt

    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "Document not found")

    record = {
        "feedback_id": str(uuid.uuid4()),
        "doc_id": doc_id,
        "version_id": body.version_id,
        "doc_type": doc.doc_type,
        "rating": body.rating,
        "field": body.field,
        "correction": body.correction,
        "timestamp": _dt.datetime.utcnow().isoformat(),
    }

    out_dir = Path("data/ft_collected")
    out_dir.mkdir(parents=True, exist_ok=True)
    with open(out_dir / "feedback.jsonl", "a", encoding="utf-8") as fh:
        fh.write(json.dumps(record, ensure_ascii=False) + "\n")

    return FeedbackResponse(feedback_id=record["feedback_id"], saved=True)


# ---------------------------------------------------------------------------
# Wake-word detection — WebSocket endpoint
# ---------------------------------------------------------------------------
# The browser streams raw PCM-16 audio chunks (16 kHz, mono, 100 ms each =
# 3200 bytes per chunk).  openWakeWord runs inference on every chunk and
# sends a JSON event back when it detects "start_clerk" or "over_clerk".
#
# Models are loaded once (lazy singleton) from data/wake_word_models/.
# Train them once with:  python scripts/train_wake_words.py
# ---------------------------------------------------------------------------

_oww_model = None          # lazy singleton — loaded on first WebSocket connection
_OWW_MODEL_DIR = Path("data/wake_word_models")
_OWW_THRESHOLD = 0.5       # detection confidence threshold (0–1)
_OWW_CHUNK_SAMPLES = 1600  # 100 ms at 16 000 Hz


def _get_oww():
    """Load openWakeWord model (once). Returns None if models not yet trained."""
    global _oww_model
    if _oww_model is not None:
        return _oww_model
    start_model = _OWW_MODEL_DIR / "start_clerk.onnx"
    over_model  = _OWW_MODEL_DIR / "over_clerk.onnx"
    if not start_model.exists() or not over_model.exists():
        logging.warning(
            "wake_word: model files not found in %s — run scripts/train_wake_words.py first",
            _OWW_MODEL_DIR,
        )
        return None
    try:
        from openwakeword.model import Model as OWWModel
        _oww_model = OWWModel(
            wakeword_models=[str(start_model), str(over_model)],
            inference_framework="onnx",
        )
        logging.info("wake_word: models loaded from %s", _OWW_MODEL_DIR)
    except Exception as exc:
        logging.error("wake_word: failed to load models: %s", exc)
        _oww_model = None
    return _oww_model


@app.websocket("/ws/wake-word")
async def wake_word_ws(websocket: WebSocket):
    """Stream PCM-16 audio from the browser; fire events on wake-word detection.

    Browser sends raw bytes:  Int16Array  (16 kHz, mono, 100 ms = 3200 bytes)
    Server sends JSON events: {"event": "activate"}   — "Start Clerk" detected
                              {"event": "deactivate"} — "Over Clerk"  detected
                              {"event": "model_unavailable"}  — models not trained yet
    """
    import numpy as np

    await websocket.accept()
    model = _get_oww()
    if model is None:
        await websocket.send_json({"event": "model_unavailable",
                                   "message": "Wake word models not trained. Run scripts/train_wake_words.py"})
        await websocket.close()
        return

    logging.info("wake_word: client connected %s", websocket.client)
    try:
        while True:
            chunk = await websocket.receive_bytes()
            if len(chunk) < 2:
                continue
            audio = np.frombuffer(chunk, dtype=np.int16).astype(np.float32) / 32768.0
            preds = model.predict(audio)
            start_score = float(preds.get("start_clerk", 0))
            over_score  = float(preds.get("over_clerk",  0))
            if start_score >= _OWW_THRESHOLD:
                logging.debug("wake_word: START CLERK detected (score=%.2f)", start_score)
                await websocket.send_json({"event": "activate",   "score": round(start_score, 3)})
            elif over_score >= _OWW_THRESHOLD:
                logging.debug("wake_word: OVER CLERK detected (score=%.2f)", over_score)
                await websocket.send_json({"event": "deactivate", "score": round(over_score, 3)})
    except WebSocketDisconnect:
        logging.info("wake_word: client disconnected")
    except Exception as exc:
        logging.warning("wake_word: error — %s", exc)
        try:
            await websocket.close()
        except Exception:
            pass


