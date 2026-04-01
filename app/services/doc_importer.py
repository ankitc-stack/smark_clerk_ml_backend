"""doc_importer.py — Extract sections from uploaded DOCX, PDF (digital or scanned), and images.

Public API:
    extract_sections_from_docx(path) -> list[dict]
    extract_sections_from_pdf(path)  -> list[dict]
    extract_sections_from_image(path) -> list[dict]
    generate_plain_docx(sections, title, out_path) -> None

Each function returns a list of:
    {"type": str, "text": str, "confidence": float}
    Entries with confidence < 0.6 also include "low_confidence": True.
"""
from __future__ import annotations

import io
import re
import logging

import fitz  # PyMuPDF — already in requirements
from PIL import Image

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Surya OCR lazy singletons (models ~1.5 GB, downloaded once on first use)
# ---------------------------------------------------------------------------

_surya_models: "dict | None" = None


def _get_surya_models() -> dict:
    global _surya_models
    if _surya_models is None:
        from surya.foundation import FoundationPredictor
        from surya.detection import DetectionPredictor
        from surya.recognition import RecognitionPredictor
        foundation = FoundationPredictor()
        _surya_models = {
            "det_predictor": DetectionPredictor(),
            "rec_predictor": RecognitionPredictor(foundation_predictor=foundation),
        }
    return _surya_models


# ---------------------------------------------------------------------------
# Layout-aware OCR — extracts bounding boxes for spatial analysis
# ---------------------------------------------------------------------------

def _ocr_with_layout(img: Image.Image) -> list[dict]:
    """Run Surya OCR and return line-level bounding boxes, sorted top-to-bottom.

    Returns list of:
        {"text": str, "x_min": float, "x_max": float,
         "y_min": float, "y_max": float, "x_center": float, "conf": float}
    """
    m = _get_surya_models()
    results = m["rec_predictor"]([img], None, m["det_predictor"])
    boxes = []
    for line in (results[0].text_lines if results else []):
        x1, y1, x2, y2 = line.bbox
        boxes.append({
            "text":     line.text,
            "x_min":    float(x1),
            "x_max":    float(x2),
            "y_min":    float(y1),
            "y_max":    float(y2),
            "x_center": float((x1 + x2) / 2),
            "conf":     float(line.confidence),
        })
    boxes.sort(key=lambda b: (b["y_min"], b["x_min"]))
    # Infer bold by bbox height: lines taller than 1.25× median are likely bold/heading text.
    if boxes:
        import statistics as _stats
        heights = [b["y_max"] - b["y_min"] for b in boxes if b["y_max"] - b["y_min"] > 2]
        if heights:
            _med_h = _stats.median(heights)
            for b in boxes:
                b["inferred_bold"] = (b["y_max"] - b["y_min"]) > _med_h * 1.25
    return boxes


def _group_lines(ocr_boxes: list[dict]) -> list[dict]:
    """Group OCR boxes sharing the same visual row into display-row descriptors.

    Two boxes are in the same visual row if their y ranges overlap significantly
    (ref_number on left + date on right often appear as separate OCR boxes at same height).

    Returns list of:
        {"text": str,             # full row text (left → right)
         "left_text": str,        # left-cluster text (for inline rows)
         "right_text": str,       # right-cluster text (for inline rows)
         "alignment": str,        # "left" | "center" | "right" | "inline"
         "conf": float,
         "x_min": float, "x_max": float, "y_mid": float, "page_width": float}
    """
    if not ocr_boxes:
        return []

    page_width = max(b["x_max"] for b in ocr_boxes) or 1.0

    # Greedily group overlapping-Y boxes into visual rows
    rows: list[list[dict]] = []
    current: list[dict] = [ocr_boxes[0]]
    for b in ocr_boxes[1:]:
        row_y_max = max(x["y_max"] for x in current)
        row_h = max(1.0, row_y_max - min(x["y_min"] for x in current))
        # Box belongs to current row if its top is within the row's Y band
        if b["y_min"] < row_y_max - row_h * 0.2:
            current.append(b)
        else:
            rows.append(sorted(current, key=lambda x: x["x_min"]))
            current = [b]
    rows.append(sorted(current, key=lambda x: x["x_min"]))

    result = []
    for row in rows:
        x_min = min(b["x_min"] for b in row)
        x_max = max(b["x_max"] for b in row)
        y_mid = sum((b["y_min"] + b["y_max"]) / 2 for b in row) / len(row)
        conf = min(b["conf"] for b in row)
        row_width = x_max - x_min
        row_cx = (x_min + x_max) / 2

        # Split into left / right clusters for inline detection
        left_boxes  = [b for b in row if b["x_center"] < page_width * 0.45]
        right_boxes = [b for b in row if b["x_center"] > page_width * 0.55]

        left_text  = " ".join(b["text"] for b in sorted(left_boxes,  key=lambda b: b["x_min"]))
        right_text = " ".join(b["text"] for b in sorted(right_boxes, key=lambda b: b["x_min"]))
        full_text  = " ".join(b["text"] for b in row)

        # Alignment: narrow centered text → "center"; two side clusters → "inline"
        is_narrow = row_width < page_width * 0.60
        has_left  = bool(left_boxes)
        has_right = bool(right_boxes)

        if has_left and has_right and len(row) >= 2:
            alignment = "inline"
        elif is_narrow and page_width * 0.28 < row_cx < page_width * 0.72:
            alignment = "center"
        elif row_cx > page_width * 0.65:
            alignment = "right"
        else:
            alignment = "left"

        # Two-column split: if left and right clusters have a large gap (>30% page width),
        # emit as two separate rows so signee blocks and multi-column content are not merged.
        # ref_number+date pairs will still be correctly detected by their individual patterns.
        if alignment == "inline" and left_boxes and right_boxes:
            left_x_max  = max(b["x_max"] for b in left_boxes)
            right_x_min = min(b["x_min"] for b in right_boxes)
            gap = right_x_min - left_x_max
            if gap > page_width * 0.28:
                lx_min = min(b["x_min"] for b in left_boxes)
                lx_max = max(b["x_max"] for b in left_boxes)
                lx_cx  = (lx_min + lx_max) / 2
                rx_min = min(b["x_min"] for b in right_boxes)
                rx_max = max(b["x_max"] for b in right_boxes)
                l_align = "center" if page_width * 0.28 < lx_cx < page_width * 0.72 else "left"
                r_align = "right" if (rx_min + rx_max) / 2 > page_width * 0.65 else "center"
                _col_bold = sum(1 for b in row if b.get("inferred_bold", False)) > len(row) / 2
                base = {"conf": conf, "y_mid": y_mid, "page_width": page_width,
                        "left_text": "", "right_text": "", "inferred_bold": _col_bold}
                result.append({**base, "text": left_text,  "alignment": l_align,
                                "x_min": lx_min, "x_max": lx_max})
                result.append({**base, "text": right_text, "alignment": r_align,
                                "x_min": rx_min, "x_max": rx_max})
                continue

        # Row is "bold" if the majority of its text boxes were inferred bold
        bold_count = sum(1 for b in row if b.get("inferred_bold", False))
        row_inferred_bold = bold_count > len(row) / 2

        result.append({
            "text": full_text,
            "left_text": left_text,
            "right_text": right_text,
            "alignment": alignment,
            "conf": conf,
            "x_min": x_min,
            "x_max": x_max,
            "y_mid": y_mid,
            "page_width": page_width,
            "inferred_bold": row_inferred_bold,
        })
    return result


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_sections_from_docx(path: str) -> list[dict]:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn as _qn
    doc = DocxDocument(path)
    # Iterate body XML children in document order so tables appear at correct position.
    body = doc.element.body
    items: list[dict] = []   # {"kind": "para"|"table", "text": str}
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = "".join(r.text or "" for r in child.iter(_qn("w:t"))).strip()
            if text:
                items.append({"kind": "para", "text": text})
        elif tag == "tbl":
            # Convert DOCX table → markdown pipe-table string
            md_rows: list[str] = []
            for row in child.iter(_qn("w:tr")):
                cells = []
                for cell in row.iter(_qn("w:tc")):
                    cell_text = "".join(r.text or "" for r in cell.iter(_qn("w:t"))).strip()
                    cells.append(cell_text)
                if cells:
                    md_rows.append("| " + " | ".join(cells) + " |")
            if md_rows:
                # Insert separator after header row
                n_cols = len(md_rows[0].strip("|").split("|"))
                sep = "|" + "|".join("-------" for _ in range(n_cols)) + "|"
                md_rows.insert(1, sep)
                items.append({"kind": "table", "text": "\n".join(md_rows)})

    paragraphs = [it["text"] for it in items if it["kind"] == "para"]
    detected = _detect_sections(paragraphs)

    # Apply paragraph-level bold hint from original DOCX formatting.
    # If all non-empty runs in a paragraph are bold, mark the detected section bold.
    _bold_paras: set[str] = set()
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and p.runs and all(bool(r.bold) for r in p.runs if r.text.strip()):
            _bold_paras.add(t)
    for sec in detected:
        if sec["text"] in _bold_paras:
            sec["bold"] = True

    # Re-insert table_block entries at the correct position based on paragraph ordering.
    # For each table in items, find the nearest preceding paragraph in items to anchor it.
    if any(it["kind"] == "table" for it in items):
        # Build a mapping: paragraph text → index in detected sections
        result: list[dict] = list(detected)
        insert_offset = 0
        para_cursor = 0  # tracks which paragraph we've passed in items
        for it in items:
            if it["kind"] == "para":
                para_cursor += 1
            elif it["kind"] == "table":
                # Insert after the last para_cursor-th paragraph section seen so far
                insert_pos = min(para_cursor + insert_offset, len(result))
                result.insert(insert_pos, {
                    "type": "table_block",
                    "text": it["text"],
                    "confidence": 0.95,
                })
                insert_offset += 1
        return result

    return detected


# ---------------------------------------------------------------------------
# PDF extraction — hybrid: digital text first, OCR for scanned pages
# ---------------------------------------------------------------------------

_OCR_CHAR_THRESHOLD = 50  # characters per page below which we assume it is scanned


def extract_sections_from_pdf(path: str) -> list[dict]:
    doc = fitz.open(path)
    all_sections: list[dict] = []
    for page in doc:
        text = page.get_text("text") or ""
        if len(text.strip()) < _OCR_CHAR_THRESHOLD:
            # Scanned page — render at 2× zoom (~144 dpi) and use layout-aware OCR
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            page_sections = _extract_layout_sections(img)
            all_sections.extend(page_sections)
        else:
            # Digital page — fast text extraction
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            all_sections.extend(_detect_sections(lines))
    return _merge_cross_page(all_sections)


def _merge_cross_page(sections: list[dict]) -> list[dict]:
    """Merge consecutive same-type single-instance sections across page boundaries."""
    _SINGLE = {"reference_number", "date", "receiver_block", "subject", "signee_block", "letterhead"}
    merged: list[dict] = []
    for sec in sections:
        if merged and merged[-1]["type"] == sec["type"] and sec["type"] in _SINGLE:
            merged[-1]["text"] += "\n" + sec["text"]
            merged[-1]["confidence"] = max(merged[-1]["confidence"], sec["confidence"])
        else:
            merged.append(dict(sec))
    return merged


# ---------------------------------------------------------------------------
# Image extraction (JPG, PNG, TIFF, BMP …)
# ---------------------------------------------------------------------------

def _detect_tables_img2table(img_path: str) -> list[dict]:
    """Use img2table to detect bordered tables in an image.

    Returns [{"y1": int, "y2": int, "markdown": str}] for each detected table.
    Falls back to [] on any import/runtime error (img2table is optional).
    """
    try:
        from img2table.document import Image as ImgDoc
        from img2table.ocr import TesseractOCR
        _i2t_ocr = TesseractOCR()
        tables_raw = ImgDoc(img_path).extract_tables(
            ocr=_i2t_ocr, implicit_rows=True, borderless_tables=False
        ) or []
        result = []
        for t in tables_raw:
            if t.df is None or t.df.empty:
                continue
            df = t.df.fillna("")
            rows_clean = [
                [str(v).replace("\n", " ").strip() for v in row]
                for _, row in df.iterrows()
            ]
            md: list[str] = []
            for i, row in enumerate(rows_clean):
                md.append("| " + " | ".join(row) + " |")
                if i == 0:
                    md.append("|" + "|".join("-------" for _ in row) + "|")
            result.append({"y1": t.bbox.y1, "y2": t.bbox.y2, "markdown": "\n".join(md)})
        return result
    except Exception as _e:
        logger.debug("img2table table detection skipped: %s", _e)
        return []


def extract_sections_from_image(path: str) -> list[dict]:
    img = Image.open(path)
    tables = _detect_tables_img2table(path)
    return _extract_layout_sections(img, tables=tables, img_path=path)


def _extract_layout_sections(
    img: Image.Image,
    tables: list[dict] | None = None,
    img_path: str = "",
) -> list[dict]:
    """Full layout-aware OCR + section detection for a single image.

    If `tables` is provided (from img2table), OCR boxes that fall inside a
    detected table region are suppressed, and table_block sections are
    injected at the correct document position.
    """
    boxes = _ocr_with_layout(img)
    tables = tables or []

    if tables:
        # Filter OCR boxes that lie inside a detected table bbox
        def _in_table_region(box: dict) -> bool:
            y_mid = (box["y_min"] + box["y_max"]) / 2
            return any(tbl["y1"] <= y_mid <= tbl["y2"] for tbl in tables)

        non_table_boxes = [b for b in boxes if not _in_table_region(b)]
        rows = _group_lines(non_table_boxes)
        text_sections = _detect_sections_layout(rows)
        text_sections = _split_ref_date(text_sections)
        text_sections = _split_embedded_sub_paras(text_sections)
        _split_signee_from_last_para(text_sections)

        # Each section carries _y_mid set during detection (first row's y_mid).
        # Build combined list: (y, kind, section_or_table)
        combined: list[tuple[float, str, dict]] = []
        for sec in text_sections:
            y = sec.pop("_y_mid", 0.0)
            combined.append((y, "sec", sec))
        for tbl in tables:
            y_mid = (tbl["y1"] + tbl["y2"]) / 2
            combined.append((y_mid, "table", {"type": "table_block", "text": tbl["markdown"], "confidence": 0.9}))
        combined.sort(key=lambda x: x[0])

        return [item for _, _, item in combined]

    # No tables detected — original path
    rows = _group_lines(boxes)
    sections = _detect_sections_layout(rows)
    sections = _split_ref_date(sections)
    sections = _split_embedded_sub_paras(sections)
    _split_signee_from_last_para(sections)
    for s in sections:
        s.pop("_y_mid", None)
    return sections


# ---------------------------------------------------------------------------
# Plain DOCX generation (for download when no template available)
# ---------------------------------------------------------------------------

_RE_SUB_PARA = re.compile(r"^\s*\d+\.\d+[\.\)]?\s+\S")   # e.g. "4.1 text", "4.1. text", "4.1) text"
_RE_PARA_NUM = re.compile(r"^(\s*\d+[\.\)]\s+)(.*)", re.DOTALL)   # split "1. text" into prefix + text
_RE_RANK_PREFIX = re.compile(
    r"^(Maj(?:\s+Gen)?|Lt\s+(?:Col|Gen)|Brig(?:adier)?|Gen(?:eral)?|Col(?:onel)?"
    r"|Capt(?:ain)?|Sub\s+Lt|Gde|Hav|Sep|Cfn|Sgt|Cpl|Lt\s+Col)\b", re.IGNORECASE
)


def _format_signee_lines(text: str) -> list[str]:
    """Split a single-line signee '(Name) Rank Appt for Unit' into 3-4 separate lines.

    If text already contains newlines, just return stripped lines as-is.
    """
    if "\n" in text:
        return [l.strip() for l in text.splitlines() if l.strip()]

    lines: list[str] = []
    rest = text.strip()

    # 1. Extract "(Name)" part
    paren_close = rest.find(")")
    if paren_close >= 0:
        lines.append(rest[: paren_close + 1].strip())
        rest = rest[paren_close + 1 :].strip()

    if not rest:
        return lines or [text]

    # 2. Extract leading rank token
    m = _RE_RANK_PREFIX.match(rest)
    if m:
        lines.append(m.group(0).strip())
        rest = rest[m.end() :].strip()

    if not rest:
        return lines or [text]

    # 3. Split on "for " to separate appointment from "for [Senior Officer] [Unit]"
    for_match = re.search(r"\bfor\s+", rest, re.IGNORECASE)
    if for_match and for_match.start() > 0:
        appt = rest[: for_match.start()].strip()
        for_part = rest[for_match.start() :].strip()
        if appt:
            lines.append(appt)
        if for_part:
            lines.append(for_part)
    else:
        lines.append(rest)

    return lines if lines else [text]


def _lexical_para_alignment(state: dict, para_index: int = 0):
    """Return WD_ALIGN_PARAGRAPH for the nth Lexical paragraph, or None if unset/left."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    paras = (state.get("root") or {}).get("children") or []
    if para_index >= len(paras):
        return None
    fmt = str(paras[para_index].get("format") or "").lower()
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(fmt)  # returns None for "left" or ""


def _parse_markdown_table(text: str) -> list[list[str]] | None:
    """Parse a markdown pipe-table into a 2-D list of strings (header row first).

    Returns None if the text doesn't look like a markdown table.
    Skips the separator row (e.g. |---|---|).
    """
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Separator row (|---|---|) → skip; must contain at least one dash
        if "-" in line and re.match(r"^\|[-| :]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows if rows else None


def _has_inline_format(state: dict) -> bool:
    """Return True if any Lexical text node has non-zero format flags or inline CSS styling."""
    for para in (state.get("root") or {}).get("children") or []:
        for node in para.get("children") or []:
            if node.get("type") == "text":
                if node.get("format", 0) or node.get("style", ""):
                    return True
    return False


def _apply_runs_to_para(p, lx_para_node: dict) -> None:
    """Add python-docx formatted runs to paragraph `p` from one Lexical paragraph node.

    Preserves bold, italic, underline (format bitmask) and font-family, font-size,
    color, background-color (highlight) from CSS style string on each text node.
    """
    from docx.shared import Pt
    FLAG_BOLD, FLAG_ITALIC, FLAG_UNDERLINE = 1, 2, 8
    for node in (lx_para_node.get("children") or []):
        if node.get("type") != "text":
            continue
        text = node.get("text", "")
        if not text:
            continue
        fmt = node.get("format", 0)
        run = p.add_run(text)
        if fmt & FLAG_BOLD:      run.bold = True
        if fmt & FLAG_ITALIC:    run.italic = True
        if fmt & FLAG_UNDERLINE: run.underline = True
        css = node.get("style", "")
        if css:
            m = re.search(r"font-size:\s*(\d+)pt", css)
            if m:
                run.font.size = Pt(int(m.group(1)))
            m = re.search(r"font-family:\s*([^;]+)", css)
            if m:
                run.font.name = m.group(1).strip().strip("'\"")
            m = re.search(r"color:\s*#?([0-9a-fA-F]{6})", css)
            if m:
                try:
                    from docx.shared import RGBColor
                    hx = m.group(1)
                    run.font.color.rgb = RGBColor(
                        int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
                    )
                except Exception:
                    pass
            if re.search(r"background-color:", css):
                try:
                    from docx.enum.text import WD_COLOR_INDEX
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    pass


def _add_hr(doc) -> None:
    """Add a thin horizontal rule paragraph (bottom border on a blank line)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(_qn("w:val"), "single")
    bottom.set(_qn("w:sz"), "6")
    bottom.set(_qn("w:space"), "1")
    bottom.set(_qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_plain_docx(sections: list[dict], title: str, out_path: str) -> None:
    """Generate a simple DOCX from extracted sections using python-docx."""
    from docx import Document as DocxDoc
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    doc.core_properties.title = title

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Group consecutive ref+date sections so they render inline (ref left, date right)
    sections = _pair_ref_date_for_render(sections)

    for sec in sections:
        sec_type = sec.get("type", "paragraph")
        _align_override = sec.get("alignment")   # per-section override (may be None)
        _para_count_before = len(doc.paragraphs)

        # Inline ref+date pair — rendered as a single tab-separated paragraph
        if sec_type == "_ref_date_pair":
            p = doc.add_paragraph()
            p.add_run(sec["ref_text"])
            # tab stop at right margin pushes date to right
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:pos"), "9072")  # ~6.3 inches at 1440 twips/inch
            tabs.append(tab)
            pPr.append(tabs)
            p.add_run("\t" + sec["date_text"])
            continue  # alignment override not applicable to inline pairs

        text = sec.get("text", "").strip()
        if not text:
            continue

        if sec_type == "letterhead":
            rs = sec.get("richtext_state")
            lh_lines = [l for l in text.splitlines() if l.strip()]
            for i, line in enumerate(lh_lines):
                p = doc.add_paragraph(line.strip())
                if rs:
                    align = _lexical_para_alignment(rs, i)
                    if align is not None:
                        p.alignment = align

        elif sec_type == "reference_number":
            doc.add_paragraph(text)

        elif sec_type == "date":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif sec_type == "receiver_block":
            rs = sec.get("richtext_state")
            rb_lines = text.splitlines()
            for i, line in enumerate(rb_lines):
                p = doc.add_paragraph(line.strip() if line.strip() else "")
                if rs:
                    align = _lexical_para_alignment(rs, i)
                    if align is not None:
                        p.alignment = align

        elif sec_type == "subject":
            clean = re.sub(r"^\s*(?:SUBJECT|Sub(?:ject)?)\s*[:–\-]\s*", "", text,
                           flags=re.IGNORECASE).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rs = sec.get("richtext_state")
            if rs and _has_inline_format(rs):
                # Subject is always bold+underline; honour additional formatting per-run
                for lx_para in (rs.get("root") or {}).get("children") or []:
                    _apply_runs_to_para(p, lx_para)
                for run in p.runs:
                    run.bold = True       # enforce military style: subject always bold
                    run.underline = True
            else:
                run = p.add_run(clean)
                run.bold = True
                run.underline = True

        elif sec_type == "signee_block":
            doc.add_paragraph("")  # blank line before signee
            for line in _format_signee_lines(text):
                doc.add_paragraph(line)

        elif sec_type == "security_classification":
            p = doc.add_paragraph(text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

        elif sec_type == "salutation":
            doc.add_paragraph("")  # blank line before salutation
            p = doc.add_paragraph()
            if sec.get("align") == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.strip())
            if sec.get("bold"):
                run.bold = True
            if sec.get("underline"):
                run.underline = True

        elif sec_type == "precedence":
            # JSSD Appendix E: precedence keyword is right-aligned in superscription
            p = doc.add_paragraph(text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.bold = True

        elif sec_type == "noo":
            # "NOT ON ORIGINAL" — centered, bold italic stamp on copy letters
            p = doc.add_paragraph(text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.italic = True

        elif sec_type == "remarks_block":
            # RECOMMENDED / NOT RECOMMENDED — horizontal rule, then centered bold underline
            _add_hr(doc)
            p = doc.add_paragraph(text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.underline = True

        elif sec_type == "endorsement":
            # SANCTION / NOT SANCTION — horizontal rule, then centered bold underline
            _add_hr(doc)
            p = doc.add_paragraph(text.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.underline = True

        elif sec_type == "table_block":
            doc.add_paragraph("")  # blank line before table
            # Reconstruct table text from Lexical state — linebreak nodes must become \n
            # (lexical_to_plain_text space-joins text nodes, losing row separators)
            rs = sec.get("richtext_state")
            if rs:
                _lines: list[str] = []
                for _lx_para in (rs.get("root") or {}).get("children") or []:
                    for _node in (_lx_para.get("children") or []):
                        if _node.get("type") == "text":
                            _lines.append(_node.get("text") or "")
                        elif _node.get("type") == "linebreak":
                            _lines.append("\n")
                    _lines.append("\n")  # paragraph boundary
                table_text = "".join(_lines)
            else:
                table_text = text
            table_rows = _parse_markdown_table(table_text)
            if table_rows:
                from docx.enum.table import WD_TABLE_ALIGNMENT
                from docx.oxml import OxmlElement as _OxmlElement
                from docx.oxml.ns import qn as _tqn
                _n_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=_n_cols)
                tbl.style = "Table Grid"
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                # Ensure explicit black borders (some Word versions ignore Table Grid)
                _tblPr = tbl._tbl.tblPr
                if _tblPr is None:
                    _tblPr = _OxmlElement("w:tblPr")
                    tbl._tbl.insert(0, _tblPr)
                _tblBorders = _OxmlElement("w:tblBorders")
                for _side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    _b = _OxmlElement(f"w:{_side}")
                    _b.set(_tqn("w:val"), "single")
                    _b.set(_tqn("w:sz"), "4")
                    _b.set(_tqn("w:space"), "0")
                    _b.set(_tqn("w:color"), "000000")
                    _tblBorders.append(_b)
                _tblPr.append(_tblBorders)
                for r_idx, row_cells in enumerate(table_rows):
                    for c_idx in range(_n_cols):
                        cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = cell_text
                        # Bold header row
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
            else:
                doc.add_paragraph(text.strip())

        elif sec_type == "copy_to":
            # GOI style (JSSD Appendix C): "Copy of the above forwarded to :-"
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Copy of the above forwarded to :-")
            label_run.bold = True
            for line in text.strip().splitlines():
                stripped = line.strip()
                # Skip the "Copy to" header line itself if already present in text
                if stripped and not re.match(r"^Copy\s+(?:of\s+the\s+above\s+forwarded\s+to|to)\b", stripped, re.IGNORECASE):
                    doc.add_paragraph("    " + stripped)

        else:
            # paragraph, copy_to, distribution_list, enclosure, annexure_block, remarks_block
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement as _OxmlElement
            _TAB_TWIPS     = "720"    # 36 pt = 0.5 inch standard military indent
            _SUB_TAB_TWIPS = "1440"   # 72 pt = 1.0 inch (sub-para, double)
            _INDENT_IN     = Pt(36)
            _SUB_INDENT_IN = Pt(72)

            rs = sec.get("richtext_state")
            if rs and _has_inline_format(rs):
                # Inline formatting present — use Lexical runs, preserve numbered-para indents
                for lx_para in (rs.get("root") or {}).get("children") or []:
                    nodes = [n for n in (lx_para.get("children") or []) if n.get("type") == "text" and n.get("text")]
                    if not nodes:
                        doc.add_paragraph("")
                        continue
                    first_text = nodes[0].get("text", "")
                    if _RE_SUB_PARA.match(first_text):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent       = _SUB_INDENT_IN
                        p.paragraph_format.first_line_indent = -_INDENT_IN
                    elif _RE_PARA_NUM.match(first_text):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent       = _INDENT_IN
                        p.paragraph_format.first_line_indent = -_INDENT_IN
                        pPr = p._p.get_or_add_pPr()
                        tabs_el = _OxmlElement("w:tabs")
                        t = _OxmlElement("w:tab")
                        t.set(_qn("w:val"), "left")
                        t.set(_qn("w:pos"), _TAB_TWIPS)
                        tabs_el.append(t)
                        pPr.append(tabs_el)
                        # Normalize multiple tabs (e.g. "1.\t\t\t\t") → single tab for DOCX.
                        # The Lexical JSON may have 4 tabs for editor visual spacing; DOCX only
                        # needs one tab to reach the single custom tab stop.
                        _lp_norm = dict(lx_para)
                        _children = list(lx_para.get("children") or [])
                        if _children and _children[0].get("type") == "text":
                            _n0 = dict(_children[0])
                            _n0["text"] = re.sub(r'^(\s*\d+[\.\)]\t)\t+', r'\1', _n0["text"])
                            _lp_norm["children"] = [_n0] + _children[1:]
                        _apply_runs_to_para(p, _lp_norm)
                        continue
                    else:
                        p = doc.add_paragraph()
                    _apply_runs_to_para(p, lx_para)
            else:
                for line in text.splitlines():
                    line = line.strip()
                    if not line:
                        doc.add_paragraph("")
                        continue

                    # Sub-paragraph items (4.1, 4.2 …) — hanging indent one level deeper
                    # "4.1" starts at 0.35" (= body of para 4), body wraps at 0.70"
                    if _RE_SUB_PARA.match(line):
                        m_sub = re.match(r'^(\s*\d+\.\d+[\.\)]?\s*)(.*)', line, re.DOTALL)
                        if m_sub:
                            sub_num  = m_sub.group(1).rstrip()
                            sub_body = m_sub.group(2).strip()
                            p = doc.add_paragraph()
                            p.paragraph_format.left_indent       = _SUB_INDENT_IN
                            p.paragraph_format.first_line_indent = -_INDENT_IN
                            pPr = p._p.get_or_add_pPr()
                            tabs_el = _OxmlElement("w:tabs")
                            t = _OxmlElement("w:tab")
                            t.set(_qn("w:val"), "left")
                            t.set(_qn("w:pos"), _SUB_TAB_TWIPS)
                            tabs_el.append(t)
                            pPr.append(tabs_el)
                            p.add_run(sub_num + "\t" + sub_body)
                        else:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.left_indent = _INDENT_IN
                        continue

                    # Numbered paragraph — hanging indent so number aligns left,
                    # text wraps at 0.35" (matches military tab stop convention)
                    m_num = _RE_PARA_NUM.match(line)
                    if m_num:
                        num_prefix = m_num.group(1).rstrip()   # e.g. "1."
                        body_text  = m_num.group(2)
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent       = _INDENT_IN
                        p.paragraph_format.first_line_indent = -_INDENT_IN
                        # Tab stop at 0.35" so the body text starts consistently
                        pPr = p._p.get_or_add_pPr()
                        tabs_el = _OxmlElement("w:tabs")
                        t = _OxmlElement("w:tab")
                        t.set(_qn("w:val"), "left")
                        t.set(_qn("w:pos"), _TAB_TWIPS)
                        tabs_el.append(t)
                        pPr.append(tabs_el)
                        p.add_run(num_prefix + "\t" + body_text)
                    else:
                        doc.add_paragraph(line)

        # Reposition block horizontally — preserves text direction (left-to-right).
        # "right"  → left_indent = 4.5" so block starts in the right half of the page
        # "left"   → reset indent to page margin (default position)
        # "center" → center paragraph alignment (symmetric; text direction unchanged)
        if _align_override:
            _pos = _align_override.lower()
            for _p in doc.paragraphs[_para_count_before:]:
                if _pos == "right":
                    _p.paragraph_format.left_indent = Inches(4.5)
                elif _pos == "left":
                    _p.paragraph_format.left_indent = Inches(0)
                elif _pos == "center":
                    _p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Footer: security classification (Appendix E — centered header + footer) ─
    # If a security_classification section is present, also stamp it in the footer
    # of every page so classified markings appear top AND bottom as required.
    sec_class_text = next(
        (s.get("text", "").strip() for s in sections if s.get("type") == "security_classification"),
        None,
    )
    if sec_class_text:
        from docx.oxml.ns import qn as _fqn
        doc_section = doc.sections[0]
        reg_footer = doc_section.footer
        reg_footer.is_linked_to_previous = False
        # Clear the default empty paragraph python-docx puts in the footer
        for ep in list(reg_footer.paragraphs):
            ep._element.getparent().remove(ep._element)
        sc_p = reg_footer.add_paragraph(sec_class_text)
        sc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sc_p.runs:
            run.bold = True

    doc.save(out_path)


# Canonical render order for military letter sections.
# Matches bp_service_letter_v1 / bp_flexible_v1 order_rules.
# Python's sort is stable: multiple "paragraph" sections keep their relative order.
_SECTION_RENDER_ORDER: dict[str, int] = {
    "precedence":              -2,
    "security_classification":  0,
    "letterhead":               1,
    "reference_number":         2,
    "date":                     3,
    "receiver_block":           4,
    "salutation":               5,
    "subject":                  6,
    "paragraph":                7,
    "table_block":              7,
    "remarks_block":            8,
    "annexure_block":          10,
    "enclosure":               11,
    "endorsement":             12,
    "signee_block":            13,
    "distribution_list":       14,
    "copy_to":                 15,
    "noo":                     16,
}


def sections_for_render(structured: dict) -> list[dict]:
    """Convert doc-engine structured data sections to the flat list[dict] format
    that generate_plain_docx() expects: [{"type": ..., "text": ..., "alignment": ...}].

    Sections are sorted by canonical military-letter order so the DOCX always renders
    with subject before paragraphs, signee_block after all paragraphs, etc., regardless
    of the insertion order returned by the doc-engine.
    """
    from app.services.render_adapter import _section_text
    raw = []
    for sec in (structured.get("sections") or []):
        if not isinstance(sec, dict):
            continue
        _rt = (sec.get("content") or {}).get("richtext") or {}
        _rs = _rt.get("state") if isinstance(_rt.get("state"), dict) else None
        raw.append({
            "type": sec.get("type", "paragraph"),
            "text": _section_text(sec),
            "alignment": sec.get("alignment"),
            "richtext_state": _rs,
        })
    # Stable sort preserves relative order of multiple sections of the same type (e.g. paragraphs).
    raw.sort(key=lambda s: _SECTION_RENDER_ORDER.get(s["type"], 7))
    return raw


def _pair_ref_date_for_render(sections: list[dict]) -> list[dict]:
    """If consecutive reference_number + date sections exist, merge into a _ref_date_pair
    so generate_plain_docx can render them inline (ref left, date right) like the original."""
    result: list[dict] = []
    i = 0
    while i < len(sections):
        sec = sections[i]
        if (
            sec.get("type") == "reference_number" and
            i + 1 < len(sections) and
            sections[i + 1].get("type") == "date"
        ):
            result.append({
                "type": "_ref_date_pair",
                "ref_text": sec["text"],
                "date_text": sections[i + 1]["text"],
            })
            i += 2
        else:
            result.append(sec)
            i += 1
    return result


# ---------------------------------------------------------------------------
# Regex anchors (shared by both detection paths)
# ---------------------------------------------------------------------------
_MONTHS = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
)
_RE_REF_1        = re.compile(r"^\s*No\.?\s+\d[\w/\-]+", re.IGNORECASE)
_RE_REF_2        = re.compile(r"^\s*[A-Z]{2,}/\d+")
_RE_REF_3        = re.compile(r"^\s*\d[\d\w]*/[\w/\-]+")   # digits-first e.g. 12345/o00/Trails/01
_RE_DATE_LONG    = re.compile(rf"\b\d{{1,2}}\s+{_MONTHS}\s*\d{{4}}\b", re.IGNORECASE)   # \s* allows "Jan2026"
_RE_DATE_SHORT   = re.compile(r"\b\d{1,2}/\d{1,2}/\d{2,4}\b")
_RE_DATE_ISO     = re.compile(r"\b\d{4}-\d{2}-\d{2}\b")
_RE_SUBJECT      = re.compile(r"^\s*(?:SUBJECT|Sub(?:ject)?)\s*[:–\-]", re.IGNORECASE)
_RE_SUBJECT_CAPS = re.compile(r"^[A-Z][A-Z\s\(\)/\-\.,0-9]*[A-Z]$")
_RE_NUMBERED_PARA= re.compile(r"^\s*\d+[\.\)]\s+\S")
_RE_COPY_TO      = re.compile(r"^\s*Copy\s+to\b", re.IGNORECASE)
_RE_DISTRIB      = re.compile(r"^\s*Distr(?:ibution)?\b", re.IGNORECASE)
_RE_ENCL         = re.compile(r"^\s*Encl(?:osure)?\b", re.IGNORECASE)
_RE_ANNEXURE     = re.compile(r"^\s*ANNEXURE\b", re.IGNORECASE)
_RE_LONE_NUMBER    = re.compile(r"^\s*\d+[\.\)]\s*$")
_RE_TO_LINE        = re.compile(r"^\s*To\s*[,:]?\s*$", re.IGNORECASE)
_RE_TO_WHOMSOEVER  = re.compile(r"^\s*To\s+whom(?:so)?ever\s+it\s+may\s+concern", re.IGNORECASE)
_RE_SALUTATION     = re.compile(r"^\s*(?:Dear\s+(?:Sir|Ma'?am)|Sir|Ma'?am|Respected\s+Sir)\s*[,.]?\s*$", re.IGNORECASE)
_RE_SECURITY_CLASS = re.compile(r"^\s*SECURITY\s+CLASSIFICATION\b", re.IGNORECASE)
_RE_PRECEDENCE     = re.compile(r"^\s*(?:IMMEDIATE|PRIORITY|ROUTINE|FLASH|OPERATIONAL\s+IMMEDIATE)\s*$", re.IGNORECASE)
_RE_NOO            = re.compile(r"^\s*NOT\s+ON\s+ORIGINAL\s*$", re.IGNORECASE)
_RE_PERSONAL_APP   = re.compile(r"^\s*PERSONAL\s+APPLICATION\b", re.IGNORECASE)
# "Station: New Delhi ..." / "Place: ..." in signee area after body paragraphs
_RE_OFFICER_SIGNEE = re.compile(r"^\s*(?:Station|Place)\s*[:\-]?\s*\S", re.IGNORECASE)
# "Dated: Dec 2024 OIC TAIC" — "Dated" (not "Date") in body = signee date line
_RE_DATED_BODY     = re.compile(r"^\s*Dated\s*[:\-]", re.IGNORECASE)

# Standalone document-type headings — fire even at position 0 (no len(sections) guard).
# Specific enough that they can't be confused with GOI/service letter letterhead
# (which is a unit name/address, not a document type name).
_RE_DOC_TITLE = re.compile(
    r"^\s*(?:LEAVE\s+APPLICATION|MOVEMENT\s+ORDER|POSTING\s+ORDER|LEAVE\s+CERTIFICATE"
    r"|TRANSFER\s+ORDER|PROMOTION\s+ORDER|COURSE\s+NOMINATION|JOINING\s+REPORT"
    r"|REGIMENTAL\s+ORDER|CASUALTY\s+FORM|MEDICAL\s+BOARD|CHARGE\s+SHEET"
    r"|COURT\s+OF\s+INQUIRY|BONAFIDE\s+CERTIFICATE|CERTIFICATE)\s*$",
    re.IGNORECASE
)

# Recommendation / forwarding stamp boxes (appear after body, before sanction)
_RE_REMARKS_BLOCK = re.compile(
    r"^\s*(?:RECOMMENDED|NOT\s+RECOMMENDED|RECOMMENDED\s*/\s*NOT\s+RECOMMENDED"
    r"|FORWARDED|RETURNED|CONCUR(?:RED)?|NOTED|SEEN"
    r"|FOR\s+(?:NECESSARY\s+)?(?:ACTION|INFORMATION|ORDERS))\s*(?:/\s*[\w\s]+)?\s*$",
    re.IGNORECASE
)

# Sanction / endorsement stamp boxes (appear after recommendation)
_RE_ENDORSEMENT = re.compile(
    r"^\s*(?:SANCTION(?:ED)?|NOT\s+SANCTION(?:ED)?"
    r"|SANCTION(?:ED)?\s*/\s*NOT\s+SANCTION(?:ED)?)\s*$",
    re.IGNORECASE
)


def _has_date(text: str) -> bool:
    return bool(
        _RE_DATE_LONG.search(text) or _RE_DATE_SHORT.search(text) or _RE_DATE_ISO.search(text)
    )


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _is_section_header(text: str) -> bool:
    """Return True if text starts a new structural section (not a para continuation)."""
    return bool(
        _RE_SUBJECT.match(text) or _RE_TO_LINE.match(text) or
        _RE_COPY_TO.match(text) or _RE_DISTRIB.match(text) or
        _RE_ENCL.match(text) or _RE_ANNEXURE.match(text) or
        _RE_NUMBERED_PARA.match(text) or _RE_SUB_PARA.match(text) or
        _RE_LONE_NUMBER.match(text) or   # "5." alone is a new paragraph, not continuation
        re.match(r"^\s*Dated?\s*[:\-]", text, re.IGNORECASE)  # "Date:" / "Dated:" block
    )


def _is_para_continuation(prev_text: str, curr_text: str) -> bool:
    """Return True if curr_text is a wrapped continuation of prev_text's paragraph."""
    if not prev_text or not curr_text:
        return False
    if _is_section_header(curr_text):
        return False
    ends_mid     = not prev_text.rstrip().endswith(('.', '!', '?', ':'))
    starts_lower = curr_text[0].islower()
    return ends_mid or starts_lower


def _split_ref_date(sections: list[dict]) -> list[dict]:
    """If a reference_number section also contains a date, split it into
    reference_number + date so generate_plain_docx can render them inline.

    Example: "12345/o00/Trails/01 Jan2026" → ref="12345/o00/Trails/01" + date="Jan 2026"
    """
    result: list[dict] = []
    for sec in sections:
        if sec.get("type") != "reference_number":
            result.append(sec)
            continue
        text = sec["text"]
        m = _RE_DATE_LONG.search(text) or _RE_DATE_SHORT.search(text)
        if not m:
            result.append(sec)
            continue
        ref_text  = text[:m.start()].strip().rstrip("/").strip()
        date_text = text[m.start():].strip()
        _y = sec.get("_y_mid", 0.0)
        if ref_text:
            result.append({"type": "reference_number", "text": ref_text, "confidence": sec["confidence"], "_y_mid": _y})
        if date_text:
            result.append({"type": "date", "text": date_text, "confidence": 0.9, "_y_mid": _y})
    return result


_RE_EMBEDDED_SUB_PARA = re.compile(r'\s+(?=\d+\.\d+[\.\)]?\s+\S)')


def _split_embedded_sub_paras(sections: list[dict]) -> list[dict]:
    """Split sub-paragraph items (4.1, 4.2 …) that got merged into a parent paragraph.

    Handles two cases:
      - newline-separated: "4. text\\n4.1 sub"  (DOCX continuation)
      - space-separated:   "4. text 4.1 sub 4.2 sub"  (OCR continuation)
    """
    result: list[dict] = []
    for sec in sections:
        if sec.get("type") != "paragraph":
            result.append(sec)
            continue
        text = sec["text"]
        base_conf = sec.get("confidence", 0.85)

        # Build a flat list of chunks by splitting each newline-line on inline sub-para boundaries
        chunks: list[str] = []
        for raw_line in text.split("\n"):
            inline_parts = _RE_EMBEDDED_SUB_PARA.split(raw_line)
            chunks.extend(p.strip() for p in inline_parts if p.strip())

        if len(chunks) <= 1:
            result.append(sec)
            continue

        _y = sec.get("_y_mid", 0.0)
        for idx, chunk in enumerate(chunks):
            entry: dict = {
                "type": "paragraph",
                "text": chunk,
                "confidence": base_conf if idx == 0 else 0.85,
                "_y_mid": _y,
            }
            if entry["confidence"] < 0.6:
                entry["low_confidence"] = True
            result.append(entry)

    return result


_RE_SIGNEE_BRACKET = re.compile(r"\s+\(([A-Z][A-Za-z\s\.]+)\)")   # " (Name)"


def _split_signee_from_last_para(sections: list[dict]) -> None:
    """If the last numbered paragraph ends with an inline signee block
    (e.g. "5. ...requested pl (RS Bhatia) Maj Addl Offr TAIC for DG TA"),
    split it into the paragraph text + a separate signee_block.

    Modifies sections in-place.
    """
    # Only act if there is no signee_block already
    if any(s["type"] == "signee_block" for s in sections):
        return

    # Find the last paragraph section
    last_para_idx: int | None = None
    for i in range(len(sections) - 1, -1, -1):
        if sections[i]["type"] == "paragraph":
            last_para_idx = i
            break
    if last_para_idx is None:
        return

    para_text = sections[last_para_idx]["text"]
    m = _RE_SIGNEE_BRACKET.search(para_text)
    if not m:
        return

    # Heuristic: the signee part must be in the last 40% of the text length
    if m.start() < len(para_text) * 0.4:
        return

    main_text   = para_text[:m.start()].strip()
    signee_text = para_text[m.start():].strip()

    # Only split if main text is still substantial
    if len(main_text) < 5:
        return

    _y = sections[last_para_idx].get("_y_mid", 0.0)
    sections[last_para_idx]["text"] = main_text
    sections.insert(last_para_idx + 1, {
        "type": "signee_block",
        "text": signee_text,
        "confidence": 0.7,
        "_y_mid": _y,
    })


def _promote_signee(sections: list[dict]) -> None:
    """Promote trailing short unclassified paragraph sections to signee_block in-place."""
    # Collect indices of trailing low-confidence (unclassified) paragraphs
    tail: list[int] = []
    for i in range(len(sections) - 1, -1, -1):
        sec = sections[i]
        if sec["type"] == "paragraph" and sec.get("confidence", 1.0) < 0.65:
            tail.insert(0, i)
        elif sec["type"] in {"copy_to", "distribution_list", "enclosure", "annexure_block", "remarks_block", "endorsement"}:
            continue  # these may trail the signee; keep scanning
        else:
            break  # numbered para or known section — stop

    if not tail:
        return

    # Plausibility check: combined text must be ≤ 6 lines (signee is short)
    total_lines = sum(
        len([l for l in sections[i]["text"].splitlines() if l.strip()])
        for i in tail
    )
    if total_lines > 6:
        return

    merged_text = "\n".join(sections[i]["text"] for i in tail)
    sections[tail[0]]["type"] = "signee_block"
    sections[tail[0]]["text"] = merged_text
    sections[tail[0]]["confidence"] = 0.75
    sections[tail[0]].pop("low_confidence", None)
    for i in reversed(tail[1:]):
        sections.pop(i)


# ---------------------------------------------------------------------------
# Layout-aware section detection (image / scanned PDF path)
# ---------------------------------------------------------------------------

def _detect_sections_layout(row_infos: list[dict]) -> list[dict]:
    """Classify visual rows (from _group_lines) into document sections.

    Handles:
      * Inline ref_number + date on same visual row
      * Letterhead block (centered / left lines before ref/To/subject)
      * "To" line as pivot separating letterhead from receiver_block
      * ALL-CAPS or "Sub:" subject
      * Numbered paragraphs with continuation merging
      * Signee block detection
    """
    sections: list[dict] = []
    to_seen          = False
    subject_seen     = False
    ref_or_date_seen = False   # True once ref_number or date is found (no "To" needed)
    current_para_idx: int | None = None  # index of last paragraph entry
    _current_row_y: float = 0.0          # y_mid of the row being processed
    _row_bold: bool = False              # inferred from OCR bbox height for current row

    def _add(sec_type: str, text: str, confidence: float, *, align: str = "", force_bold: bool = False, force_underline: bool = False) -> None:
        nonlocal current_para_idx
        entry: dict = {"type": sec_type, "text": text, "confidence": confidence,
                       "_y_mid": _current_row_y}
        if confidence < 0.6:
            entry["low_confidence"] = True
        if _row_bold or force_bold:
            entry["bold"] = True
        if force_underline:
            entry["underline"] = True
        if align:
            entry["align"] = align
        sections.append(entry)
        current_para_idx = len(sections) - 1 if sec_type == "paragraph" else None

    def _merge_last(text: str, conf: float) -> None:
        sections[-1]["text"] += "\n" + text
        sections[-1]["confidence"] = max(sections[-1]["confidence"], conf)
        sections[-1].pop("low_confidence", None)

    # Pre-process: split rows where OCR merged multiple lines into one
    _RE_ORPHANED_PARA  = re.compile(r'^(\d+[.)].+?)\s+(\d+\.)\s*$', re.DOTALL)
    # "7. Reliver: ... Date: Dec 2024 (Signature of applicant)" → split at "Date:"
    _RE_ORPHANED_DATE  = re.compile(
        r'^(\d+[.)].+?)\s+((?:Date|Dated)\s*[:\-]\s*.+)$', re.IGNORECASE | re.DOTALL
    )
    expanded_rows: list[dict] = []
    for ri in row_infos:
        t = re.sub(r"<br\s*/?>", " ", ri["text"], flags=re.IGNORECASE).strip()
        m_num  = _RE_ORPHANED_PARA.match(t)
        m_date = _RE_ORPHANED_DATE.match(t) if not m_num else None
        if m_num and _RE_NUMBERED_PARA.match(t):
            expanded_rows.append({**ri, "text": m_num.group(1).strip()})
            expanded_rows.append({**ri, "text": m_num.group(2).strip(), "conf": 0.7})
        elif m_date and _RE_NUMBERED_PARA.match(t):
            expanded_rows.append({**ri, "text": m_date.group(1).strip()})
            expanded_rows.append({**ri, "text": m_date.group(2).strip(), "conf": 0.7})
        else:
            expanded_rows.append(ri)
    row_infos = expanded_rows

    for ri in row_infos:
        _current_row_y = float(ri.get("y_mid", 0.0))
        _row_bold      = bool(ri.get("inferred_bold", False))
        text  = re.sub(r"<br\s*/?>", " ", ri["text"], flags=re.IGNORECASE).strip()
        if not text:
            continue
        align = ri["alignment"]
        conf  = max(0.3, float(ri.get("conf", 0.7)))

        # ── PRECEDENCE (IMMEDIATE / PRIORITY / ROUTINE — right-aligned superscription) ─
        if _RE_PRECEDENCE.match(text):
            _add("precedence", text, 0.95)
            continue

        # ── Security classification line (top/bottom of letter) ─────────────
        if _RE_SECURITY_CLASS.match(text):
            _add("security_classification", text, 0.95)
            continue

        # ── "To whomsoever it may Concern" — certificate / NOC opener ────────
        # Maps to salutation (blueprint order 5, after subject=4) so title comes first
        if _RE_TO_WHOMSOEVER.match(text):
            _add("salutation", text, 0.95, align="center", force_bold=True, force_underline=True)
            continue

        # ── Inline row: separate left cluster (ref) and right cluster (date) ────
        if align == "inline":
            left  = ri["left_text"].strip()
            right = ri["right_text"].strip()
            # Numbered paragraphs (e.g. "3. Leave applied -") must never be treated as ref/date
            left_is_numbered = bool(_RE_NUMBERED_PARA.match(left) or _RE_SUB_PARA.match(left))
            left_is_ref   = (not left_is_numbered) and bool(_RE_REF_1.match(left)  or _RE_REF_2.match(left) or _RE_REF_3.match(left))
            right_is_date = _has_date(right) and len(right.strip()) < 30   # tighter threshold
            right_is_ref  = bool(_RE_REF_1.match(right) or _RE_REF_2.match(right) or _RE_REF_3.match(right))
            left_is_date  = (not left_is_numbered) and _has_date(left) and len(left.strip()) < 30
            if not left_is_numbered and (left_is_ref or right_is_date or right_is_ref or left_is_date):
                if left:
                    _add("reference_number" if left_is_ref  else "date", left,  0.9)
                if right:
                    _add("date"             if right_is_date else "reference_number", right, 0.9)
                ref_or_date_seen = True
                continue
            # Inline but not ref/date — fall through with full text

        # ── Reference number (standalone short line) ──────────────────────────
        if not subject_seen and len(sections) < 15 and (
            _RE_REF_1.match(text) or _RE_REF_2.match(text) or _RE_REF_3.match(text)
        ):
            _add("reference_number", text, 0.9)
            ref_or_date_seen = True
            continue

        # ── Date (short standalone line with date pattern) ────────────────────
        if not subject_seen and len(sections) < 15 and len(text) < 60 and _has_date(text):
            _add("date", text, 0.9)
            ref_or_date_seen = True
            continue

        # ── "To" line → pivot to receiver block ──────────────────────────────
        if _RE_TO_LINE.match(text):
            to_seen = True
            ref_or_date_seen = True   # "To" also anchors the receiver zone
            _add("receiver_block", text, 0.9)
            continue

        # ── Receiver block continuation (after To OR after ref/date found) ────
        if (to_seen or ref_or_date_seen) and not subject_seen:
            is_subj_prefixed = bool(_RE_SUBJECT.match(text))
            is_subj_caps = (
                5 < len(text) < 120 and
                _RE_SUBJECT_CAPS.match(text.strip()) and
                not _RE_NUMBERED_PARA.match(text) and
                len(sections) >= 1    # relaxed from >= 2; ref_or_date already found
            )
            if not is_subj_prefixed and not is_subj_caps:
                if sections and sections[-1]["type"] == "receiver_block":
                    _merge_last(text, conf)
                else:
                    _add("receiver_block", text, 0.75)
                continue

        # ── Salutation ("Sir,", "Dear Sir," — after receiver zone) ──────────
        if not subject_seen and _RE_SALUTATION.match(text):
            _add("salutation", text, 0.95)
            continue

        # ── Personal Application heading → treated as subject section ─────────
        if not subject_seen and _RE_PERSONAL_APP.match(text):
            _add("subject", text, 0.95)
            subject_seen = True
            continue

        # ── Known document-type heading (LEAVE APPLICATION, MOVEMENT ORDER …) ─
        # Fires even at position 0 — specific enough to not collide with letterhead
        if not subject_seen and _RE_DOC_TITLE.match(text):
            _add("subject", text, 0.95)
            subject_seen = True
            continue

        # ── Subject (explicit "Sub:" / "SUBJECT:" prefix) ────────────────────
        if _RE_SUBJECT.match(text):
            _add("subject", text, 0.95)
            subject_seen = True
            continue

        # ── Subject (ALL-CAPS line — relaxed to >= 1 section already found) ──
        if (
            not subject_seen and
            5 < len(text) < 120 and
            _RE_SUBJECT_CAPS.match(text.strip()) and
            not _RE_NUMBERED_PARA.match(text) and
            len(sections) >= 1
        ):
            _add("subject", text, 0.82)
            subject_seen = True
            continue

        # ── Letterhead (top block, before any ref/date/To/subject) ───────────
        if not to_seen and not subject_seen and not ref_or_date_seen:
            if sections and sections[-1]["type"] == "letterhead":
                _merge_last(text, conf)
            else:
                _add("letterhead", text, 0.8 if align == "center" else 0.72)
            continue

        # ── Structural body sections ──────────────────────────────────────────
        if _RE_COPY_TO.match(text):
            _add("copy_to", text, 0.9)
            continue
        if _RE_DISTRIB.match(text):
            _add("distribution_list", text, 0.9)
            continue
        if _RE_ENCL.match(text):
            _add("enclosure", text, 0.9)
            continue
        if _RE_ANNEXURE.match(text):
            _add("annexure_block", text, 0.9)
            continue

        # ── NOT ON ORIGINAL stamp ─────────────────────────────────────────────
        if _RE_NOO.match(text):
            _add("noo", text, 0.95)
            continue

        # ── Recommendation stamp (RECOMMENDED / NOT RECOMMENDED …) ───────────
        if _RE_REMARKS_BLOCK.match(text):
            _add("remarks_block", text, 0.90)
            continue

        # ── Endorsement / sanction stamp (SANCTION / NOT SANCTION …) ─────────
        if _RE_ENDORSEMENT.match(text):
            _add("endorsement", text, 0.90)
            continue

        # ── Table (markdown pipe-table row from OCR) ──────────────────────────
        if text.startswith("|") and text.count("|") >= 3:
            # Accumulate consecutive pipe-table lines into a single table_block
            if sections and sections[-1]["type"] == "table_block":
                sections[-1]["text"] += "\n" + text
            else:
                _add("table_block", text, 0.9)
            current_para_idx = None
            continue

        # ── Numbered paragraph ────────────────────────────────────────────────
        if _RE_NUMBERED_PARA.match(text):
            _add("paragraph", text, 0.9)
            continue

        # ── Sub-paragraph (4.1, 4.2 …) — own section, not a continuation ─────
        if subject_seen and _RE_SUB_PARA.match(text):
            _add("paragraph", text, 0.85)
            continue

        # ── Paragraph continuation (wrapped body text after a paragraph) ──────
        if subject_seen and current_para_idx is not None:
            if _is_para_continuation(sections[current_para_idx]["text"], text):
                sections[current_para_idx]["text"] += " " + text
                continue

        # ── Officer station / dated lines (signee area after body) ───────────
        # "Station: New Delhi (RS Bhatia) Maj" / "Dated: Dec 2024 OIC TAIC"
        if subject_seen and (_RE_OFFICER_SIGNEE.match(text) or _RE_DATED_BODY.match(text)):
            if sections and sections[-1]["type"] == "signee_block":
                _merge_last(text, conf)
            else:
                _add("signee_block", text, 0.85)
            continue

        # ── Unclassified body / pre-subject leftover ──────────────────────────
        if subject_seen:
            _add("paragraph", text, 0.5)
        else:
            # Pre-subject unclassified — letterhead continuation
            if sections and sections[-1]["type"] == "letterhead":
                _merge_last(text, conf)
            else:
                _add("letterhead", text, 0.6)

    # ── Post-process: signee detection ────────────────────────────────────────
    _promote_signee(sections)

    logger.debug(
        "doc_importer: detected %d sections (layout-aware) from %d rows",
        len(sections), len(row_infos),
    )
    return sections


# ---------------------------------------------------------------------------
# Text-only section detection (used for DOCX / digital PDF pages)
# ---------------------------------------------------------------------------

def _detect_sections(paragraphs: list[str]) -> list[dict]:
    """Classify paragraphs into section types and return structured list."""
    # Strip HTML tags (e.g. <br> from img2table markdown output)
    paragraphs = [re.sub(r"<[^>]+>", " ", p).strip() for p in paragraphs if p.strip()]
    # Pre-process: OCR often splits "1.  Text" → ["1.", "Text"]. Merge them back.
    merged_paras: list[str] = []
    skip_next = False
    for j, p in enumerate(paragraphs):
        if skip_next:
            skip_next = False
            continue
        if _RE_LONE_NUMBER.match(p) and j + 1 < len(paragraphs):
            merged_paras.append(p.rstrip() + " " + paragraphs[j + 1])
            skip_next = True
        else:
            merged_paras.append(p)
    paragraphs = merged_paras

    sections: list[dict] = []
    subject_idx: int | None = None
    last_para_idx: int | None = None   # index of the most recent "paragraph" section
    signee_candidates: list[int] = []

    for i, para in enumerate(paragraphs):
        if not para:
            continue

        sec_type: str | None = None
        confidence: float = 0.4

        # --- precedence ---
        if _RE_PRECEDENCE.match(para):
            sec_type, confidence = "precedence", 0.95
            last_para_idx = None

        # --- security_classification ---
        elif _RE_SECURITY_CLASS.match(para):
            sec_type, confidence = "security_classification", 0.95
            last_para_idx = None

        # --- "To whomsoever it may Concern" — certificate/NOC opener ---
        elif _RE_TO_WHOMSOEVER.match(para):
            sections.append({"type": "salutation", "text": para, "confidence": 0.95,
                              "bold": True, "underline": True, "align": "center"})
            last_para_idx = None
            continue

        # --- reference_number ---
        elif i < 12 and (_RE_REF_1.match(para) or _RE_REF_2.match(para) or _RE_REF_3.match(para)):
            sec_type, confidence = "reference_number", 0.9
            last_para_idx = None

        # --- date ---
        elif i < 12 and _has_date(para) and len(para) < 60:
            sec_type, confidence = "date", 0.9
            last_para_idx = None

        # --- subject ---
        elif _RE_SUBJECT.match(para):
            sec_type, confidence = "subject", 0.95
            subject_idx = len(sections)
            last_para_idx = None

        # --- subject (ALL-CAPS short line after at least 3 earlier paragraphs) ---
        elif (
            subject_idx is None and i >= 3 and
            5 < len(para) < 100 and
            _RE_SUBJECT_CAPS.match(para.strip()) and
            not _RE_NUMBERED_PARA.match(para)
        ):
            sec_type, confidence = "subject", 0.8
            subject_idx = len(sections)
            last_para_idx = None

        # --- salutation ---
        elif subject_idx is None and _RE_SALUTATION.match(para):
            sec_type, confidence = "salutation", 0.95
            last_para_idx = None

        # --- personal application heading → subject section ---
        elif subject_idx is None and _RE_PERSONAL_APP.match(para):
            sec_type, confidence = "subject", 0.95
            subject_idx = len(sections)
            last_para_idx = None

        # --- known document-type heading (LEAVE APPLICATION, MOVEMENT ORDER …) ---
        elif subject_idx is None and _RE_DOC_TITLE.match(para):
            sec_type, confidence = "subject", 0.95
            subject_idx = len(sections)
            last_para_idx = None

        # --- numbered paragraphs ---
        elif _RE_NUMBERED_PARA.match(para):
            sec_type, confidence = "paragraph", 0.9

        # --- sub-paragraph (4.1, 4.2 …) ---
        elif _RE_SUB_PARA.match(para):
            sec_type, confidence = "paragraph", 0.85

        # --- copy_to ---
        elif _RE_COPY_TO.match(para):
            sec_type, confidence = "copy_to", 0.9
            last_para_idx = None

        # --- distribution_list ---
        elif _RE_DISTRIB.match(para):
            sec_type, confidence = "distribution_list", 0.9
            last_para_idx = None

        # --- enclosure ---
        elif _RE_ENCL.match(para):
            sec_type, confidence = "enclosure", 0.9
            last_para_idx = None

        # --- annexure_block ---
        elif _RE_ANNEXURE.match(para):
            sec_type, confidence = "annexure_block", 0.9
            last_para_idx = None

        # --- noo (Not on Original stamp) ---
        elif _RE_NOO.match(para):
            sec_type, confidence = "noo", 0.95
            last_para_idx = None

        # --- recommendation stamp (RECOMMENDED / NOT RECOMMENDED …) ---
        elif _RE_REMARKS_BLOCK.match(para):
            sec_type, confidence = "remarks_block", 0.90
            last_para_idx = None

        # --- endorsement / sanction stamp (SANCTION / NOT SANCTION …) ---
        elif _RE_ENDORSEMENT.match(para):
            sec_type, confidence = "endorsement", 0.90
            last_para_idx = None

        # --- table_block: markdown pipe-table (OCR/PDF text path) ---
        elif para.startswith("|") and para.count("|") >= 3:
            # Merge consecutive pipe rows into same table_block section
            if sections and sections[-1]["type"] == "table_block":
                sections[-1]["text"] += "\n" + para
                continue
            sec_type, confidence = "table_block", 0.9
            last_para_idx = None

        # --- receiver_block (block before subject that isn't ref/date) ---
        elif subject_idx is None:
            sec_type, confidence = "receiver_block", 0.75
            last_para_idx = None

        # --- officer station / dated lines in signee area ---
        elif subject_idx is not None and (
            _RE_OFFICER_SIGNEE.match(para) or _RE_DATED_BODY.match(para)
        ):
            if sections and sections[-1]["type"] == "signee_block":
                sections[-1]["text"] += "\n" + para
                continue
            sec_type, confidence = "signee_block", 0.85
            last_para_idx = None

        # --- body paragraph (unmatched text after subject) ---
        else:
            # Check if this is a continuation of the previous paragraph
            if last_para_idx is not None and _is_para_continuation(
                sections[last_para_idx]["text"], para
            ):
                sections[last_para_idx]["text"] += " " + para
                continue  # merged — don't add a new entry
            sec_type, confidence = "paragraph", 0.5
            signee_candidates.append(len(sections))

        entry: dict = {"type": sec_type, "text": para, "confidence": confidence}
        if confidence < 0.6:
            entry["low_confidence"] = True
        sections.append(entry)
        if sec_type == "paragraph":
            last_para_idx = len(sections) - 1

    # --- signee_block: trailing short body paragraphs ---
    if signee_candidates:
        signee_start = len(signee_candidates) - 1
        while signee_start > 0:
            prev_idx = signee_candidates[signee_start - 1]
            curr_idx = signee_candidates[signee_start]
            if curr_idx - prev_idx == 1 and len(sections[prev_idx]["text"].splitlines()) <= 3:
                signee_start -= 1
            else:
                break
        indices = signee_candidates[signee_start:]
        if len([l for l in sections[indices[0]]["text"].splitlines() if l.strip()]) <= 5:
            merged_text = "\n".join(sections[i]["text"] for i in indices)
            sections[indices[0]]["type"] = "signee_block"
            sections[indices[0]]["text"] = merged_text
            sections[indices[0]]["confidence"] = 0.75
            sections[indices[0]].pop("low_confidence", None)
            for i in reversed(indices[1:]):
                sections.pop(i)

    # --- merge consecutive same-type blocks ---
    _MERGE_TYPES = {"receiver_block", "signee_block"}
    merged: list[dict] = []
    for sec in sections:
        if merged and merged[-1]["type"] == sec["type"] and sec["type"] in _MERGE_TYPES:
            merged[-1]["text"] = merged[-1]["text"] + "\n" + sec["text"]
            merged[-1]["confidence"] = max(merged[-1]["confidence"], sec["confidence"])
            merged[-1].pop("low_confidence", None)
            if merged[-1]["confidence"] < 0.6:
                merged[-1]["low_confidence"] = True
        else:
            merged.append(dict(sec))
    sections = merged

    sections = _split_ref_date(sections)
    sections = _split_embedded_sub_paras(sections)
    _split_signee_from_last_para(sections)

    logger.debug("doc_importer: detected %d sections from %d paragraphs", len(sections), len(paragraphs))
    return sections


# ---------------------------------------------------------------------------
# Slot-driven DOCX generation (Leave Certificate, Movement Order)
# ---------------------------------------------------------------------------

def _build_leave_cert_para(slots: dict) -> str:
    """Build the body paragraph text for a leave certificate from extracted slots."""
    army_no = slots.get("army_no", "")
    rank = slots.get("rank", "")
    person_name = slots.get("person_name", "")
    unit = slots.get("unit", "")
    att_unit = slots.get("att_unit", "")
    leave_type = slots.get("leave_type", "")
    days = slots.get("days", "")
    from_date = slots.get("from_date", "")
    to_date = slots.get("to_date", "")
    vill = slots.get("leave_vill", "")
    teh = slots.get("leave_teh", "")
    dist = slots.get("leave_dist", "")
    state = slots.get("leave_state", "")
    pin = slots.get("leave_pin", "")
    contact = slots.get("contact_no", "")
    prefix_date = slots.get("prefix_date", "")
    suffix_date = slots.get("suffix_date", "")

    att_part = f" att with {att_unit}" if att_unit else ""

    if prefix_date and suffix_date:
        prefix_suffix = (
            f" with permission to prefix on {prefix_date}"
            f" and suffix on {suffix_date} being Holiday/ Sunday"
        )
    elif prefix_date:
        prefix_suffix = f" with permission to prefix on {prefix_date}"
    elif suffix_date:
        prefix_suffix = f" with permission to suffix on {suffix_date} being Holiday/ Sunday"
    else:
        prefix_suffix = ""

    addr_parts = []
    if vill:
        addr_parts.append(f"village {vill}")
    if teh:
        addr_parts.append(f"Tehsil {teh}")
    if dist:
        addr_parts.append(f"District {dist}")
    if state:
        addr_parts.append(f"State {state}")
    addr = ", ".join(addr_parts)
    if pin:
        addr += f" - {pin}"

    para1 = (
        f"1.\t\t\t\tNo. {army_no} Rank {rank} Name {person_name} of {unit}{att_part}"
        f" is hereby granted {leave_type} Leave for {days} days"
        f" with effect from {from_date} to {to_date}.{prefix_suffix}"
    )
    if addr:
        para1 += f" The individual is permitted to proceed to {addr}."

    paras = [para1]
    if contact:
        paras.append(f"2.\t\t\t\tContact No: {contact}")
    return "\n\n".join(paras)


def _build_movement_order_para(slots: dict) -> str:
    """Build the body paragraph text for a movement order from extracted slots."""
    army_no = slots.get("army_no", "")
    rank = slots.get("rank", "")
    person_name = slots.get("person_name", "")
    unit = slots.get("unit", "")
    att_unit = slots.get("att_unit", "")
    destination = slots.get("destination", "")
    departure_date = slots.get("departure_date", "")
    departure_time = slots.get("departure_time", "")
    route = slots.get("route", "MR")
    destination_desc = slots.get("destination_desc", "Known to the indl")
    remarks = slots.get("remarks", "Proceeding on temp duty")

    att_part = f", att with {att_unit}" if att_unit else ""
    time_part = f" ({departure_time})" if departure_time else ""

    para1 = (
        f"1.\t\t\t\t{rank} {person_name}, No {army_no}, of {unit}{att_part}"
        f" will proceed on temporary duty to {destination}"
        f" on {departure_date}{time_part}."
    )
    para2 = f"2.\t\t\t\t{destination_desc}"
    para3 = f"3.\t\t\t\t{remarks}. The journey will be performed by {route}."
    return "\n\n".join([para1, para2, para3])


def generate_slot_docx(doc_data: dict, doc_type: str, title: str, out_path: str) -> bool:
    """For slot-driven document types (LEAVE_CERTIFICATE, MOVEMENT_ORDER), generate a
    DOCX from extracted _slots directly instead of blueprint paragraph placeholders.

    Returns True if handled, False if doc_type is not slot-driven.
    """
    doc_type_up = (doc_type or "").upper()
    if doc_type_up not in ("LEAVE_CERTIFICATE", "MOVEMENT_ORDER"):
        return False

    slots = doc_data.get("_slots") or {}

    if doc_type_up == "LEAVE_CERTIFICATE":
        para_text = _build_leave_cert_para(slots)
    else:
        para_text = _build_movement_order_para(slots)

    # Rebuild sections: keep all non-paragraph sections from doc_data, replace paragraph
    from app.services.render_adapter import _section_text
    render_secs: list[dict] = []
    for sec in (doc_data.get("sections") or []):
        if not isinstance(sec, dict):
            continue
        sec_type = sec.get("type", "")
        if sec_type == "paragraph":
            continue  # will add slot-rendered paragraph below
        _rt = (sec.get("content") or {}).get("richtext") or {}
        _rs = _rt.get("state") if isinstance(_rt.get("state"), dict) else None
        render_secs.append({
            "type": sec_type,
            "text": _section_text(sec),
            "alignment": sec.get("alignment"),
            "richtext_state": _rs,
        })

    # Add slot-rendered paragraph
    render_secs.append({"type": "paragraph", "text": para_text})

    # Sort by canonical order
    render_secs.sort(key=lambda s: _SECTION_RENDER_ORDER.get(s["type"], 7))

    import os
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    generate_plain_docx(render_secs, title or "Document", out_path)
    return True


# ---------------------------------------------------------------------------
# Rich Lexical body extractor — preserves DOCX run formatting
# ---------------------------------------------------------------------------

def extract_rich_lexical_body(
    raw_paras: list,
    skip_titles: set,
    skip_sig: set,
    font_family: str = "Times New Roman",
    font_size: int = 12,
) -> "tuple[dict | None, dict | None]":
    """Build Lexical root states from python-docx paragraph objects.

    Preserves bold / italic / underline from DOCX runs AND paragraph alignment
    (center / right / left) from each DOCX paragraph.

    Grouping rules:
      - Paragraphs starting with a digit (numbered paras) or 'Station'/'Dated'
        → start a NEW Lexical paragraph (paragraph spacing).
      - Paragraphs with a different alignment than the current group
        → start a NEW Lexical paragraph (alignment break).
      - Everything else → joined with a linebreak node (no extra spacing).
      - Stops paragraph collection at a 'Distr' line; collects remaining
        non-empty lines as the distribution_list Lexical state.

    Returns (paragraph_lx, distribution_lx) — either can be None.
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    except ImportError:
        _WD_ALIGN = None

    _style = (
        f"font-family:{font_family};"
        f"font-size:{font_size}pt;"
        f"color:#000000;"
    )

    _PARA_START    = re.compile(r"^\d+[\.\t\)]")
    _SECTION_START = re.compile(
        r"^(?:Station|Dated|Place)\s*[:\.]?\s", re.IGNORECASE
    )

    # Lexical format bitmask constants
    _BOLD      = 1
    _ITALIC    = 2
    _UNDERLINE = 8

    def _get_align(para) -> str:
        if _WD_ALIGN is None:
            return ""
        try:
            a = para.alignment
            if a == _WD_ALIGN.CENTER:  return "center"
            if a == _WD_ALIGN.RIGHT:   return "right"
            if a == _WD_ALIGN.JUSTIFY: return "justify"
        except Exception:
            pass
        return ""

    def _run_node(text: str, fmt: int) -> dict:
        return {
            "type": "text", "version": 1,
            "text": text, "format": fmt,
            "detail": 0, "mode": "normal", "style": _style,
        }

    def _linebreak() -> dict:
        return {"type": "linebreak", "version": 1}

    def _para_children(para) -> list[dict]:
        children = []
        for run in para.runs:
            if not run.text:
                continue
            fmt = 0
            if run.bold:
                fmt |= _BOLD
            if run.italic:
                fmt |= _ITALIC
            if run.underline:
                fmt |= _UNDERLINE
            children.append(_run_node(run.text, fmt))
        if not children:
            children = [_run_node("", 0)]
        return children

    def _flush(children: list, align: str = "") -> dict:
        return {
            "type": "paragraph", "version": 1,
            "format": align, "indent": 0, "direction": "ltr",
            "children": children,
        }

    lexical_paras: list[dict] = []
    current: list[dict] = []
    current_align: str = ""
    distr_lines: list[str] = []
    distr_label: str = ""
    in_distr = False

    for dp in raw_paras:
        text = dp.text.strip()
        if not text:
            continue
        if not in_distr and (text.upper() in skip_titles or text in skip_sig):
            continue
        if not in_distr and re.match(r'^Distr\s*[\t ]*[:\.]?', text, re.IGNORECASE):
            in_distr = True
            distr_label = text   # keep "Distr :" label for bold heading
            continue

        if in_distr:
            if text not in skip_sig:
                distr_lines.append(text)
            continue

        align = _get_align(dp)
        children = _para_children(dp)

        if current:
            # Start a new Lexical paragraph on: numbered para, Station/Dated, or alignment change
            if _PARA_START.match(text) or _SECTION_START.match(text) or align != current_align:
                lexical_paras.append(_flush(current, current_align))
                current = children[:]
                current_align = align
            else:
                current.append(_linebreak())
                current.extend(children)
        else:
            current = children[:]
            current_align = align

    if current:
        lexical_paras.append(_flush(current, current_align))

    para_lx = (
        {"root": {"type": "root", "version": 1, "children": lexical_paras}}
        if lexical_paras else None
    )

    # Build distribution_list Lexical state from collected lines
    distr_lx: "dict | None" = None
    if distr_label or distr_lines:
        distr_paras = []
        # First paragraph: bold "Distr :" label
        if distr_label:
            distr_paras.append(_flush([_run_node(distr_label, _BOLD)], ""))
        # Second paragraph: recipient lines joined with linebreaks
        if distr_lines:
            rec_children = []
            for i, line in enumerate(distr_lines):
                if i > 0:
                    rec_children.append(_linebreak())
                rec_children.append(_run_node(line, 0))
            distr_paras.append(_flush(rec_children, ""))
        distr_lx = {
            "root": {
                "type": "root", "version": 1,
                "children": distr_paras,
            }
        }

    return para_lx, distr_lx
